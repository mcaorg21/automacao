#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
| file: /sites/lancamento-prepostos/main.py

| projeto: automacao-python
| data: 2026-06-17
| autor: GitHub Copilot
"""

import json
import os
import base64
import pdb
import re
import sys
import shutil
import time
import unicodedata
import importlib.util
import requests as _requests
import xlrd
from docx import Document
from docx2pdf import convert
from xlutils.copy import copy as xl_copy
from PyPDF2 import PdfMerger
from datetime import datetime, timedelta
from pprint import pprint
from urllib.parse import urlparse
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

#from cpj_app.app_functions_part1 import NUMERO_RECIBO

sys.path.append('C:\\www\\automacao')

from cpj_api import (
    set_api_credentials,
    api_login,
    api_logout,
    api_buscar_processo_tarefa,
    api_buscar_processo_tarefa_filter,
    api_buscar_processo_tarefa_por_data,
    api_buscar_processo_por_pj,
    api_buscar_documentos_pj,
    api_baixar_documento,
    api_buscar_processo_por_ficha,
    api_atualizar_tarefa,
    api_atualizar_processo,
    api_buscar_lancamentos_bmg
)

BASE_PATH = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(BASE_PATH, 'config.json')
CONFIG_DAYCOVAL_PATH = os.path.join(BASE_PATH, 'config_daycoval.json')
CONFIG_IFOOD_PATH = os.path.join(BASE_PATH, 'config_ifood.json')
CONFIG_BMG_PATH = os.path.join(BASE_PATH, 'config_bmg.json')
DOWNLOADS_PATH = os.path.join(BASE_PATH, 'downloads')
COOKIES_JSON_PATH = os.path.join(BASE_PATH, 'cookies.json')
COOKIES_JSON_IFOOD_PATH = os.path.join(BASE_PATH, 'cookies_ifood.json')
PLANILHA_ORIGINAL_PREPOSTO_PATH = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'planilha_modelo_preposto_original.xls')
PLANILHA_MODELO_PATH = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'planilha_modelo_preposto.xls')
DESTINO_PLANILHA_BMG = PLANILHA_MODELO_PATH
CPJ_REEMBOLSO_BMG_PATH = os.path.abspath(r'C:\www\automacao\sites\cpj-reembolso-bmg\main.py')
RECIBO_PADRAO_ORIGINAL = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'recibo_padrao_original.docx')
DESTINO_RECIBO_PADRAO_DOCX = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'recibo_padrao.docx')
DESTINO_RECIBO_PADRAO_PDF = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'recibo_padrao.pdf')
DESTINO_RECIBO_PADRAO_BMG_PDF = DESTINO_RECIBO_PADRAO_PDF
DOCX_MODELO_BMG = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'descritivo_padrao_original.docx')
DESTINO_DESCRITIVO_PADRAO_DOCX = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'descritivo_padrao.docx')
DESTINO_DESCRITIVO_PADRAO_PDF = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'descritivo_padrao.pdf')
DESTINO_DESCRITIVO_RECIBO_PADRAO_FINAL_PDF = os.path.join(BASE_PATH, 'documentos_padrao', 'banco_bmg', 'descritivo_recibo_padrao_final.pdf')
DESTINO_DESCRITIVO_RECIBO_PADRAO_FINAL_BMG_PDF = DESTINO_DESCRITIVO_RECIBO_PADRAO_FINAL_PDF
RECIBO_PADRAO_IFOOD_DOCX = os.path.join(BASE_PATH, 'documentos_padrao', 'ifood', 'recibo_padrao.docx')
RECIBO_PADRAO_IFOOD_COPY_DOCX = os.path.join(BASE_PATH, 'documentos_padrao', 'ifood', 'recibo_padrao_copy.docx')
RECIBO_PADRAO_IFOOD_COPY_PDF = os.path.join(BASE_PATH, 'documentos_padrao', 'ifood', 'recibo_padrao_copy.pdf')
CHROME_PROFILE_ROOT = os.path.join(BASE_PATH, 'chrome_profile')
CHROME_USER_DATA_DIR = os.path.join(CHROME_PROFILE_ROOT, 'User Data')
CHROME_PROFILE_DIRECTORY = 'Default'


def carregar_funcoes_cpj_reembolso_bmg():
    """Importa dinamicamente funções do módulo cpj-reembolso-bmg/main.py."""
    if not os.path.exists(CPJ_REEMBOLSO_BMG_PATH):
        raise FileNotFoundError(f'Módulo cpj-reembolso-bmg não encontrado: {CPJ_REEMBOLSO_BMG_PATH}')

    spec = importlib.util.spec_from_file_location('cpj_reembolso_bmg_main', CPJ_REEMBOLSO_BMG_PATH)
    module = importlib.util.module_from_spec(spec)
    loader = spec.loader
    if loader is None:
        raise ImportError(f'Não foi possível carregar o módulo: {CPJ_REEMBOLSO_BMG_PATH}')

    old_argv = sys.argv
    hoje = datetime.now().strftime('%d/%m/%Y')
    sys.argv = [old_argv[0] if old_argv else 'cpj_reembolso_bmg', '1', hoje, hoje]
    try:
        loader.exec_module(module)
    finally:
        sys.argv = old_argv

    return module


cpj_reembolso_bmg = carregar_funcoes_cpj_reembolso_bmg()
login_web_exyon_bmg = cpj_reembolso_bmg.login_web_exyon_bmg
selecionar_unidade = cpj_reembolso_bmg.selecionar_unidade
executar_script_menu = cpj_reembolso_bmg.executar_script_menu
executar_script_menu_lateral = cpj_reembolso_bmg.executar_script_menu_lateral
executar_script_classificacao = cpj_reembolso_bmg.executar_script_classificacao
executar_preenchimento_formulario = cpj_reembolso_bmg.executar_preenchimento_formulario
anexar_pdfs_formulario = cpj_reembolso_bmg.anexar_pdfs_formulario
buscar_processo_alternativo = cpj_reembolso_bmg.buscar_processo_alternativo
buscar_quantidade_processos = cpj_reembolso_bmg.buscar_quantidade_processos
finalizar_processo = cpj_reembolso_bmg.finalizar_processo
verificar_lancamentos = cpj_reembolso_bmg.verificar_lancamentos

def formatar_numero_processo_cnj(valor: object) -> str:
    """Formata um número de processo para o padrão CNJ com pontos e hífen."""
    if valor is None:
        return ''

    texto = str(valor).strip()
    if not texto:
        return ''

    digits = ''.join(ch for ch in texto if ch.isdigit())
    if not digits:
        return texto

    if len(digits) < 20:
        digits = digits.zfill(20)
    elif len(digits) > 20:
        digits = digits[:20]

    return f'{digits[:7]}-{digits[7:9]}.{digits[9:13]}.{digits[13:14]}.{digits[14:16]}.{digits[16:20]}'


API_BASE_URL = 'https://app.leviatan.com.br/dcncadv/cpj/agnes'
API_LOGIN = 'api'
API_PASSWORD = '2025'
DAYCOVAL_URL_CONSULTA = 'https://spjw.daycoval.com.br:8282/Form/Processo/ProcessoLeve.aspx?status=64'
IFOOD_URL_CONSULTA = r'https://auth.ifoodcorp.com.br/realms/corporate/protocol/openid-connect/auth?scope=openid&state=2QnmLZh-bRNnoIiDJzR9FafKQ-PK5f0gl3Zu6zLvdkY.nfw9MnIVOXE.TzLgDjfZQ9q_BrOtWpvUxw.eyJydSI6Imh0dHBzOi8vaWZvb2QuYmNsZWdhbC5pby9zaWduaW4tb2lkYyIsInJ0IjoiY29kZSIsInJtIjoiZm9ybV9wb3N0Iiwic3QiOiJDZkRKOE1YWjRlcnhvR1ZDdTJVaWRna0NvOGFBN1Y3TktrNHE4N2Rqekw0dnBLU01RTU9PT1B4azhpOVRCY0oyRHZvZGtRUVRtQVRxQnZjb3dhRmZfWWRxdW9PR0ZJSl92MjNrMnNoYUZtRGxkVHZ1NlpxVU82amZwelVYZ1BjSk53aTBRbzdBZjFUMWJGT2VicXhWZEJQdFFyZkU5TnNUNnNhaGk5dF9NS3hhN3Bjek91cXQ1dk9RQlFrM0E2bFA4Wi03T0pTSkN0VDJneGRRQ0lMbVFTMjFFY1kwRnRWNmxFUFB0NHVnZmU4ZHhMenhtUHRaeUU4RHlmT0EwMGVILVNCSk1xQ0tqd3p2NmExX1RncWw1aXpRdTkzeEppVGNVYjhJdnB1OTd1NXJGRFdfblVFb0xnNlJpTmJIejBEaXR2THJobjBkZTRNSHJyX2xSWURJcWJxVlNFQ1FxTkk2Q3RpSFRXODE5bmlYaEdxanlKbEJzMGVMTGNWVkdEVjdWbzdWdFEifQ&response_type=code&client_id=bc-legal&redirect_uri=https%3A%2F%2Fsso.bclegal.io%2Frealms%2FiFood%2Fbroker%2Fifoodlover%2Fendpoint&nonce=MXGMhtZos_upkuIAtrtv5w'


def open_chrome_browser():
    """Abre o navegador Chrome usando Selenium (padrão omni-pde-fsp-trc)."""
    try:
        print('\nAbrindo navegador Chrome...')

        chrome_options = Options()
        # chrome_options.add_argument('--headless')  # Descomente para rodar sem interface gráfica
        chrome_options.add_argument('--window-size=1200,800')
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')
        chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])
        chrome_options.add_argument('--ignore-ssl-errors')
        chrome_options.add_argument('--ignore-certificate-errors')
        chrome_options.add_argument('--no-first-run')
        chrome_options.add_argument('--no-default-browser-check')

        os.makedirs(CHROME_USER_DATA_DIR, exist_ok=True)
        os.makedirs(os.path.join(CHROME_USER_DATA_DIR, CHROME_PROFILE_DIRECTORY), exist_ok=True)
        chrome_options.add_argument(f'--user-data-dir={CHROME_USER_DATA_DIR}')
        chrome_options.add_argument(f'--profile-directory={CHROME_PROFILE_DIRECTORY}')

        try:
            print('Tentando usar ChromeDriver do PATH do sistema...')
            driver = webdriver.Chrome(options=chrome_options)
            print('✓ Chrome aberto com sucesso usando ChromeDriver do PATH!')
            return driver
        except Exception as path_error:
            print(f'⚠ ChromeDriver não encontrado no PATH: {path_error}')

        try:
            print('Tentando instalar ChromeDriver automaticamente com webdriver-manager...')

            cache_path = os.path.join(os.path.expanduser('~'), '.wdm')
            if os.path.exists(cache_path):
                try:
                    shutil.rmtree(cache_path)
                    print('✓ Cache limpo, baixando ChromeDriver novamente...')
                except Exception as cache_error:
                    print(f'⚠ Erro ao limpar cache: {cache_error}')

            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
            print('✓ Chrome aberto com sucesso usando webdriver-manager!')
            return driver

        except Exception as manager_error:
            print(f'✗ Erro com webdriver-manager: {manager_error}')

        raise Exception('Não foi possível inicializar o ChromeDriver')

    except Exception as e:
        print(f'\n✗ ERRO CRÍTICO: Não foi possível abrir o Chrome')
        print(f'Erro: {e}')
        raise


def salvar_cookies(driver, cookies_path: str = COOKIES_JSON_PATH):
    """Salva cookies atuais do navegador em arquivo JSON."""
    try:
        cookies = driver.get_cookies()
        with open(cookies_path, 'w', encoding='utf-8') as f:
            json.dump(cookies, f, ensure_ascii=False, indent=2)
        print(f'✓ {len(cookies)} cookie(s) salvos em: {cookies_path}')
        return True
    except Exception as e:
        print(f'⚠ Erro ao salvar cookies: {e}')
        return False


def salvar_cookies_ifood(driver, cookies_path: str = COOKIES_JSON_IFOOD_PATH):
    """Salva cookies do iFood no arquivo cookies_ifood.json."""
    return salvar_cookies(driver, cookies_path)


def carregar_cookies(driver, url_base: str, cookies_path: str = COOKIES_JSON_PATH):
    """Carrega cookies salvos no navegador para o domínio informado."""
    try:
        if not os.path.exists(cookies_path):
            print(f'⚠ Arquivo de cookies não encontrado: {cookies_path}')
            return False

        with open(cookies_path, 'r', encoding='utf-8') as f:
            cookies = json.load(f)

        print(f'\nCarregando {len(cookies)} cookie(s) de {cookies_path}...')

        driver.get(url_base)
        time.sleep(1)

        try:
            driver.execute_cdp_cmd('Network.enable', {})
        except Exception:
            pass

        adicionados = 0
        for cookie in cookies:
            cookie_name = cookie.get('name', 'sem_nome')
            cookie_value = cookie.get('value', '')
            cookie_path = cookie.get('path') or '/'
            cookie_secure = bool(cookie.get('secure', False))
            cookie_http_only = bool(cookie.get('httpOnly', False))
            cookie_domain = cookie.get('domain')
            host_only = bool(cookie.get('hostOnly', False))

            payload = {
                'url': url_base,
                'name': cookie_name,
                'value': cookie_value,
                'path': cookie_path,
                'secure': cookie_secure,
                'httpOnly': cookie_http_only,
            }

            if cookie_domain and not host_only:
                payload['domain'] = cookie_domain

            if 'expirationDate' in cookie:
                payload['expires'] = int(cookie['expirationDate'])

            samesite_map = {
                'no_restriction': 'None',
                'none': 'None',
                'lax': 'Lax',
                'strict': 'Strict',
            }
            same_site = samesite_map.get(str(cookie.get('sameSite', '')).lower())
            if same_site:
                payload['sameSite'] = same_site

            try:
                result = driver.execute_cdp_cmd('Network.setCookie', payload)
                if result.get('success', False):
                    adicionados += 1
                    print(f'  ✓ Cookie injetado via CDP: {cookie_name}')
                    continue
            except Exception as cdp_error:
                print(f'  ⚠ CDP falhou para {cookie_name}: {cdp_error}')

            try:
                selenium_cookie = {
                    'name': cookie_name,
                    'value': cookie_value,
                    'path': cookie_path,
                    'secure': cookie_secure,
                }
                if cookie_domain and not host_only:
                    selenium_cookie['domain'] = cookie_domain
                if 'expirationDate' in cookie:
                    selenium_cookie['expiry'] = int(cookie['expirationDate'])
                if same_site:
                    selenium_cookie['sameSite'] = same_site
                driver.add_cookie(selenium_cookie)
                adicionados += 1
                print(f'  ✓ Cookie injetado via add_cookie: {cookie_name}')
            except Exception as cookie_error:
                print(f'  ⚠ Cookie ignorado ({cookie_name}): {cookie_error}')
        
        print(f'✓ {adicionados}/{len(cookies)} cookie(s) carregado(s) com sucesso!')
        return adicionados > 0

    except Exception as e:
        print(f'✗ Erro ao carregar cookies: {e}')
        return False


def _set_cookie_via_cdp(driver, url_base: str, cookie: dict) -> bool:
    """Tenta definir o cookie via Chrome DevTools Protocol quando add_cookie falha."""
    try:
        samesite_map = {
            'no_restriction': 'None',
            'lax': 'Lax',
            'strict': 'Strict',
        }
        payload = {
            'url': url_base,
            'name': cookie['name'],
            'value': cookie['value'],
            'path': cookie.get('path', '/'),
            'secure': cookie.get('secure', False),
            'httpOnly': cookie.get('httpOnly', False),
        }
        if cookie.get('domain'):
            payload['domain'] = cookie['domain']
        if 'expirationDate' in cookie:
            payload['expires'] = int(cookie['expirationDate'])
        same_site = samesite_map.get(cookie.get('sameSite', '').lower())
        if same_site:
            payload['sameSite'] = same_site

        result = driver.execute_cdp_cmd('Network.setCookie', payload)
        return result.get('success', False)
    except Exception:
        return False


def carregar_cookies_ifood(driver, url_base: str = 'https://ifood.bclegal.io/', cookies_path: str = COOKIES_JSON_IFOOD_PATH):
    """Carrega cookies do iFood no navegador usando cookies_ifood.json."""
    return carregar_cookies(driver, url_base=url_base, cookies_path=cookies_path)


def aguardar_loading_ifood(driver, timeout: int = 20) -> None:
    """Aguarda o overlay de loading do iFood deixar de estar visivel."""
    def loading_invisivel(_driver):
        try:
            loading = _driver.find_element(By.ID, 'loading')
        except Exception:
            return True

        return _driver.execute_script(
            """
            const element = arguments[0];
            if (!element) {
                return true;
            }

            const style = window.getComputedStyle(element);
            return style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0';
            """,
            loading,
        )

    WebDriverWait(driver, timeout).until(loading_invisivel)


def carregar_config_bmg() -> dict:
    """Carrega as configurações de BMG do arquivo config_bmg.json."""
    try:
        if not os.path.exists(CONFIG_BMG_PATH):
            print(f'⚠ Arquivo de configuração BMG não encontrado: {CONFIG_BMG_PATH}')
            return {}

        with open(CONFIG_BMG_PATH, 'r', encoding='utf-8') as f:
            config_bmg = json.load(f)

        print(f'✓ Configuração BMG carregada de: {CONFIG_BMG_PATH}')
        return config_bmg if isinstance(config_bmg, dict) else {}
    except Exception as e:
        print(f'✗ Falha ao ler config_bmg.json: {e}')
        return {}


def salvar_config_bmg(config_bmg: dict) -> None:
    """Salva as configurações de BMG no arquivo config_bmg.json."""
    try:
        with open(CONFIG_BMG_PATH, 'w', encoding='utf-8') as f:
            json.dump(config_bmg, f, ensure_ascii=False, indent=4)
        print(f'✓ Configuração BMG salva em: {CONFIG_BMG_PATH}')
    except Exception as e:
        print(f'✗ Falha ao salvar config_bmg.json: {e}')


def zerar_config_bmg() -> None:
    """Zera os valores de config_bmg.json."""
    config_bmg = {
        'numero_recibo': '',
        'data_inicial': '',
        'data_final': '',
        'iniciado_em': '',
    }
    salvar_config_bmg(config_bmg)
    print('✓ Values de config_bmg.json zerados')


def abrir_chrome_e_acessar_url(url_destino: str, url_base_cookies: str = None):
    """Abre o Chrome, tenta reutilizar cookies e acessa a URL desejada."""
    driver = open_chrome_browser()

    if url_base_cookies:
        carregar_cookies(driver, url_base=url_base_cookies)

    driver.get(url_destino)
    salvar_cookies(driver)

    return driver


def login_daycoval(driver, usuario: str, senha: str, timeout: int = 20) -> bool:
    """Realiza login no Daycoval e aguarda elemento que indica sessão autenticada."""
    try:
        
        wait = WebDriverWait(driver, timeout)

        driver.delete_all_cookies()

        campo_usuario = wait.until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_tbarMain_pgcMain_hucLogin_txtUser_I"]'))
        )
        campo_usuario.clear()
        campo_usuario.send_keys(usuario)

        campo_senha = wait.until(
            EC.presence_of_element_located((By.ID, 'ctl00_tbarMain_pgcMain_hucLogin_txtPass_I'))
        )
        campo_senha.clear()
        campo_senha.send_keys(senha)

        botao_entrar = wait.until(
            EC.element_to_be_clickable((By.ID, 'ctl00_tbarMain_pgcMain_hucLogin_btnLogin_CD'))
        )
        botao_entrar.click()

        wait.until(
            EC.presence_of_element_located((By.ID, 'ctl00_tbarMain_pgcMain_arpProcesso1_HC'))
        )

        salvar_cookies(driver)
        print('    ✓ Login Daycoval realizado com sucesso!')
        return True

    except Exception as e:
        print(f'    ✗ Falha no login Daycoval: {e}')
        return False

def login_ifood(driver, usuario: str, senha: str, timeout: int = 20) -> bool:
    """Realiza login no iFood e aguarda elemento que indica sessão autenticada."""
    try:

        wait = WebDriverWait(driver, timeout)
        try:
            cookies_loaded = carregar_cookies_ifood(driver)

            if not cookies_loaded:
                driver.delete_all_cookies()
                print('    ⚠ Nenhum cookie do iFood foi injetado; prosseguindo com login manual.')
            else:
                driver.get('https://ifood.bclegal.io/')
                try:

                    menu = wait.until(
                        EC.element_to_be_clickable((By.CLASS_NAME, 'rz-header'))
                    )

                    driver.get('https://ifood.bclegal.io/ifoodv5/(S())/default.aspx?RedirectV8=solicitacao/redireciona.aspx')

                    try:
                        menu_corporativo = wait.until(
                            EC.element_to_be_clickable((By.ID, 'social-ifoodlover'))
                        )

                        menu_corporativo.click()
                    except: 
                        pass

                    global SESSION_ID
                    SESSION_ID = session_id = driver.execute_script(""" return getSessionId();""")

                    time.sleep(3)  # Aguarda a página carregar

                    driver.switch_to.default_content()
                    WebDriverWait(driver, 20).until(
                        EC.frame_to_be_available_and_switch_to_it((By.ID, 'mainFrame'))
                    )

                    driver.execute_script("""__doPostBack('ctl00$ContentPlaceHolder1$rptMenu$ctl35$lnkCriarNovaSolicitacao','')""")  # Rola até o final da página

                    global URL_CADASTRO
                    URL_CADASTRO = f'https://ifood.bclegal.io/ifoodv5/(S({session_id}))/solicitacao/reembolso/Cadastro.aspx?nwf=sim&nivel='
                    driver.get(URL_CADASTRO)

                    if driver.page_source.find('Reembolso') != -1:
                        print('    ✓ Sessão iFood autenticada via cookies!')
                        return True   
                    
                    driver.get('https://ifood.bclegal.io/')

                    driver.delete_all_cookies() 

                except:
                    pass
        
        except Exception as e:
            driver.delete_all_cookies()
            print(f'    ⚠ Falha ao carregar cookies do iFood: {e}')

        botao_corporativo = wait.until(
            EC.element_to_be_clickable((By.ID, 'social-ifoodlover'))
        )
        botao_corporativo.click()

        driver.get('https://ifood.bclegal.io/')

        botao_corporativo = wait.until(
            EC.element_to_be_clickable((By.ID, 'social-ifoodlover'))
        )
        botao_corporativo.click()

        campo_usuario = wait.until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="username"]'))
        )
        campo_usuario.clear()
        campo_usuario.send_keys(usuario)

        campo_senha = wait.until(
            EC.presence_of_element_located((By.ID, 'password'))
        )
        campo_senha.clear()
        campo_senha.send_keys(senha)

        botao_remember = wait.until(
            EC.element_to_be_clickable((By.ID, 'rememberMe'))
        )
        botao_remember.click()
        time.sleep(1)
        botao_remember.click()

        pdb.set_trace()
        botao_entrar = wait.until(
            EC.element_to_be_clickable((By.ID, 'kc-login'))
        )

        botao_entrar.click()

        
        try:

            campo_codigo_unico = wait.until(
                EC.presence_of_element_located((By.ID, 'smsCode'))
            )

            campo_codigo_unico.clear()
            codigo_unico = input("Informe o código único:")
            campo_codigo_unico.send_keys(codigo_unico)

            botao_entrar = wait.until(
                EC.element_to_be_clickable((By.ID, 'sendSmsCodeBtn'))
            )

            botao_entrar.click()

            try:
                menu = wait.until(
                    EC.element_to_be_clickable((By.CLASS_NAME, 'rz-header'))
                )

                if 'ESCRITÓRIO DIAS COSTA SOCIEDADE DE ADVOGADOS' in menu.text:
                    return True
            
            except:
                return False
                

        except Exception as e:
            print(f'    ⚠ Fluxo de código único não foi necessário ou não foi encontrado: {e}')

        print('    ✓ Login Ifood realizado com sucesso!')
        return True

    except Exception as e:
        print(f'    ✗ Falha no login Ifood: {e}')
        return False


def configurar_api() -> None:
    set_api_credentials(
        base_url=API_BASE_URL,
        login=API_LOGIN,
        password=API_PASSWORD,
    )


def normalizar_texto_busca(texto: str) -> str:
    if not texto:
        return ''
    texto = unicodedata.normalize('NFKD', texto)
    texto = texto.encode('ascii', 'ignore').decode('ascii')
    return texto.lower().strip()


def formatar_data_hora_tarefa(valor) -> str:
    """Formata data/hora da tarefa aceitando datetime ou string ISO."""
    if isinstance(valor, datetime):
        return valor.strftime(' %Y-%m-%d %H:%M:%S')

    if isinstance(valor, str):
        valor_limpo = valor.strip()
        if valor_limpo.endswith('Z'):
            valor_limpo = valor_limpo[:-1] + '+00:00'

        try:
            return datetime.fromisoformat(valor_limpo).strftime(' %d/%m/%Y %H:%M:%S')
        except Exception:
            return f' {valor_limpo}'

    return ''


def parse_data_para_datetime(data_str: str):
    """Converte uma string de data para datetime, suportando vários formatos."""
    if not data_str:
        return None

    data_str = str(data_str).strip()

    try:
        if '/' in data_str:
            return datetime.strptime(data_str, '%d/%m/%Y')
        if 'T' in data_str:
            return datetime.fromisoformat(data_str.replace('Z', '+00:00'))
        return datetime.strptime(data_str, '%Y-%m-%d')
    except Exception:
        try:
            return datetime.fromisoformat(data_str)
        except Exception:
            return None


def formatar_data_lancamento(data_str) -> str:
    """Formata data de lançamento para DD/MM/YYYY."""
    if isinstance(data_str, datetime):
        return data_str.strftime('%d/%m/%Y')

    try:
        data = parse_data_para_datetime(data_str)
        if data:
            return data.strftime('%d/%m/%Y')
    except Exception:
        pass

    return str(data_str or '')


def atualizar_planilha_preposto(lancamentos_processados: list, destino_path: str) -> None:
    """Copia o template original e preenche a nova planilha sem alterar a formatação."""
    if not lancamentos_processados:
        print('⚠ Nenhum lançamento processado para escrever na planilha.')
        return

    try:
        if not os.path.exists(PLANILHA_ORIGINAL_PREPOSTO_PATH):
            raise FileNotFoundError(f'Template original não encontrado: {PLANILHA_ORIGINAL_PREPOSTO_PATH}')

        workbook_template = xlrd.open_workbook(PLANILHA_ORIGINAL_PREPOSTO_PATH, formatting_info=True)
        workbook = xl_copy(workbook_template)
        sheet = workbook.get_sheet(0)

        headers = ['Cod_Causa', 'Processo', 'Tipo_de_Lancamento', 'Valor', 'Data_Ocorrencia']
        for col_idx, header in enumerate(headers):
            sheet.write(0, col_idx, header)

        for row_idx, lancamento in enumerate(lancamentos_processados, start=1):
            sheet.write(row_idx, 0, lancamento.get('numero_integracao', ''))
            sheet.write(row_idx, 1, lancamento.get('numero_processo_pesq', ''))
            sheet.write(row_idx, 2, 'DESPESA COM PREPOSTO')
            sheet.write(row_idx, 3, '75')
            sheet.write(row_idx, 4, lancamento.get('data_lancamento', ''))

        os.makedirs(os.path.dirname(destino_path), exist_ok=True)
        workbook.save(destino_path)
        print(f'✓ Planilha salva em: {destino_path}')

    except Exception as e:
        print(f'✗ Erro ao atualizar planilha preposto: {e}')


def formatar_data_mes_ano_atual() -> str:
    hoje = datetime.now()
    meses = [
        'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
    ]
    dia = hoje.day
    mes = meses[hoje.month - 1]
    ano = hoje.year
    return f'{dia} de {mes} de {ano}'


def atualizar_recibo_padrao(numero_recibo: str, destino_path_docx: str = None, destino_path_pdf: str = None) -> None:
    if destino_path_docx is None:
        destino_path_docx = DESTINO_RECIBO_PADRAO_DOCX
    if destino_path_pdf is None:
        destino_path_pdf = DESTINO_RECIBO_PADRAO_PDF

    try:
        if not os.path.exists(RECIBO_PADRAO_ORIGINAL):
            print(f'✗ Recibo padrão original não encontrado: {RECIBO_PADRAO_ORIGINAL}')
            return

        os.makedirs(os.path.dirname(destino_path_docx), exist_ok=True)
        shutil.copy2(RECIBO_PADRAO_ORIGINAL, destino_path_docx)

        document = Document(destino_path_docx)
        data_mes_ano = formatar_data_mes_ano_atual()

        for paragraph in document.paragraphs:
            if '|TITULO_INSERIDO|' in paragraph.text or '|DATA_MES_DO_ANO|' in paragraph.text:
                texto = paragraph.text
                texto = texto.replace('|TITULO_INSERIDO|', str(numero_recibo))
                texto = texto.replace('|DATA_MES_DO_ANO|', data_mes_ano)
                paragraph.text = texto

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '|TITULO_INSERIDO|' in cell.text or '|DATA_MES_DO_ANO|' in cell.text:
                        texto = cell.text
                        texto = texto.replace('|TITULO_INSERIDO|', str(numero_recibo))
                        texto = texto.replace('|DATA_MES_DO_ANO|', data_mes_ano)
                        cell.text = texto

        document.save(destino_path_docx)
        print(f'✓ Recibo DOCX atualizado em: {destino_path_docx}')

        try:
            convert(destino_path_docx, destino_path_pdf)
            print(f'✓ Recibo PDF gerado em: {destino_path_pdf}')
        except Exception as e:
            print(f'✗ Erro ao converter recibo para PDF: {e}')

    except Exception as e:
        print(f'✗ Erro ao atualizar recibo padrão: {e}')


def _replace_docx_text(document, search: str, replace: str) -> None:
    for paragraph in document.paragraphs:
        if search in paragraph.text:
            for run in paragraph.runs:
                if search in run.text:
                    run.text = run.text.replace(search, replace)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search in paragraph.text:
                        for run in paragraph.runs:
                            if search in run.text:
                                run.text = run.text.replace(search, replace)


def formatar_data_por_extenso(data: datetime) -> str:
    meses = [
        'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
    ]
    return f'{data.day} de {meses[data.month - 1]} de {data.year}'


def gerar_recibo_ifood(numero_processo: str, valor: str, valor_extenso: str, texto: str, data_por_extenso: str, destino_docx: str = None, destino_pdf: str = None) -> bool:
    if destino_docx is None:
        destino_docx = RECIBO_PADRAO_IFOOD_COPY_DOCX
    if destino_pdf is None:
        destino_pdf = RECIBO_PADRAO_IFOOD_COPY_PDF

    try:
        if not os.path.exists(RECIBO_PADRAO_IFOOD_DOCX):
            print(f'✗ Recibo padrão iFood não encontrado: {RECIBO_PADRAO_IFOOD_DOCX}')
            return False

        shutil.copy2(RECIBO_PADRAO_IFOOD_DOCX, destino_docx)
        document = Document(destino_docx)

        replacements = {
            '{NUMERO_PROCESSO}': numero_processo,
            '{VALOR}': f'{valor},00',
            '{VALOR_POR_EXTENSO}': valor_extenso,
            '{TIPO_PRESENCA}': texto,
            '{DATA_POR_EXTENSO}': data_por_extenso,
        }

        for search, replace in replacements.items():
            _replace_docx_text(document, search, replace)

        document.save(destino_docx)

        try:
            convert(destino_docx, destino_pdf)
            print(f'✓ Recibo iFood gerado em PDF: {destino_pdf}')
            return True
        except Exception as e:
            print(f'✗ Erro ao converter recibo iFood para PDF: {e}')
            return False

    except Exception as e:
        print(f'✗ Erro ao gerar recibo iFood: {e}')
        return False


def concatenar_pdfs(arquivos: list[str], output_path: str) -> bool:
    try:
        merger = PdfMerger()
        for arquivo in arquivos:
            merger.append(arquivo)
        with open(output_path, 'wb') as f_out:
            merger.write(f_out)
        merger.close()
        print(f'✓ PDFs concatenados em: {output_path}')
        return True
    except Exception as e:
        print(f'✗ Erro ao concatenar PDFs: {e}')
        return False


def registrar_tarefa_sem_ata(tarefa: dict, id_processo, caminho_saida: str, motivo: str = 'Erro') -> None:
    """Registra em JSON as tarefas em que nenhuma ATA foi encontrada."""
    try:
        registros_existentes = []
        if os.path.exists(caminho_saida):
            with open(caminho_saida, 'r', encoding='utf-8') as f:
                conteudo_bruto = f.read().strip()
                if not conteudo_bruto:
                    conteudo = []
                else:
                    conteudo = json.loads(conteudo_bruto)
                if isinstance(conteudo, list):
                    registros_existentes = conteudo

        id_tramitacao = tarefa.get('id_tramitacao')
        ag_data_hora = tarefa.get('ag_data_hora')
        evento = tarefa.get('evento')
        update_data_hora = tarefa.get('update_data_hora')

        for registro_existente in registros_existentes:
            if not isinstance(registro_existente, dict):
                continue

            mesmo_processo = registro_existente.get('id_processo') == id_processo
            mesmo_id_tramitacao = (
                id_tramitacao is not None
                and registro_existente.get('id_tramitacao') == id_tramitacao
            )
            mesma_combinacao = (
                id_tramitacao is None
                and registro_existente.get('id_tramitacao') is None
                and
                registro_existente.get('ag_data_hora') == ag_data_hora
                and registro_existente.get('evento') == evento
                and registro_existente.get('update_data_hora') == update_data_hora
            )

            if mesmo_processo and (mesmo_id_tramitacao or mesma_combinacao):
                print(f'    ✓ Tarefa sem ATA já registrada em: {caminho_saida}')
                return

        registro = {
            'id_processo': id_processo,
            'ag_data_hora': ag_data_hora,
            'evento': evento,
            'id_tramitacao': id_tramitacao,
            'update_data_hora': update_data_hora,
            'tarefa': tarefa,
            'registrado_em': datetime.now().isoformat(),
        }
        registros_existentes.append(registro)

        with open(caminho_saida, 'w', encoding='utf-8') as f:
            json.dump(registros_existentes, f, ensure_ascii=False, indent=2, default=str)

        print(f'    ✓ Tarefa sem ATA registrada em: {caminho_saida}')
    except Exception as e:
        print(f'    ⚠ Erro ao registrar tarefa sem ATA: {e}')


def fill_descritivo_pdf(lancamentos_processados: list, numero_recibo: str, valor_somado: str | None = None, destino_path_docx: str = None, destino_path_pdf: str = None) -> None:
    """Preenche o documento descritivo e converte para PDF."""
    if destino_path_docx is None:
        destino_path_docx = DESTINO_DESCRITIVO_PADRAO_DOCX
    if destino_path_pdf is None:
        destino_path_pdf = DESTINO_DESCRITIVO_PADRAO_PDF

    try:
        if not os.path.exists(DOCX_MODELO_BMG):
            print(f'✗ Modelo descritivo original não encontrado: {DOCX_MODELO_BMG}')
            return

        if not lancamentos_processados:
            print('⚠ Nenhum lançamento processado para preencher o descritivo.')
            return

        print('\nPreenchendo documento descritivo...')
        os.makedirs(os.path.dirname(destino_path_docx), exist_ok=True)
        shutil.copy2(DOCX_MODELO_BMG, destino_path_docx)

        document = Document(destino_path_docx)
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        mes_hoje = datetime.now().strftime('%m/%Y')
        if valor_somado is None:
            valor_somado = f'{len(lancamentos_processados) * 75:.2f}'.replace('.', ',')

        print(f'✓ Número recibo: {numero_recibo}')
        print(f'✓ Valor somado: {valor_somado}')
        print(f'✓ Registros a preencher: {len(lancamentos_processados)}')

        for table_idx, table in enumerate(document.tables):
            print(f'Processando tabela {table_idx + 1}...')
            for row in table.rows:
                for cell in row.cells:
                    if '{{NUMERO_RECIBO}}' in cell.text:
                        cell.text = cell.text.replace('{{NUMERO_RECIBO}}', str(numero_recibo))
                    if '{{DATA_HOJE}}' in cell.text:
                        cell.text = cell.text.replace('{{DATA_HOJE}}', data_hoje)
                    if '{{MES_HOJE}}' in cell.text:
                        cell.text = cell.text.replace('{{MES_HOJE}}', mes_hoje)
                    if '{{VALOR_SOMADO}}' in cell.text:
                        cell.text = cell.text.replace('{{VALOR_SOMADO}}', valor_somado)

        if len(document.tables) > 2:
            table = document.tables[2]
            print('Preenchendo registros na tabela 3...')
            for idx, lancamento in enumerate(lancamentos_processados):
                row_idx = idx + 1
                if row_idx >= len(table.rows):
                    table.add_row()
                row = table.rows[row_idx]
                if len(row.cells) >= 5:
                    row.cells[1].text = lancamento.get('data_lancamento', '')
                    row.cells[2].text = lancamento.get('numero_integracao', '')
                    row.cells[3].text = 'DESPESA COM PREPOSTO'
                    row.cells[4].text = '75,00'
                else:
                    print(f'  ⚠ Linha {row_idx} não tem colunas suficientes para preencher o descritivo.')
        else:
            print('⚠ Modelo descritivo não possui tabela 3 para preencher os lançamentos.')

        document.save(destino_path_docx)
        print(f'✓ Documento Word salvo em: {destino_path_docx}')

        print('Convertendo para PDF...')
        convert(destino_path_docx, destino_path_pdf)
        print(f'✓ PDF salvo em: {destino_path_pdf}')
        print('✓ Descritivo preenchido com sucesso!')

    except Exception as e:
        print(f'✗ Erro ao preencher descritivo: {e}')
        import traceback
        traceback.print_exc()
        raise


def merge_descritivo_e_recibo_pdf(descritivo_pdf_path: str = None, recibo_pdf_path: str = None, output_pdf_path: str = None) -> None:
    """Mescla descritivo e recibo em um único PDF com descritivo primeiro."""
    if descritivo_pdf_path is None:
        descritivo_pdf_path = DESTINO_DESCRITIVO_PADRAO_PDF
    if recibo_pdf_path is None:
        recibo_pdf_path = DESTINO_RECIBO_PADRAO_PDF
    if output_pdf_path is None:
        output_pdf_path = DESTINO_DESCRITIVO_RECIBO_PADRAO_FINAL_PDF

    try:
        if not os.path.exists(descritivo_pdf_path):
            print(f'✗ PDF do descritivo não encontrado: {descritivo_pdf_path}')
            return
        if not os.path.exists(recibo_pdf_path):
            print(f'✗ PDF do recibo não encontrado: {recibo_pdf_path}')
            return

        os.makedirs(os.path.dirname(output_pdf_path), exist_ok=True)

        merger = PdfMerger()
        merger.append(descritivo_pdf_path)
        merger.append(recibo_pdf_path)

        with open(output_pdf_path, 'wb') as f_out:
            merger.write(f_out)

        merger.close()
        print(f'✓ PDF final mesclado salvo em: {output_pdf_path}')

    except Exception as e:
        print(f'✗ Erro ao mesclar PDFs: {e}')
        import traceback
        traceback.print_exc()
        raise

def anexar_pdfs_formulario(driver, tentativa=-1):
    """Anexa os PDFs da pasta pdf_merge nos inputs correspondentes do formulário"""
    try:
        print('\nAnexando PDFs ao formulário...')

        # Lista para armazenar processos alternativos
        processos_alternativos = []

        continuar_anexar = True

        # Procura o input correspondente no formulário
        encontrado = False
        indice = 1
        indice_simbolo = 2
        
        
        while True:
            try:

                # Tenta encontrar o input cau_X
                cau_input = driver.find_element(By.ID, f'cau_{indice}')
                cau_value = cau_input.get_attribute('value')
                
                # Tenta encontrar o input Pleito_X
                pleito_input = driver.find_element(By.ID, f'Pleito_{indice}')
                pleito_value = pleito_input.get_attribute('value').strip().replace('.','')
                
                # Encontra o numero do processo importado no sistema
                numero_processo_sistema = driver.find_element(By.ID, f'pro_{indice}')
                numero_processo_sistema_value = numero_processo_sistema.get_attribute('value')

                # Verifica se os valores correspondem
                    
                elemento = driver.find_element(By.XPATH, f'/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[{indice_simbolo}]')
                html_interno = elemento.get_attribute('innerHTML')
                    # and numero_processo == numero_processo_sistema_value
                if '/images/check_green.gif' not in html_interno:

                    numero_processo = driver.find_element( By.XPATH, f"//input[@type='hidden' and @id='pro_{indice}']").get_attribute("value")
                    numero_integracao = driver.find_element( By.XPATH, f"//input[@type='hidden' and @id='cau_{indice}']").get_attribute("value")  

                    if '#3146c7' in html_interno and tentativa == -1:
                        print('  ⚠ Processo com status não encontrado, nova tentativa de importacao com formatação acontecerá')

                        processos_alternativos.append({
                            "numero_processo_pesq": formatar_numero_processo_cnj(numero_processo),
                            "numero_integracao": numero_integracao,
                            "remover": "nao"
                        })

                    elif '#3146c7' in html_interno and tentativa > -1:

                        quantidade_processos = buscar_quantidade_processos(numero_integracao)

                        if quantidade_processos == 1:

                            processos_alternativos.append({
                                "numero_processo_pesq": formatar_numero_processo_cnj(numero_processo),
                                "numero_integracao": numero_integracao,
                                "remover": "sim"
                            })

                        elif quantidade_processos == 2:

                            if tentativa == 0:

                                novo_numero_processo = buscar_processo_alternativo(numero_integracao, numero_processo, numero_processo, tentativa)
                                
                                processos_alternativos.append({
                                    "numero_processo_pesq": re.sub(r'[^0-9]', '', str(novo_numero_processo)),
                                    "numero_integracao": numero_integracao,
                                    "remover": "nao"
                                })

                            elif tentativa == 1:

                                processos_alternativos.append({
                                    "numero_processo_pesq": formatar_numero_processo_cnj(novo_numero_processo) ,
                                    "numero_integracao": numero_integracao,
                                    "remover": "nao"
                                })

                                pdb.set_trace() #debug para caso específico 1- 3146c7

                            elif tentativa == 2:

                                pdb.set_trace() #debug para caso específico 2- 3146c7

                        else:
                            pdb.set_trace() #debug para caso específico 3146c7
                            #processo_alternativo = buscar_processo_alternativo(numero_integracao, numero_processo, numero_processo, tentativa)

                    elif '#00ff21' in html_interno:

                        pdb.set_trace() #debug para caso específico 00ff21

                    else:
                        print('XXXXXXXXXXXXXXXXXXX TRATAR STATUS NOVO XXXXXXXXXXXXXXXXXXX')
                        pdb.set_trace()

                    continuar_anexar = False
  
                print(f'  ✓ Correspondência encontrada no índice {indice}')
                    
                # Anexa o PDF no input docLancfin_X
                doc_input = driver.find_element(By.ID, f'docLancfin_{indice}')
                
                # Se for o primeiro índice (cau_1), faz merge com descritivo
                if indice == 1:
                    print('  → Primeira posição: fazendo merge com descritivo...')

                    # Verifica se o campo já está preenchido
                    campo_preenchido = doc_input.get_attribute('value')
                    if campo_preenchido:
                        print(f'  ⚠ Campo docLancfin_{indice} já preenchido, pulando...')
                        encontrado = True

                    pdf_path = DESTINO_DESCRITIVO_RECIBO_PADRAO_FINAL_BMG_PDF
                
                else: 
                    pdf_path = DESTINO_DESCRITIVO_PADRAO_PDF

                doc_input.send_keys(pdf_path)
                        
                print(f'  ✓ PDF anexado: {os.path.basename(pdf_path)}')

                indice += 1
                indice_simbolo +=  4
 
            except:

                if continuar_anexar == False:
                    return {"retorno": False, "processos_alternativos": processos_alternativos}
                
                # Se não encontrar mais inputs, sai do loop
                break
        
    
        print('\n✓ Todos os PDFs foram processados!')
        return {"retorno": True, "processos_alternativos": []}
        
        
    except Exception as e:
        print(f'✗ Erro ao anexar PDFs: {e}')
        import traceback
        traceback.print_exc()
        return {"retorno": True, "processos_alternativos": []}


def main() -> None:
    print('=' * 70)
    print('LANCAMENTO-PREPOSTOS')
    print('=' * 70)
    print(f'Iniciado em: {datetime.now():%Y-%m-%d %H:%M:%S}')
    print(f'Config path: {CONFIG_PATH}')

    configurar_api()
    print('✓ API CPJ configurada')

    # Placeholder inicial para o fluxo do projeto.
    # A partir daqui, a automacao pode autenticar e consultar tarefas/processos.
    token = api_login()
    if token:

        tipo_titulo = False

        print('✓ Login na API CPJ realizado com sucesso')
        
        if tipo_titulo == False:

            integracao = False
            paths_ata_audiencia = []
            tarefas_processamento_atualizadas = []
            driver_daycoval = None
            os.makedirs(DOWNLOADS_PATH, exist_ok=True)

            data_inicial = datetime.strptime("2026-06-18", "%Y-%m-%d")
            data_final = datetime.today() - timedelta(days=3)
            data = data_inicial
            caminho_tarefas_json = os.path.join(BASE_PATH, 'tarefas.json')

            if integracao == True:

                if os.path.exists(caminho_tarefas_json):
                    try:
                        os.remove(caminho_tarefas_json)
                        print(f'✓ Arquivo limpo antes de iniciar: {caminho_tarefas_json}')
                    except Exception as e:
                        print(f'⚠ Não foi possível limpar {caminho_tarefas_json}: {e}')

                while data <= data_final:

                    filtro_tarefas = {
                        "_and": [
                            {
                                "evento": {
                                    "_in": "DAUD,$RA"
                                }
                            },
                            # {
                            #     "id_processo": {
                            #         "_eq": 293768
                            #     }
                            # },
                            {
                                "ag_data_hora": {
                                    "_gte": f"{data.strftime('%Y-%m-%d')}T00:00:00.000-03:00"
                                }
                            },
                            {
                                "ag_data_hora": {
                                    "_lte": f"{data.strftime('%Y-%m-%d')}T23:59:59.999-03:00"
                                }
                            }
                        ]
                    }

                    tarefas = api_buscar_processo_tarefa_filter(
                        filter_data=filtro_tarefas,
                        limit=1000,
                    )

                    data += timedelta(days=1)

                    tarefas_brutas = tarefas if tarefas is not None else []
                    tarefas_por_processo = {}
                    tarefas_sem_id = []
                    qtd_duplicadas = 0

                    for tarefa in tarefas_brutas:
                        id_processo = tarefa.get('id_processo')

                        if not id_processo:
                            tarefas_sem_id.append(tarefa)
                            continue

                        if id_processo in tarefas_por_processo:
                            qtd_duplicadas += 1

                        # Mantem a ocorrencia mais recente do mesmo id_processo.
                        tarefas_por_processo[id_processo] = tarefa

                    tarefas_unificadas = list(tarefas_por_processo.values()) + tarefas_sem_id
                    tarefas_processamento = tarefas_unificadas
                    try:
                        tarefas_existentes = []
                        if os.path.exists(caminho_tarefas_json):
                            with open(caminho_tarefas_json, 'r', encoding='utf-8') as f_tarefas:
                                dados_existentes = json.load(f_tarefas)
                                if isinstance(dados_existentes, list):
                                    tarefas_existentes = dados_existentes
                                elif isinstance(dados_existentes, dict):
                                    tarefas_existentes = dados_existentes.get('tarefas', [])

                        tarefas_para_salvar = tarefas_existentes + tarefas_unificadas

                        with open(caminho_tarefas_json, 'w', encoding='utf-8') as f_tarefas:
                            json.dump(tarefas_para_salvar, f_tarefas, ensure_ascii=False, indent=2, default=str)

                        with open(caminho_tarefas_json, 'r', encoding='utf-8') as f_tarefas:
                            tarefas_processamento = json.load(f_tarefas)

                        print(
                            f'✓ tarefas salvas em: {caminho_tarefas_json} '
                            f'(total={len(tarefas_brutas)}, unificadas={len(tarefas_unificadas)}, duplicadas={qtd_duplicadas}, acumuladas={len(tarefas_para_salvar)})'
                        )
                    except Exception as e:
                        print(f'⚠ Não foi possível salvar tarefas.json: {e}. Usando dados em memória.')

                
                if tarefas_processamento:
                        total_tarefas = len(tarefas_processamento)
                        print(f'✓ Busca retornou {total_tarefas} tarefa(s) após unificação')

                        for idx, tarefa in enumerate(tarefas_processamento, start=1):
                            id_processo = tarefa.get('id_processo')

                            if not id_processo:
                                print(f'  [{idx}/{total_tarefas}] ⚠ Tarefa sem id_processo, pulando...')
                                continue

                            dados_processo = api_buscar_processo_por_pj(id_processo)

                            if dados_processo:
                                tarefa['processo_detalhes'] = dados_processo
                                print(f'  [{idx}/{total_tarefas}] id_processo={id_processo} -> processo(s): {len(dados_processo)}')
                            else:
                                tarefa['processo_detalhes'] = []
                                print(f'  [{idx}/{total_tarefas}] ⚠ Nenhum processo encontrado para id_processo={id_processo}')
                                continue

                            tarefa['grupo_trabalho'] = dados_processo[0].get('grupo_trabalho') 

                            tarefas_processamento_atualizadas.append(tarefa)

                            try:
                                with open(caminho_tarefas_json, 'w', encoding='utf-8') as f_tarefas:
                                    json.dump(
                                        tarefas_processamento_atualizadas,
                                        f_tarefas,
                                        ensure_ascii=False,
                                        indent=2,
                                        default=str,
                                    )
                                print(
                                    f'✓ tarefas.json atualizado com {len(tarefas_processamento_atualizadas)} tarefa(s) '
                                    f'após processar id_processo={id_processo}'
                                )
                            except Exception as e:
                                print(f'⚠ Não foi possível atualizar tarefas.json após append: {e}')

                tarefas_ordenadas = sorted(tarefas_processamento_atualizadas, key=lambda x: x["grupo_trabalho"])
            
            else:
                
                try:
                    with open(caminho_tarefas_json, 'r', encoding='utf-8') as f_tarefas:
                        tarefas_json = json.load(f_tarefas)

                    if isinstance(tarefas_json, dict):
                        tarefas_ordenadas = tarefas_json.get('tarefas', [])
                    elif isinstance(tarefas_json, list):
                        tarefas_ordenadas = tarefas_json
                    else:
                        tarefas_ordenadas = []

                    tarefas_ordenadas = sorted(
                        tarefas_ordenadas,
                        key=lambda x: x.get('grupo_trabalho') or 0,
                    )
                except Exception as e:
                    print(f'⚠ Não foi possível carregar tarefas.json para processamento: {e}')
                    tarefas_ordenadas = sorted(
                        tarefas_processamento_atualizadas,
                        key=lambda x: x.get('grupo_trabalho') or 0,
                    )


            #processa tarefas 
            for idx, tarefa in enumerate(tarefas_ordenadas, start=1):

                    cliente_grupo_trabalho = tarefa.get('processo_detalhes', [{}])[0].get('grupo_trabalho')
                    id_processo = tarefa.get('id_processo')
                    dados_processo = tarefa.get('processo_detalhes', [{}])
                    id_tramitacao_atual = tarefa.get('id_tramitacao')

                    #REMOVERRRRRRRRRRRR
                    if cliente_grupo_trabalho != 88:
                        continue

                    #IFOOD #88 de grupo de trabalho
                    if cliente_grupo_trabalho == 88:

                        #carrega dados do ifood_lancamentos_sucesso.json para evitar reprocessamento
                        caminho_sucesso_ifood = os.path.join(BASE_PATH, 'ifood_lancamentos_sucesso.json')
                        id_tramitacao_sucesso = set()
                        try:
                            if os.path.exists(caminho_sucesso_ifood):
                                with open(caminho_sucesso_ifood, 'r', encoding='utf-8') as f_sucesso:
                                    conteudo_sucesso = f_sucesso.read().strip()
                                    if conteudo_sucesso:
                                        dados_sucesso = json.loads(conteudo_sucesso)
                                        if isinstance(dados_sucesso, list):
                                            for item_sucesso in dados_sucesso:
                                                if not isinstance(item_sucesso, dict):
                                                    continue
                                                id_tramitacao_item = item_sucesso.get('id_tramitacao')
                                                if id_tramitacao_item is None and isinstance(item_sucesso.get('tarefa'), dict):
                                                    id_tramitacao_item = item_sucesso['tarefa'].get('id_tramitacao')
                                                if id_tramitacao_item is not None:
                                                    id_tramitacao_sucesso.add(id_tramitacao_item)

                        except Exception as sucesso_error:
                            print(f'⚠ Não foi possível carregar ifood_lancamentos_sucesso.json: {sucesso_error}')
                        
                        if id_tramitacao_atual in id_tramitacao_sucesso:
                            print(f'  [{idx}/{len(tarefas_ordenadas)}] ✓ id_tramitacao já processado com sucesso: {id_tramitacao_atual}')
                            continue

                        filtro_tarefa = {
                            "_and": [
                                {
                                    "_or": [
                                    {
                                        "evento": {
                                        "_eq": "ATA"
                                        }
                                    },
                                    {
                                        "evento": {
                                        "_eq": "IA AT"
                                        }
                                    },
                                    {
                                        "evento": {
                                        "_eq": "IA ATA"
                                        }
                                    }
                                    ]
                                },
                                {
                                    "id_processo": {
                                        "_eq": id_processo
                                    }
                                },
                                {
                                    "ag_data_hora": {
                                        "_gte": tarefa['ag_data_hora']
                                    }
                                }
                            ]
                        }

                        tarefas_ata = api_buscar_processo_tarefa_filter(
                            filter_data=filtro_tarefa,
                            limit=1,
                        )

                        qtd_ata = len(tarefas_ata) if tarefas_ata else 0
                        caminho_sem_ata_json = os.path.join(BASE_PATH, 'ifood_tarefas_sem_ata.json')

                        if qtd_ata == 0:

                            registrar_tarefa_sem_ata(
                                tarefa=tarefa,
                                id_processo=id_processo,
                                caminho_saida=caminho_sem_ata_json,
                                motivo = 'Nenhuma tarefa de ATA encontrada para o processo'
                            )

                            continue

                        if tarefas_ata:
                            documentos_pj = api_buscar_documentos_pj(10, id_processo) or []

                            id_ged_filtrados = ""

                            if len(documentos_pj) == 0:
                                print(f'    ⚠ Nenhuma ata encontrada para id_processo={id_processo}')

                                registrar_tarefa_sem_ata(
                                    tarefa=tarefa,
                                    id_processo=id_processo,
                                    caminho_saida=caminho_sem_ata_json,
                                    motivo = 'Nenhum documento encontrado para o processo'
                                )

                                continue

                                
                            for documento in documentos_pj:
                                if not isinstance(documento, dict):
                                    continue

                                path_documento = documento.get('path')
                                if not path_documento:
                                    continue

                                path_normalizado = normalizar_texto_busca(path_documento)

                                if 'ata_audiencia' in path_normalizado:
                                    id_ged_filtrados = documento.get('id_ged')
                                    break
                            
                            if id_ged_filtrados == '':
                                print(f'    ⚠ Nenhuma ata encontrada para id_processo={id_processo}')

                                registrar_tarefa_sem_ata(
                                    tarefa=tarefa,
                                    id_processo=id_processo,
                                    caminho_saida=caminho_sem_ata_json,
                                    motivo = 'Nenhum documento de ata encontrado anexado no processo'
                                )

                                continue
                            
                            if id_ged_filtrados:
                                print(f'    ✓ Arquivos com "ata de audiencia": {id_ged_filtrados}')

                                arquivo_local = os.path.join(
                                    DOWNLOADS_PATH,
                                    f'id_processo_{id_processo}_id_ged_{id_ged_filtrados}.pdf'
                                )
                                sucesso_download = api_baixar_documento(id_ged_filtrados, arquivo_local)

                                if sucesso_download and os.path.exists(arquivo_local):
                                    with open(arquivo_local, 'rb') as f_pdf:
                                        b64 = base64.b64encode(f_pdf.read()).decode('utf-8')

                                    resp_audiencia = _requests.post(
                                        'https://n8n-diascosta.up.railway.app/webhook/ocr-guia',
                                        json={'base64': b64, 'mime_type': 'application/pdf', 'prompt' : """Você é um analisador de atas de audiência judiciais.

                                                Sua tarefa é analisar o texto da ata e identificar se a audiência possui participação presencial.

                                                Regras:

                                                - Retorne TRUE quando houver qualquer indicação de presença física ou possibilidade de participação presencial, incluindo termos como:
                                                - presencial
                                                - pregão presencial
                                                - sala de audiência
                                                - fórum
                                                - comparecimento presencial
                                                - audiência híbrida
                                                - virtual/presencial
                                                - pregão virtual e presencial

                                                - Retorne FALSE quando a audiência for exclusivamente virtual/remota, contendo apenas referências a videoconferência, webconferência, Teams, Zoom, Meet ou participação remota, sem qualquer menção à modalidade presencial.
                                                - Se houver a possibilidade de participação virtual, mesmo que a audiência seja híbrida, retorne false para "audiencia_presencial".
                                                - Se não houver informação suficiente para determinar a modalidade, retorne FALSE.
                                                - Se a audicencia for híbrida, ou seja, permitir participação tanto presencial quanto virtual, considere como "audiencia_presencial": false.

                                                IMPORTANTE:
                                                - Analise todo o texto.
                                                - Não faça suposições.
                                                - Retorne apenas JSON válido.
                                                - Não inclua explicações, comentários ou markdown.
                                                

                                                Formato obrigatório de saída:

                                                {
                                                "audiencia_presencial": true
                                                }"""},
                                        timeout=500
                                    )

                                    if resp_audiencia.status_code == 200:
                                        try:
                                            texto = resp_audiencia.json()[0]['output'][0]['content'][0]['text']
                                            audiencia_presencial = json.loads(texto)['audiencia_presencial']
                                        except Exception:
                                            audiencia_presencial = resp_audiencia.text

                                        tarefa['audiencia_presencial'] = audiencia_presencial

                                        try:
                                            with open(CONFIG_IFOOD_PATH, 'r', encoding='utf-8') as f:
                                                config_ifood = json.load(f)
                                        except Exception as e:
                                            print(f'  ⚠ Erro ao carregar config_ifood.json: {e}')
                                            config_ifood = {
                                                'audiencia_virtual': {'valor': 50},
                                                'audiencia_presencial': {'valor': 100},
                                            }

                                        url_ifood= config_ifood.get('url', '').strip()
                                        usuario_ifood = config_ifood.get('usuario', '').strip()
                                        senha_ifood = config_ifood.get('senha', '').strip()

                                        dados_audiencia = config_ifood.get('audiencia_virtual')
                                        if audiencia_presencial:
                                            dados_audiencia= config_ifood.get('audiencia_presencial')

                                        if url_ifood:
                                            driver_ifood_ativo = False
                                            if driver_ifood_ativo is not None:
                                                try:
                                                    _ = driver_ifood.current_window_handle
                                                    driver_ifood_ativo = True
                                                except Exception:
                                                    driver_ifood = None

                                            if not driver_ifood_ativo:
                                                print(f'    ✓ Abrindo Ifood: {url_ifood}')
                                                driver_ifood = abrir_chrome_e_acessar_url(
                                                    url_destino=url_ifood
                                                )

                                                if usuario_ifood and senha_ifood:
                                                    
                                                    login = login_ifood(driver_ifood, usuario_ifood, senha_ifood)

                                                    if login:
                                                        salvar_cookies_ifood(driver_ifood, COOKIES_JSON_IFOOD_PATH)
                                                    else:
                                                        pdb.set_trace() #nao logado

                                                else:
                                                    print('    ⚠ Usuário/senha Ifood ausentes no config_ifood.json')
                                                    pdb.set_trace() #debug login_ifood(driver_ifood, usuario_ifood, senha_ifood)
                                            else:
                                                print('    ✓ Reutilizando driver Ifood já aberto')

                                            numero_processo = str(dados_processo[0].get('numero_processo', '')).strip()
                                            if numero_processo:

                                                driver_ifood.get(URL_CADASTRO)

                                                select_centro_custo = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_ddlRegionalCelula')
                                                        )
                                                    )

                                                try:
                                                    select = Select(select_centro_custo)
                                                    select.select_by_value('071B708D-AF63-4F08-83F3-9EC136378D9C')
                                                    print('    ✓ Centro de custo selecionado no Ifood: 071B708D-AF63-4F08-83F3-9EC136378D9C')
                                                except Exception as select_error:
                                                    print(f'    ⚠ Falha ao selecionar centro de custo: {select_error}')


                                                select_coligada = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_ddlColigada')
                                                        )
                                                    )

                                                try:
                                                    select = Select(select_coligada)
                                                    select.select_by_value('9610')
                                                    print('    ✓ Coligada selecionada')
                                                    time.sleep(2)
                                                except Exception as select_error:
                                                    print(f'    ⚠ Falha ao selecionar centro de custo: {select_error}')

                                                text_area = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtObservacaoPagamento')
                                                        )
                                                    )
                                                
                                                text_area.send_keys('TRATE-SE DE REEMBOLSO DE CUSTAS E PREPOSTOS')

                                                try:
                                                    campo_numero_processo = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtNrProcesso')
                                                        )
                                                    )
                                                    campo_numero_processo.clear()
                                                    campo_numero_processo.send_keys(numero_processo)
                                                    print(f'    ✓ Número do processo preenchido no Ifood: {numero_processo}')

                                                    tipo_despesa = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_ddlTipoDespesa')
                                                        )
                                                    )

                                                    try:
                                                        select = Select(tipo_despesa)
                                                        select.select_by_value(dados_audiencia['codigo'])
                                                        print('    ✓ Tipo de despesa selecionada: ' + dados_audiencia['texto'])
                                                    except Exception as select_error:
                                                        print(f'    ⚠ Falha ao tipo de despesa: {select_error}')


                                                    data_audiencia_formatada = formatar_data_lancamento(tarefas_ata[0]['ag_data_hora'])
                                                    data_audiencia = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtDataDespesa_I')
                                                        )
                                                    )

                                                    data_audiencia.send_keys(Keys.CONTROL, "a")
                                                    data_audiencia.send_keys(Keys.BACKSPACE)
                                                    data_audiencia.send_keys(data_audiencia_formatada[0:2])
                                                    data_audiencia.send_keys(data_audiencia_formatada[3:5])
                                                    data_audiencia.send_keys(data_audiencia_formatada[6:]) 

                                                    campo_numero_documento = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtNumeroDocumento')
                                                        )
                                                    )
                                                    campo_numero_documento.clear()
                                                    campo_numero_documento.send_keys('N/A')
                                                    print(f'    ✓ Número do documento: {'N/A'}')

                                                    campo_valor_audiencia = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtValorDespesa')
                                                        )
                                                    )
                                                    campo_valor_audiencia.clear()
                                                    campo_valor_audiencia.send_keys(str(dados_audiencia['valor']) + '.00')
                                                    print(f'    ✓ Valor audiência: {dados_audiencia['valor']}')

                                                    text_area_descricao = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_txtDescricaoDespesa')
                                                        )
                                                    )
                                                
                                                    text_area_descricao.send_keys(f'AUDIÊNCIA REALIZADA NO DIA {data_audiencia_formatada}')
                                                    print(f'    ✓ Campo descrição preenchido.')


                                                    campo_responsavel = WebDriverWait(driver_ifood, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphPrincipal_wcResponsavel_cmbResponsavel_I')
                                                        )
                                                    )

                                                    campo_responsavel.send_keys('Ariane Soares da Silva Coutinho')
                                                    campo_responsavel.click()

                                                    print(f'    ✓ Número do documento: {'N/A'}')

                                                    doc_numero = numero_processo
                                                    doc_valor = str(dados_audiencia['valor'])
                                                    doc_extenso = str(dados_audiencia.get('extenso', ''))
                                                    doc_texto = str(dados_audiencia.get('texto', ''))
                                                    doc_data_extenso = formatar_data_por_extenso(datetime.now())
                                                    pdf_concatenado = None

                                                    if gerar_recibo_ifood(
                                                        numero_processo=doc_numero,
                                                        valor=doc_valor,
                                                        valor_extenso=doc_extenso,
                                                        texto=doc_texto,
                                                        data_por_extenso=doc_data_extenso,
                                                    ):
                                                        pasta_arquivos_ifood = os.path.join(BASE_PATH, 'arquivos_gerados', 'ifood')
                                                        os.makedirs(pasta_arquivos_ifood, exist_ok=True)
                                                        pdf_concatenado = os.path.join(
                                                            pasta_arquivos_ifood,
                                                            f'arquivo_concatenado_{id_processo}_{id_ged_filtrados}.pdf'
                                                        )
                                                        if concatenar_pdfs([RECIBO_PADRAO_IFOOD_COPY_PDF, arquivo_local], pdf_concatenado):
                                                            try:
                                                                upload_input = WebDriverWait(driver_ifood, 20).until(
                                                                    EC.presence_of_element_located(
                                                                        (By.ID, 'ctl00_cphPrincipal_oInputPathFile')
                                                                    )
                                                                )
                                                                upload_input.send_keys(pdf_concatenado)
                                                                print(f'    ✓ Arquivo concatenado anexado: {pdf_concatenado}')

                                                                

                                                                try:
                                                                    
                                                                    botao_adicionar = WebDriverWait(driver_ifood, 20).until(
                                                                        EC.element_to_be_clickable(
                                                                            (By.ID, 'ctl00_cphPrincipal_btnAdicionarCasosRelacionados')
                                                                        )
                                                                    )
                                                                    botao_adicionar.click()
                                                                    print('    ✓ Clique em Adicionar realizado com sucesso')

                                                                    time.sleep(5)

                                                                    botao_adicionar = WebDriverWait(driver_ifood, 20).until(
                                                                        EC.element_to_be_clickable(
                                                                            (By.ID, 'ctl00_cphPrincipal_btnAdicionarCasosRelacionados')
                                                                        )
                                                                    )
                                                                    botao_adicionar.click()

                                                                except Exception as adcionar_error:
                                                                    print(f'    ⚠ Falha ao clicar em Adicionar : {adcionar_error}')
                                                                    pdb.set_trace()  #debug adicionar

                                                                try:
                                                                    aguardar_loading_ifood(driver_ifood, 60)

                                                                    botao_enviar = WebDriverWait(driver_ifood, 20).until(
                                                                        EC.element_to_be_clickable(
                                                                            (By.ID, 'ctl00_btnSalvar')
                                                                        )
                                                                    )
                                                                    
                                                                    botao_enviar.click()
                                                                    print('    ✓ Clique em Enviar realizado com sucesso')

                                                                    aguardar_loading_ifood(driver_ifood, 60)

                                                                    try:

                                                                        tela_sucesso = WebDriverWait(driver_ifood, 20).until(
                                                                            EC.element_to_be_clickable(
                                                                                (By.CLASS_NAME, 'swal-text')
                                                                            )
                                                                        )

                                                                        if 'Documento salvo com sucesso!' in tela_sucesso.text:
                                                                            print('    ✓ Documento salvo com sucesso')

                                                                            try:
                                                                                botao_ok = WebDriverWait(driver_ifood, 20).until(
                                                                                    EC.element_to_be_clickable(
                                                                                        (By.CLASS_NAME, 'swal-button--confirm')
                                                                                    )
                                                                                )

                                                                                botao_ok.click()

                                                                                print('    ✓ Clique em OK realizado com sucesso')

                                                                                botao_iniciar_aprovacao = WebDriverWait(driver_ifood, 20).until(
                                                                                    EC.element_to_be_clickable(
                                                                                        (By.ID, 'ctl00_btnIniciarFluxo')
                                                                                    )
                                                                                )

                                                                                botao_iniciar_aprovacao.click()

                                                                                print('    ✓ Clique em Iniciar Aprovação realizado com sucesso')

                                                                            except Exception as iniciar_aprovacao_error:
                                                                                print(f'    ⚠ Falha ao clicar em Iniciar Aprovação: {iniciar_aprovacao_error}')
                                                                                pdb.set_trace()  #debug iniciar aprovacao

                                                                            try:
                                                                                WebDriverWait(driver_ifood, 10).until(EC.alert_is_present())
                                                                                driver_ifood.switch_to.alert.accept()
                                                                                print('    ✓ Alerta de confirmação do fluxo aceito com sucesso')
                                                                            except Exception as alert_error:
                                                                                print(f'    ⚠ Falha ao confirmar alerta de início do fluxo: {alert_error}')
                                                                                pdb.set_trace()  #debug alerta fluxo aceito 
                                                                            
                                                                            #aguardar_loading_ifood(driver_ifood, 60)

                                                                            time.sleep(5)  # Aguarda o alerta de sucesso aparecer
                                                                            
                                                                            try:
                                                                                alert_sucesso = WebDriverWait(driver_ifood, 5).until(EC.alert_is_present())
                                                                                mensagem_sucesso = (alert_sucesso.text or '').strip()
                                                                                if mensagem_sucesso == 'O fluxo foi iniciado com sucesso para aprovação e liberação.':
                                                                                    alert_sucesso.accept()
                                                                                    print('    ✓ OK no alerta de sucesso do fluxo realizado')

                                                                                    try:
                                                                                        registros_sucesso = []

                                                                                        if os.path.exists(caminho_sucesso_ifood):
                                                                                            with open(caminho_sucesso_ifood, 'r', encoding='utf-8') as f_sucesso:
                                                                                                conteudo_sucesso = f_sucesso.read().strip()
                                                                                                if conteudo_sucesso:
                                                                                                    dados_sucesso = json.loads(conteudo_sucesso)
                                                                                                    if isinstance(dados_sucesso, list):
                                                                                                        registros_sucesso = dados_sucesso

                                                                                        registro = dict(tarefa)
                                                                                        registro['salvo_em'] = datetime.now().isoformat()
                                                                                        registros_sucesso.append(registro)

                                                                                        with open(caminho_sucesso_ifood, 'w', encoding='utf-8') as f_sucesso:
                                                                                            json.dump(registros_sucesso, f_sucesso, ensure_ascii=False, indent=2, default=str)

                                                                                        pdb.set_trace() #debug final

                                                                                        print(f'    ✓ Tarefa salva em: {caminho_sucesso_ifood}')
                                                                                    except Exception as salvar_sucesso_error:
                                                                                        print(f'    ⚠ Falha ao salvar tarefa de sucesso do iFood: {salvar_sucesso_error}')

                                                                                else:
                                                                                    print(f'    ⚠ Alerta inesperado após iniciar fluxo: {mensagem_sucesso}')
                                                                                    pdb.set_trace()  #debug alerta inesperado
                                                                            except Exception:
                                                                                pdb.set_trace()  #debug alerta fluxo sucesso
                                                                                pass

                                                                            
                                                                            

                                                                        else:
                                                                            print('    ⚠ Falha ao salvar documento')
                                                                            pdb.set_trace()  #debug erro ao salvar documento
                                                                    
                                                                    except: 
                                                                        pdb.set_trace()  #debug tela sucesso


                                                                except Exception as enviar_error:
                                                                    print(f'    ⚠ Falha ao clicar em Enviar: {enviar_error}')
                                                                    pdb.set_trace()  #debug enviar

                                                            except Exception as upload_error:
                                                                print(f'    ⚠ Falha ao anexar arquivo: {upload_error}')
                                                        else:
                                                            print('    ⚠ Falha ao concatenar PDFs')
                                                    else:
                                                        print('    ⚠ Falha ao gerar recibo iFood')

                                                    try:
                                                        if os.path.exists(RECIBO_PADRAO_IFOOD_COPY_DOCX):
                                                            os.remove(RECIBO_PADRAO_IFOOD_COPY_DOCX)
                                                        if os.path.exists(RECIBO_PADRAO_IFOOD_COPY_PDF):
                                                            os.remove(RECIBO_PADRAO_IFOOD_COPY_PDF)
                                                        if pdf_concatenado and os.path.exists(pdf_concatenado):
                                                            os.remove(pdf_concatenado)
                                                        if os.path.exists(arquivo_local):
                                                            os.remove(arquivo_local)
                                                        print('    ✓ Arquivos temporários e original removidos')
                                                    except Exception as cleanup_error:
                                                        print(f'    ⚠ Erro ao limpar arquivos temporários: {cleanup_error}')                                                   
                                                except Exception as e:
                                                    pdb.set_trace()  # Debug: Captura de exceção durante o preenchimento do Daycoval
                                                    print(f'    ⚠ Não foi possível preencher número do processo no Daycoval: {e}')
                                            else:
                                                print('    ⚠ numero_processo vazio para preenchimento no Daycoval')
                                        else:
                                            print('    ⚠ URL Daycoval não encontrada no config_daycoval.json')
                                    else:
                                        print(f'    ✗ OCR falhou. Status: {resp_audiencia.status_code}')


                    #DAYCOVAL #78 de grupo de trabalho
                    elif cliente_grupo_trabalho == 78:
                        
                        try:
                            with open(CONFIG_DAYCOVAL_PATH, 'r', encoding='utf-8') as f:
                                config_daycoval = json.load(f)
                        except Exception as e:
                            print(f'  ⚠ Erro ao carregar config_daycoval.json: {e}')
                            config_daycoval = {
                                'valor_audiencia_virtual': 50,
                                'valor_audiencia_presencial': 100,
                            }

                        valor_audiencia_virtual = config_daycoval.get('valor_audiencia_virtual', 50)
                        valor_audiencia_presencial = config_daycoval.get('valor_audiencia_presencial', 100)

                        filtro_tarefa = {
                            "_and": [
                                {
                                    "_or": [
                                    {
                                        "evento": {
                                        "_eq": "ATA"
                                        }
                                    },
                                    {
                                        "evento": {
                                        "_eq": "IA AT"
                                        }
                                    },
                                    {
                                        "evento": {
                                        "_eq": "IA ATA"
                                        }
                                    }
                                    ]
                                },
                                {
                                    "id_processo": {
                                        "_eq": id_processo
                                    }
                                },
                                {
                                    "ag_data_hora": {
                                        "_gte": tarefa['ag_data_hora']
                                    }
                                }
                            ]
                        }

                        tarefas_ata = api_buscar_processo_tarefa_filter(
                            filter_data=filtro_tarefa,
                            limit=1,
                        )

                        qtd_ata = len(tarefas_ata) if tarefas_ata else 0
                        caminho_sem_ata_json = os.path.join(BASE_PATH, 'daycoval_tarefas_sem_ata.json')

                        if qtd_ata == 0:

                            registrar_tarefa_sem_ata(
                                tarefa=tarefa,
                                id_processo=id_processo,
                                caminho_saida=caminho_sem_ata_json,
                                motivo = 'Nenhuma tarefa de ATA encontrada para o processo'
                            )

                            continue

                        print(f'  [{idx}/{len(tarefas_ordenadas)}] id_processo={id_processo} -> ATA: {qtd_ata} tarefa(s)')

                        if tarefas_ata:
                            documentos_pj = api_buscar_documentos_pj(10, id_processo) or []

                            id_ged_filtrados = ""

                            if len(documentos_pj) == 0:
                                print(f'    ⚠ Nenhuma ata encontrada para id_processo={id_processo}')

                                registrar_tarefa_sem_ata(
                                    tarefa=tarefa,
                                    id_processo=id_processo,
                                    caminho_saida=caminho_sem_ata_json,
                                    motivo = 'Nenhum documento encontrado para o processo'
                                )

                                continue

                                
                            for documento in documentos_pj:
                                if not isinstance(documento, dict):
                                    continue

                                path_documento = documento.get('path')
                                if not path_documento:
                                    continue

                                path_normalizado = normalizar_texto_busca(path_documento)
                                
                                if 'ata_audiencia' in path_normalizado:
                                    id_ged_filtrados = documento.get('id_ged')
                                    break
                            
                            if id_ged_filtrados == '':
                                print(f'    ⚠ Nenhuma ata encontrada para id_processo={id_processo}')

                                registrar_tarefa_sem_ata(
                                    tarefa=tarefa,
                                    id_processo=id_processo,
                                    caminho_saida=caminho_sem_ata_json,
                                    motivo = 'Nenhum documento de ata encontrado anexado no processo'
                                )

                                continue
                                    
                            if id_ged_filtrados:
                                print(f'    ✓ Arquivos com "ata de audiencia": {id_ged_filtrados}')

                                arquivo_local = os.path.join(
                                    DOWNLOADS_PATH,
                                    f'id_processo_{id_processo}_id_ged_{id_ged_filtrados}.pdf'
                                )
                                sucesso_download = api_baixar_documento(id_ged_filtrados, arquivo_local)

                                if sucesso_download and os.path.exists(arquivo_local):
                                    with open(arquivo_local, 'rb') as f_pdf:
                                        b64 = base64.b64encode(f_pdf.read()).decode('utf-8')

                                    resp_audiencia = _requests.post(
                                        'https://n8n-diascosta.up.railway.app/webhook/ocr-guia',
                                        json={'base64': b64, 'mime_type': 'application/pdf', 'prompt' : """Você é um analisador de atas de audiência judiciais.

                                                Sua tarefa é analisar o texto da ata e identificar se a audiência possui participação presencial.

                                                Regras:

                                                - Retorne TRUE quando houver qualquer indicação de presença física ou possibilidade de participação presencial, incluindo termos como:
                                                - presencial
                                                - pregão presencial
                                                - sala de audiência
                                                - fórum
                                                - comparecimento presencial
                                                - audiência híbrida
                                                - virtual/presencial
                                                - pregão virtual e presencial

                                                - Retorne FALSE quando a audiência for exclusivamente virtual/remota, contendo apenas referências a videoconferência, webconferência, Teams, Zoom, Meet ou participação remota, sem qualquer menção à modalidade presencial.
                                                - Se houver a possibilidade de participação virtual, mesmo que a audiência seja híbrida, retorne false para "audiencia_presencial".
                                                - Se não houver informação suficiente para determinar a modalidade, retorne FALSE.
                                                - Se a audicencia for híbrida, ou seja, permitir participação tanto presencial quanto virtual, considere como "audiencia_presencial": false.

                                                IMPORTANTE:
                                                - Analise todo o texto.
                                                - Não faça suposições.
                                                - Retorne apenas JSON válido.
                                                - Não inclua explicações, comentários ou markdown.
                                                

                                                Formato obrigatório de saída:

                                                {
                                                "audiencia_presencial": true
                                                }"""},
                                        timeout=500
                                    )

                                    if resp_audiencia.status_code == 200:
                                        try:
                                            texto = resp_audiencia.json()[0]['output'][0]['content'][0]['text']
                                            audiencia_presencial = json.loads(texto)['audiencia_presencial']
                                        except Exception:
                                            audiencia_presencial = resp_audiencia.text

                                        tarefa['audiencia_presencial'] = audiencia_presencial

                                        valor = valor_audiencia_virtual
                                        if audiencia_presencial:
                                            valor = valor_audiencia_presencial

                                        url_daycoval = config_daycoval.get('url', '').strip()
                                        usuario_daycoval = config_daycoval.get('usuario', '').strip()
                                        senha_daycoval = config_daycoval.get('senha', '').strip()

                                        if url_daycoval:
                                            driver_daycoval_ativo = False
                                            if driver_daycoval is not None:
                                                try:
                                                    _ = driver_daycoval.current_window_handle
                                                    driver_daycoval_ativo = True
                                                except Exception:
                                                    driver_daycoval = None

                                            if not driver_daycoval_ativo:
                                                print(f'    ✓ Abrindo Daycoval: {url_daycoval}')
                                                driver_daycoval = abrir_chrome_e_acessar_url(
                                                    url_destino=url_daycoval,
                                                    url_base_cookies=url_daycoval,
                                                )

                                                if usuario_daycoval and senha_daycoval:
                                                    login_daycoval(driver_daycoval, usuario_daycoval, senha_daycoval)
                                                else:
                                                    print('    ⚠ Usuário/senha Daycoval ausentes no config_daycoval.json')
                                                    pdb.set_trace() #debug login_daycoval(driver_daycoval, usuario_daycoval, senha_daycoval)
                                            else:
                                                print('    ✓ Reutilizando driver Daycoval já aberto')
                                            
                                            driver_daycoval.get(DAYCOVAL_URL_CONSULTA)

                                            numero_processo = str(dados_processo[0].get('numero_processo', '')).strip()
                                            if numero_processo:
                                                try:
                                                    campo_numero_processo = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_ctl00_cphMain_pgcFiltros_cphFiltros_HucFoneticoDossieProcesso1_txtNumProc_I')
                                                        )
                                                    )
                                                    campo_numero_processo.clear()
                                                    campo_numero_processo.send_keys(numero_processo)
                                                    print(f'    ✓ Número do processo preenchido no Daycoval: {numero_processo}')

                                                    botao_consultar = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_ctl00_ctl00_tbarMain_pgcMain_pnlConsultar_HTC_mnuRibConsultar_DXI0_T')
                                                        )
                                                    )
                                                    botao_consultar.click()

                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_ctl00_cphMain_cphGrid_pgcGrid_cphGrid_grdMain_DXDataRow0')
                                                        )
                                                    )
                                                    print('    ✓ Consulta Daycoval carregada com sucesso')

                                                    item_menu = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.XPATH, '/html/body/form/div[6]/div[2]/div/div[2]/div[2]/div[2]/div[1]/div[1]/ul/li[11]')
                                                        )
                                                    )
                                                    item_menu.click()
                                                    print('    ✓ Clique no item da grade realizado com sucesso')

                                                    item_financeiro = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.XPATH, '/html/body/form/div[6]/div[2]/div/div[2]/div[2]/div[2]/div[1]/div[2]/div/ul/li[21]/div')
                                                        )
                                                    )
                                                    item_financeiro.click()
                                                    print('    ✓ Clique em Financeiro realizado com sucesso')

                                                    driver_daycoval.switch_to.default_content()
                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.frame_to_be_available_and_switch_to_it(
                                                            (By.ID, 'ctl00_ctl00_ctl00_popupMain_CIF0')
                                                        )
                                                    )
                                                    print('    ✓ Contexto alterado para iframe de popup')

                                                    lancamento_ja_realizado = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphGridDetail_grdMain_DXMainTable')
                                                        )
                                                    )

                                                    valor_audiencia_texto = str(valor).strip()+',00'
                                                    data_hora_tarefa = formatar_data_hora_tarefa(tarefa.get('ag_data_hora'))
                                                    texto_complemento = ('AUDIENCIA PRESENCIAL' if audiencia_presencial else 'AUDIENCIA VIRTUAL') + data_hora_tarefa

                                                    if valor_audiencia_texto in lancamento_ja_realizado.text and texto_complemento in lancamento_ja_realizado.text:
                                                        print(f'    ✓ Lançamento já realizado para valor {valor_audiencia_texto} e complemento "{texto_complemento}"')
                                                        continue

                                                    botao_novo_lancamento = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphGridDetail_dnavDeposito_DXI10_T')
                                                        )
                                                    )

                                                    botao_novo_lancamento.click()
                                                    print('    ✓ Clique em Novo Lancamento realizado com sucesso')

                                                    time.sleep(2)  # Aguarda o carregamento do iframe do popup

                                                    driver_daycoval.switch_to.default_content()
                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.frame_to_be_available_and_switch_to_it(
                                                            (By.ID, 'ctl00_ctl00_ctl00_popupMain_CIF1')
                                                        )
                                                    )

                                                    campo_valor_audiencia = WebDriverWait(driver_daycoval, 10).until(
                                                            EC.presence_of_element_located(
                                                                (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_spNumValor_I')
                                                            )
                                                        )

                                                    campo_valor_audiencia.clear()
                                                    campo_valor_audiencia.send_keys(valor_audiencia_texto)
                                                    print(f'    ✓ Valor da audiencia preenchido: {valor_audiencia_texto}')
                                                    time.sleep(0.5)

                                                    checkbox_pagador_favorecido = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_chkPagadorFavorecido_S_D')
                                                        )
                                                    )
                                                    if not checkbox_pagador_favorecido.is_selected():
                                                        checkbox_pagador_favorecido.click()
                                                    print('    ✓ Checkbox Pagador/Favorecido marcado')
                                                    time.sleep(0.5)

                                                    campo_lookup_outra_parte= WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_cliEscritFornLookupOutroCCE_I')
                                                        )
                                                    )
                                                    campo_lookup_outra_parte.click()
                                                    campo_lookup_outra_parte.send_keys(Keys.CONTROL, 'a')
                                                    campo_lookup_outra_parte.send_keys(Keys.DELETE)

                                                    campo_lookup_outra_parte.send_keys('FINANCEIRA CONSIGNADO')
                                                    print('    ✓ Texto enviado para campo Outro CCE: FINANCEIRA CONSIGNADO')
                                                    time.sleep(0.5)

                                                    campo_nome_favorecido = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_cbkpExecutar_NomeFavorecido_I')
                                                        )
                                                    )
                                                    campo_nome_favorecido.click()
                                                    campo_nome_favorecido.send_keys(Keys.CONTROL, 'a')
                                                    campo_nome_favorecido.send_keys(Keys.DELETE)
                                                    campo_nome_favorecido.send_keys('DIAS COSTA SOCIEDADE DE ADVOGADOS')
                                                    print('    ✓ Texto enviado para Nome Favorecido')
                                                    time.sleep(0.5)

                                                    campo_plano_conta = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_planoContaCCLookup1_I')
                                                        )
                                                    )
                                                    # campo_plano_conta.click()
                                                    # campo_plano_conta.send_keys(Keys.CONTROL, 'a')
                                                    # campo_plano_conta.send_keys(Keys.DELETE)
                                                    campo_plano_conta.send_keys('HON PREPOSTO - 2310')
                                                    print('    ✓ Texto enviado para Plano de Conta')
                                                    time.sleep(0.5)

                                                
                                                    texto_historico = 'PREPOSTO PRESENCIAL' if audiencia_presencial else 'PREPOSTO VIRTUAL'
                                                    campo_historico = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_historicoCCLookup1_I')
                                                        )
                                                    )
                                                    campo_historico.click()
                                                    # campo_historico.send_keys(Keys.CONTROL, 'a')
                                                    # campo_historico.send_keys(Keys.DELETE)
                                                    campo_historico.send_keys(texto_historico)
                                                    print(f'    ✓ Texto enviado para Historico: {texto_historico}')
                                                    time.sleep(0.5)

                                                    campo_tipo_doc = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_tipoDocCCLookup1_I')
                                                        )
                                                    )
                                                    campo_tipo_doc.click()
                                                    campo_tipo_doc.send_keys(Keys.CONTROL, 'a')
                                                    campo_tipo_doc.send_keys(Keys.DELETE)
                                                    campo_tipo_doc.send_keys('RECIBO')
                                                    print('    ✓ Texto enviado para Tipo Doc: RECIBO')
                                                    time.sleep(0.5)

                                                    campo_centro_custo = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_cCustoLookup1_I')
                                                        )
                                                    )
                                                    campo_centro_custo.click()
                                                    campo_centro_custo.send_keys(Keys.CONTROL, 'a')
                                                    campo_centro_custo.send_keys(Keys.DELETE)
                                                    campo_centro_custo.send_keys('11000.000 - FINANCEIRA CONSIGNADO')
                                                    print('    ✓ Texto enviado para Centro de Custo: 11000.000 - FINANCEIRA CONSIGNADO')

                                                    campo_complemento = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_cphMain_fvwMain_COMPLEMENTOTextEdit_I')
                                                        )
                                                    )
                                                    campo_complemento.click()
                                                    campo_complemento.send_keys(Keys.CONTROL, 'a')
                                                    campo_complemento.send_keys(Keys.DELETE)
                                                    campo_complemento.send_keys(texto_complemento)
                                                    print(f'    ✓ Texto enviado para Complemento: {texto_complemento}')

                                                    campo_historico.send_keys(texto_historico)

                                                    print('    ✓ Validando campos preenchidos no HTML antes de salvar...')
                                                    campos_para_validar = [
                                                        ('Valor da audiencia', campo_valor_audiencia, valor_audiencia_texto),
                                                        ('Outro CCE', campo_lookup_outra_parte, 'FINANCEIRA CONSIGNADO'),
                                                        ('Nome Favorecido', campo_nome_favorecido, 'DIAS COSTA SOCIEDADE DE ADVOGADOS'),
                                                        ('Plano de Conta', campo_plano_conta, 'HON PREPOSTO - 2310'),
                                                        ('Historico', campo_historico, texto_historico),
                                                        ('Tipo Doc', campo_tipo_doc, 'RECIBO'),
                                                        ('Centro de Custo', campo_centro_custo, '11000.000 - FINANCEIRA CONSIGNADO'),
                                                        ('Complemento', campo_complemento, texto_complemento),
                                                    ]

                                                    for nome_campo, elemento_campo, valor_esperado in campos_para_validar:
                                                        valor_atual = (elemento_campo.get_attribute('value') or '').strip()
                                                        esperado = str(valor_esperado).strip()

                                                        if esperado.upper() in valor_atual.upper():
                                                            print(f'    ✓ {nome_campo} validado no HTML: {valor_atual}')
                                                            continue

                                                        print(
                                                            f'    ⚠ {nome_campo} nao preenchido/correto no HTML. '
                                                            f'Atual: "{valor_atual}" | Esperado: "{esperado}"'
                                                        )
                                                        elemento_campo.click()
                                                        elemento_campo.send_keys(Keys.CONTROL, 'a')
                                                        elemento_campo.send_keys(Keys.DELETE)
                                                        elemento_campo.send_keys(esperado)
                                                        time.sleep(0.3)

                                                        valor_revalidado = (elemento_campo.get_attribute('value') or '').strip()
                                                        if esperado.upper() in valor_revalidado.upper():
                                                            print(f'    ✓ {nome_campo} repreechido com sucesso no HTML')
                                                        else:
                                                            print(
                                                                f'    ⚠ {nome_campo} ainda divergente apos repreechimento. '
                                                                f'Atual: "{valor_revalidado}"'
                                                            )

                                                    
                                                    salvar_informacoes = input("Salvar dados? (s/n): ").strip().lower()

                                                    if salvar_informacoes == 's':
                                                        #BOTAO SALVAR

                                                        botao_confirmar_lancamento = WebDriverWait(driver_daycoval, 20).until(
                                                            EC.element_to_be_clickable(
                                                                (By.ID, 'ctl00_ctl00_cphZoomMain_tbarMain_tdConfirma')
                                                            )
                                                        )
                                                        botao_confirmar_lancamento.click()
                                                        print('    ✓ Clique em Confirmar lancamento realizado com sucesso')

                                                    botao_fechar_lancamento = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_ctl00_cphZoomMain_tbarMain_tdFechar')
                                                        )
                                                    )
                                                    botao_fechar_lancamento.click()
                                                    print('    ✓ Clique em Fechar lancamento realizado com sucesso')

                                                    time.sleep(2)  # Aguarda o fechamento do popup

                                                    driver_daycoval.switch_to.default_content()
                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.frame_to_be_available_and_switch_to_it(
                                                            (By.ID, 'ctl00_ctl00_ctl00_popupMain_CIF0')
                                                        )
                                                    )

                                                    botao_detalhes = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphGridDetail_dnavDeposito_DXI17_T')
                                                        )
                                                    )
                                                    botao_detalhes.click()

                                                    time.sleep(1)  # Aguarda o carregamento do iframe de detalhes
                                                    
                                                    print('    ✓ Clique no botao detalhes realizado com sucesso')

                                                    #ctl00_cphGridDetail_dnavDeposito_DXME17_ | ctl00_cphGridDetail_dnavDeposito_DXI17i0_T
                                                    botao_documentos = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphGridDetail_dnavDeposito_DXME17_')
                                                        )
                                                    )
                                                    botao_documentos.click()
                                                    print('    ✓ Clique no botao documentos realizado com sucesso')

                                                    driver_daycoval.switch_to.default_content()
                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.frame_to_be_available_and_switch_to_it(
                                                            (By.ID, 'ctl00_ctl00_ctl00_popupMain_CIF1')
                                                        )
                                                    )

                                                    time.sleep(1)  # Aguarda o carregamento do iframe de documentos

                                                    botao_inserir = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphGridDetail_splitterHoriz_dnavMain_DXI10_Img')
                                                        )
                                                    )
                                                    botao_inserir.click()
                                                    print('    ✓ Clique no botao inserir realizado com sucesso')
                                                    time.sleep(1)  # Aguarda o carregamento do iframe de documentos

                                                    driver_daycoval.switch_to.default_content()
                                                    WebDriverWait(driver_daycoval, 20).until(
                                                        EC.frame_to_be_available_and_switch_to_it(
                                                            (By.ID, 'ctl00_ctl00_ctl00_popupMain_CIF2')
                                                        )
                                                    )

                                                    caminho_pdf_upload = os.path.abspath(arquivo_local)
                                                    campo_upload_pdf = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphZoomMain_fvwMain_uplDocto_TextBox0_Input')
                                                        )
                                                    )
                                                    campo_upload_pdf.send_keys(caminho_pdf_upload)
                                                    print(f'    ✓ PDF enviado para upload: {caminho_pdf_upload}')

                                                    campo_tipo_documento_upload = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphZoomMain_fvwMain_cbxTipoDoc_I')
                                                        )
                                                    )
                                                    campo_tipo_documento_upload.click()
                                                    # campo_tipo_documento_upload.send_keys(Keys.CONTROL, 'a')
                                                    # campo_tipo_documento_upload.send_keys(Keys.DELETE)
                                                    campo_tipo_documento_upload.send_keys('.Indefinido')
                                                    print('    ✓ Texto enviado para Tipo Doc do upload: .Indefinido')
                                                    time.sleep(0.5)

                                                    campo_autor_upload = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphZoomMain_fvwMain_cbxAutor_I')
                                                        )
                                                    )
                                                    campo_autor_upload.click()
                                                    # campo_autor_upload.send_keys(Keys.CONTROL, 'a')
                                                    # campo_autor_upload.send_keys(Keys.DELETE)
                                                    campo_autor_upload.send_keys('WALTER SILVEIRA MORAES')
                                                    print('    ✓ Texto enviado para Autor do upload: WALTER SILVEIRA MORAES')
                                                    time.sleep(0.5)

                                                    texto_descricao_upload= ('ATA DE AUDIENCIA PRESENCIAL' if audiencia_presencial else 'ATA DE AUDIENCIA VIRTUAL') + data_hora_tarefa

                                                    campo_descricao_upload = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located( 
                                                            (By.ID, 'ctl00_cphZoomMain_fvwMain_txtDescricao_I')
                                                        )
                                                    )
                                                    campo_descricao_upload.click()
                                                    # campo_descricao_upload.send_keys(Keys.CONTROL, 'a')
                                                    # campo_descricao_upload.send_keys(Keys.DELETE)
                                                    campo_descricao_upload.send_keys(texto_descricao_upload)
                                                    print(f'    ✓ Texto enviado para Descricao do upload: {texto_descricao_upload}')
                                                    time.sleep(0.5)

                                                    campo_resumo = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.presence_of_element_located(
                                                            (By.ID, 'ctl00_cphZoomMain_fvwMain_txtResumo_I')
                                                        )
                                                    )
                                                    campo_resumo.click()
                                                    # campo_resumo.send_keys(Keys.CONTROL, 'a')
                                                    # campo_resumo.send_keys(Keys.DELETE)
                                                    campo_resumo.send_keys(texto_complemento)
                                                    print(f'    ✓ Texto enviado para Resumo do upload: {texto_complemento}')
                                                    time.sleep(0.5)

                                                    print('    ✓ Validando campos de upload preenchidos no HTML...')
                                                    campos_upload_para_validar = [
                                                        ('Autor do upload', campo_autor_upload, 'WALTER SILVEIRA MORAES'),
                                                        ('Descricao do upload', campo_descricao_upload, texto_descricao_upload),
                                                        ('Resumo do upload', campo_resumo, texto_complemento),
                                                    ]

                                                    for nome_campo, elemento_campo, valor_esperado in campos_upload_para_validar:
                                                        valor_atual = (elemento_campo.get_attribute('value') or '').strip()
                                                        esperado = str(valor_esperado).strip()

                                                        if esperado.upper() in valor_atual.upper():
                                                            print(f'    ✓ {nome_campo} validado no HTML: {valor_atual}')
                                                            continue

                                                        print(
                                                            f'    ⚠ {nome_campo} nao preenchido/correto no HTML. '
                                                            f'Atual: "{valor_atual}" | Esperado: "{esperado}"'
                                                        )
                                                        elemento_campo.click()
                                                        elemento_campo.send_keys(Keys.CONTROL, 'a')
                                                        elemento_campo.send_keys(Keys.DELETE)
                                                        elemento_campo.send_keys(esperado)
                                                        time.sleep(0.3)

                                                        valor_revalidado = (elemento_campo.get_attribute('value') or '').strip()
                                                        if esperado.upper() in valor_revalidado.upper():
                                                            print(f'    ✓ {nome_campo} repreechido com sucesso no HTML')
                                                        else:
                                                            print(
                                                                f'    ⚠ {nome_campo} ainda divergente apos repreechimento. '
                                                                f'Atual: "{valor_revalidado}"'
                                                            )
                                                
                                                    salvar_upload = input("Salvar upload do documento? (s/n): ").strip().lower()

                                                    if salvar_upload == 's':
                                                        #BOTAO SALVAR UPLOAD

                                                        botao_salvar_upload = WebDriverWait(driver_daycoval, 20).until(
                                                            EC.element_to_be_clickable(
                                                                (By.ID, 'ctl00_cphZoomMain_tbarMain_tdConfirma')
                                                            )
                                                        )
                                                        botao_salvar_upload.click()
                                                        print('    ✓ Clique em Salvar upload realizado com sucesso')

                                                    pdb.set_trace() #debug após preenchimento dos campos do upload do documento no Daycoval, antes de salvar

                                                    botao_fechar_upload = WebDriverWait(driver_daycoval, 20).until(
                                                        EC.element_to_be_clickable(
                                                            (By.ID, 'ctl00_cphZoomMain_tbarMain_tdFechar')
                                                        )
                                                    )
                                                    botao_fechar_upload.click()
                                                    print('    ✓ Clique em Fechar upload realizado com sucesso')

                                                    if os.path.exists(arquivo_local):
                                                        try:
                                                            os.remove(arquivo_local)
                                                            print(f'    ✓ Arquivo removido da pasta: {arquivo_local}')
                                                        except Exception as e:
                                                            print(f'    ⚠ Não foi possível remover o arquivo local: {e}')
                                                    else:
                                                        print(f'    ⚠ Arquivo local não encontrado para remoção: {arquivo_local}')

                                                    driver_daycoval.get(DAYCOVAL_URL_CONSULTA)
                                                    print('    ✓ Retorno para URL de consulta realizado com sucesso')
                                                    pdb.set_trace() #debug após retorno para página de consulta no Daycoval, para verificar se número do processo foi preenchido corretamente

                                                except Exception as e:
                                                    pdb.set_trace()  # Debug: Captura de exceção durante o preenchimento do Daycoval
                                                    print(f'    ⚠ Não foi possível preencher número do processo no Daycoval: {e}')
                                            else:
                                                print('    ⚠ numero_processo vazio para preenchimento no Daycoval')
                                        else:
                                            print('    ⚠ URL Daycoval não encontrada no config_daycoval.json')
                                    else:
                                        print(f'    ✗ OCR falhou. Status: {resp_audiencia.status_code}')
                                else:
                                    print(f'    ✗ Falha ao baixar documento id_ged={id_ged_filtrados}')
                            else:
                                print('    ⚠ Nenhum arquivo com "ata de audiencia" encontrado no path')
                        else:
                            print(f'  [{idx}/{len(tarefas_processamento)}] id_processo={id_processo} sem tarefas ATA/IA AT/IA ATA')
                    
                    else:
                        print(f'  [{idx}/{len(tarefas_processamento)}] id_processo={id_processo} -> Grupo de trabalho não é 78, pulando busca por ATA.')
                
            else:
                print('✗ Busca não retornou dados')

                if driver_daycoval is not None:
                    try:
                        driver_daycoval.quit()
                    except Exception:
                        pass

        else:

            print ('Verificando titulos do BMG')

            #reembolsos BMG

            config_bmg = carregar_config_bmg()
            NUMERO_RECIBO = str(config_bmg.get('numero_recibo', '') or '').strip()
            data_inicial = str(config_bmg.get('data_inicial', '') or '').strip()
            data_final = str(config_bmg.get('data_final', '') or '').strip()

            if not NUMERO_RECIBO:
                print('⚠ ATENÇÃO: numero_recibo não encontrado em config_bmg.json')

            lancamentos = api_buscar_lancamentos_bmg(
                    numero_cc=1397,
                    titulo=NUMERO_RECIBO,
                    tipo_lancamento='Preposto'
                )

            lancamentos_processados = []
            if lancamentos:
                for lancamento in lancamentos:
                    ficha = lancamento.get('ficha') or lancamento.get('ficha_numero') or lancamento.get('numero_ficha')
                    if not ficha:
                        print(f'  ⚠ Lançamento sem ficha: {lancamento.get("documento", "<sem documento>")}')
                        continue

                    processos = api_buscar_processo_por_ficha(str(ficha))
                    if not processos:
                        print(f'  ⚠ Nenhum processo encontrado para ficha: {ficha}')
                        continue

                    processo_encontrado = processos[0]
                    numero_integracao = processo_encontrado.get('numero_integracao', '')
                    numero_processo_pesq = re.sub(r'[^0-9]', '', str(processo_encontrado.get('numero_processo', '')))
                    data_lancamento_formatada = formatar_data_lancamento(lancamento.get('data_lancamento', ''))

                    lancamentos_processados.append({
                        'numero_integracao': numero_integracao,
                        'numero_processo_pesq': numero_processo_pesq,
                        'data_lancamento': data_lancamento_formatada,
                    })

                print(f'✓ {len(lancamentos_processados)} lançamento(s) processado(s) com ficha e processo encontrado(s)')
                pprint(lancamentos_processados)
                
                atualizar_planilha_preposto(lancamentos_processados, DESTINO_PLANILHA_BMG)
                atualizar_recibo_padrao(NUMERO_RECIBO)

                #valor total lancamentos
                valor_total_lancamentos = len(lancamentos_processados) * 75
                str_valor_total_lancamentos = f'R$ {valor_total_lancamentos:.2f}'.replace('.', ',')
                print(f'✓ Valor total dos lançamentos BMG: {str_valor_total_lancamentos}')
                fill_descritivo_pdf(lancamentos_processados, NUMERO_RECIBO, str_valor_total_lancamentos)
                merge_descritivo_e_recibo_pdf()

                # logando no BMG e executando fluxo web
                driver = open_chrome_browser()
                login_web_exyon_bmg(driver)
                selecionar_unidade(driver)
                executar_script_menu(driver)
                executar_script_menu_lateral(driver)
                executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                tentativa_formatacao = False

                processar_pdf = anexar_pdfs_formulario(driver, tentativa = -1)

                tentativa = -1

                while processar_pdf['retorno'] == False:

                    if tentativa == -1:
                        for processo_alternativo in processar_pdf['processos_alternativos']:

                            numero_integracao = processo_alternativo.get('numero_integracao')
                            numero_processo_formatado = processo_alternativo.get('numero_processo_pesq', '')

                            for lancamento in lancamentos_processados:
                                if lancamento.get('numero_integracao') == numero_integracao:
                                    lancamento['numero_processo_pesq'] = numero_processo_formatado
                                    break

                            print(f'⚠ Processo alternativo encontrado: {processo_alternativo}')

                    atualizar_planilha_preposto(lancamentos_processados, DESTINO_PLANILHA_BMG)

                    driver.refresh()

                    executar_script_menu(driver)
                    executar_script_menu_lateral(driver)
                    executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                    executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                    tentativa += 1
                    processar_pdf = anexar_pdfs_formulario(driver, tentativa)

                    removidos = []
                    numeros_processo_para_remover = set()

                    if processar_pdf['retorno'] == False: 

                        processar_pdf = anexar_pdfs_formulario(driver, tentativa)

                        for processo_alternativo in processar_pdf['processos_alternativos']:
                            if processo_alternativo.get('remover') == 'nao':
                                numero_processo_pesq = str(processo_alternativo.get('numero_processo_pesq', '')).strip()
                                print(f'⚠ Processo alternativo encontrado (não removido): {processo_alternativo}')
                                if numero_processo_pesq:
                                    for lancamento in lancamentos_processados:
                                        if lancamento.get('numero_integracao') == processo_alternativo.get('numero_integracao'):
                                            lancamento['numero_processo_pesq'] = numero_processo_pesq
                                            break
                        
                        atualizar_planilha_preposto(lancamentos_processados, DESTINO_PLANILHA_BMG)

                        driver.refresh()

                        executar_script_menu(driver)
                        executar_script_menu_lateral(driver)
                        executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                        executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                        processar_pdf = anexar_pdfs_formulario(driver, tentativa)

                        pdb.set_trace() #debug após tentativa de anexar PDFs no BMG, antes de verificar retorno

                        # for processo_alternativo in processar_pdf['processos_alternativos']:
                        #     if processo_alternativo.get('remover') == 'sim':
                        #         numero_processo_pesq = str(processo_alternativo.get('numero_processo_pesq', '')).strip()
                        #         print(f'⚠ Processo alternativo encontrado: {processo_alternativo}')
                        #         removidos.append(processo_alternativo)
                        #         if numero_processo_pesq:
                        #             numeros_processo_para_remover.add(numero_processo_pesq)
                            
                        #     elif processo_alternativo.get('remover') == 'nao':
                        #         numero_processo_pesq = str(processo_alternativo.get('numero_processo_pesq', '')).strip()
                        #         print(f'⚠ Processo alternativo encontrado (não removido): {processo_alternativo}')
                        #         if numero_processo_pesq:
                        #             for lancamento in lancamentos_processados:
                        #                 if lancamento.get('numero_integracao') == processo_alternativo.get('numero_integracao'):
                        #                     lancamento['numero_processo_pesq'] = numero_processo_pesq
                        #                     break


                        # if numeros_processo_para_remover:
                        #     lancamentos_processados = [
                        #         lancamento for lancamento in lancamentos_processados
                        #         if str(lancamento.get('numero_processo_pesq', '')).strip() not in numeros_processo_para_remover
                        #     ]

                        # if removidos:
                        #     caminho_removidos = os.path.join(BASE_PATH, 'removidos', 'banco-bmg', f'bmg_removidos_{NUMERO_RECIBO}.json')
                        #     os.makedirs(os.path.dirname(caminho_removidos), exist_ok=True)
                        #     try:
                        #         with open(caminho_removidos, 'w', encoding='utf-8') as f_removidos:
                        #             json.dump(removidos, f_removidos, ensure_ascii=False, indent=2)
                        #         print(f'✓ Removidos salvos em: {caminho_removidos}')
                        #     except Exception as e:
                        #         print(f'⚠ Não foi possível salvar {caminho_removidos}: {e}')

                        #valor total lancamentos
                        # valor_total_lancamentos = len(lancamentos_processados) * 75
                        # str_valor_total_lancamentos = f'R$ {valor_total_lancamentos:.2f}'.replace('.', ',')
                        # print(f'✓ Valor total dos lançamentos BMG: {str_valor_total_lancamentos}')
                        # fill_descritivo_pdf(lancamentos_processados, NUMERO_RECIBO, str_valor_total_lancamentos)
                        # merge_descritivo_e_recibo_pdf()

                        # driver.refresh()

                        # executar_script_menu(driver)
                        # executar_script_menu_lateral(driver)
                        # executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                        # executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                        # tentativa += 1
                        # processar_pdf = anexar_pdfs_formulario(driver, tentativa)

                        if processar_pdf['retorno'] == False:

                            for processo_alternativo in processar_pdf['processos_alternativos']:
                                numero_processo_pesq = str(processo_alternativo.get('numero_processo_pesq', '')).strip()
                                print(f'⚠ Processo alternativo encontrado: {processo_alternativo}')
                                if numero_processo_pesq:
                                    for lancamento in lancamentos_processados:
                                        if lancamento.get('numero_integracao') == processo_alternativo.get('numero_integracao'):
                                            lancamento['numero_processo_pesq'] = numero_processo_pesq
                                            break
                            
                            atualizar_planilha_preposto(lancamentos_processados, DESTINO_PLANILHA_BMG)

                            #valor total lancamentos
                            valor_total_lancamentos = len(lancamentos_processados) * 75
                            str_valor_total_lancamentos = f'R$ {valor_total_lancamentos:.2f}'.replace('.', ',')
                            print(f'✓ Valor total dos lançamentos BMG: {str_valor_total_lancamentos}')
                            fill_descritivo_pdf(lancamentos_processados, NUMERO_RECIBO, str_valor_total_lancamentos)
                            merge_descritivo_e_recibo_pdf()

                            driver.refresh()

                            executar_script_menu(driver)
                            executar_script_menu_lateral(driver)
                            executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                            executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                            tentativa += 1
                            processar_pdf = anexar_pdfs_formulario(driver, tentativa)

                pdb.set_trace() #debug após tentativa de anexar PDFs no BMG, antes de verificar retorno
                if processar_pdf['retorno'] == True:
                        print('✓ Todos os PDFs processados com sucesso')

                        #pdb.set_trace() #debug acompanhar

                        #valor total lancamentos
                        valor_total_lancamentos = len(lancamentos_processados) * 75
                        str_valor_total_lancamentos = f'R$ {valor_total_lancamentos:.2f}'.replace('.', ',')
                        fill_descritivo_pdf(lancamentos_processados, NUMERO_RECIBO, str_valor_total_lancamentos)
                        merge_descritivo_e_recibo_pdf()

                        driver.refresh()

                        executar_script_menu(driver)
                        executar_script_menu_lateral(driver)
                        executar_script_classificacao(driver, script_select = "select('35','DESPESA');")
                        executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                        print(f'✓ Valor total dos lançamentos BMG: {str_valor_total_lancamentos}')
                        valor_somado = str_valor_total_lancamentos.replace('R$', '').replace(' ', '')

                        finalizar_processo(driver, valor_somado)

                        # verifica lançamentos para garantir que foram baixados
                        baixado = False
                        while not baixado:

                            print('\nVerificando lançamentos após submissão...')
                            
                            driver.refresh()

                            # Executar script de menu
                            executar_script_menu(driver)

                            # Executar script do menu lateral
                            executar_script_menu_lateral(driver)

                            # Executar script de classificação financeira
                            executar_script_classificacao(driver)

                            #Executa preenchimento formulario
                            executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                            #verifica se os lançamentos estão baixados, se não estiverem, tenta novamente (pode ser necessário formatar processo alternativo)
                            baixado = verificar_lancamentos(driver)

                            if baixado:
                                print('\n✓ Lançamentos verificados como baixados!')
                                baixado = True

                            else:
                                print('\n⚠ Lançamentos ainda não baixados, tentando novamente...')
                                anexar_pdfs_formulario(driver)

                                finalizar_processo_envio = finalizar_processo(driver, valor_somado)
                                
                                while finalizar_processo_envio == False:

                                    print('\n⚠ Falha ao finalizar processo, tentando novamente...')
                                    time.sleep(2)
                                    driver.refresh()

                                    # Executar script de menu
                                    executar_script_menu(driver)

                                    # Executar script do menu lateral
                                    executar_script_menu_lateral(driver)

                                    # Executar script de classificação financeira
                                    executar_script_classificacao(driver)

                                    #Executa preenchimento formulario
                                    executar_preenchimento_formulario(driver, str_valor_total_lancamentos, path = DESTINO_PLANILHA_BMG, NUMERO_RECIBO = NUMERO_RECIBO)

                                    anexar_pdfs_formulario(driver)

                                    finalizar_processo_envio = finalizar_processo(driver, str_valor_total_lancamentos)

                        zerar_config_bmg()
                        driver.quit()

            else:
                print('✗ Nenhum lançamento BMG retornado')

        api_logout()
        print('✓ Logout da API CPJ realizado')
    else:
        print('✗ Falha ao autenticar na API CPJ')


if __name__ == '__main__':
    main()
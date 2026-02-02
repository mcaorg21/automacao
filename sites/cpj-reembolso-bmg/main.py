import time
import os
import subprocess
import csv
import json
import pdb
import shutil
import sys
from datetime import datetime, timedelta
from pathlib import Path
import traceback
from pywinauto import Application, Desktop
from pywinauto.keyboard import send_keys
from pywinauto.mouse import click
import pyautogui
from PyPDF2 import PdfMerger
import xlrd
from xlwt import Workbook, XFStyle
from xlutils.copy import copy as xl_copy
from docx import Document
from docx2pdf import convert
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

# Configurações
CPJ_PATH = r'C:\CPJ3C_Client\cpj3cclient.exe'
LOGIN = 'MOAO'
PASSWORD = 'LM0G'
PLANILHA_PATH = r'C:\www\automacao\sites\cpj-reembolso-bmg\planilha'
CSV_PATH = os.path.join(PLANILHA_PATH, 'importados.csv')
JSON_PATH = os.path.join(PLANILHA_PATH, 'importados.json')

# Caminhos das imagens para reconhecimento
IMAGES_PATH = r'C:\www\automacao\sites\cpj-reembolso-bmg\images'
IMAGE_MENU_ICON = os.path.join(IMAGES_PATH, 'menu_icon.png')
IMAGE_MENU_CONTA_CORRENTE = os.path.join(IMAGES_PATH, 'menu_conta_corrente.png')
IMAGE_PRINTER_BUTTON = os.path.join(IMAGES_PATH, 'impressora.png')
IMAGE_ANALISE_LANCAMENTO = os.path.join(IMAGES_PATH, 'analise_lancamento.png')
IMAGE_RELATORIO_CONTA = os.path.join(IMAGES_PATH, 'relatorio_conta.png')
IMAGE_PASTA = os.path.join(IMAGES_PATH, 'pasta.png')

#caminho dos pdfs temporarios
PDF_FOLDER = f'C:\\Users\\{os.environ.get("USERNAME")}\\Downloads\\pdf_cpj'

# Caminho da planilha modelo
PLANILHA_MODELO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'planilha_modelo', 'bmg', 'planilha_modelo_custas.xls')
PLANILHA_MODELO_PATH_2 = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'planilha_modelo', 'bmg', 'planilha_modelo_custas_2.xls')

# Caminho do documento modelo
DOCX_MODELO = r'C:\www\automacao\sites\cpj-reembolso-bmg\documento_modelo\bmg\descritivo_custas.docx'
DOCX_OUTPUT = r'C:\www\automacao\sites\cpj-reembolso-bmg\documento_modelo\bmg\descritivo_custas_preenchido.docx'
PDF_OUTPUT = r'C:\www\automacao\sites\cpj-reembolso-bmg\documento_modelo\bmg\descritivo_custas_preenchido.pdf'

# Caminho do arquivo de configuração
CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')

# Carrega parâmetros: se tiver args usa args, senão lê do config.json
if len(sys.argv) >= 4:
    # Modo linha de comando
    NUMERO_RECIBO = sys.argv[1]
    DATA_INICIAL_PESQUISA = datetime.strptime(sys.argv[2], "%d/%m/%Y")
    DATA_FINAL_PESQUISA = datetime.strptime(sys.argv[3], "%d/%m/%Y")
else:
    # Modo config.json
    print("Nenhum argumento fornecido. Lendo configurações de config.json...")
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        config = json.load(f)
    NUMERO_RECIBO = config['numero_recibo']
    DATA_INICIAL_PESQUISA = datetime.strptime(config['data_inicial'], "%d/%m/%Y")
    DATA_FINAL_PESQUISA = datetime.strptime(config['data_final'], "%d/%m/%Y")

if not NUMERO_RECIBO.isdigit():
    raise ValueError("INFORME O NÚMERO DO RECIBO VÁLIDO")

# regra básica de sanidade
if DATA_INICIAL_PESQUISA > DATA_FINAL_PESQUISA:
    raise ValueError("DATA_INICIAL_PESQUISA não pode ser maior que DATA_FINAL_PESQUISA")

DATA_INICIAL_PESQUISA = DATA_INICIAL_PESQUISA.strftime('%d/%m/%Y')
DATA_FINAL_PESQUISA = DATA_FINAL_PESQUISA.strftime('%d/%m/%Y')

print(f"Configurações carregadas:")
print(f"  - Número Recibo: {NUMERO_RECIBO}")
print(f"  - Data Inicial: {DATA_INICIAL_PESQUISA}")
print(f"  - Data Final: {DATA_FINAL_PESQUISA}")


def valida_data(data_str: str) -> datetime:
    try:
        return datetime.strptime(data_str, "%d/%m/%Y")
    except ValueError:
        raise ValueError(
            f"Data inválida: {data_str}. Formato obrigatório: dd/mm/YYYY"
        )

def sleep_with_countdown(seconds: int, message: str = 'Aguardando'):
    """Aguarda com countdown visual"""
    for i in range(seconds, 0, -1):
        print(f'\r{message}: {i}s restantes...', end='', flush=True)
        time.sleep(1)
    print(f'\r{message}: concluído!          ')

def click_image(image_path: str, confidence: float = 0.8, timeout: int = 10, description: str = ''):
    """Procura uma imagem na tela e clica no centro dela
    
    Args:
        image_path: Caminho para o arquivo de imagem
        confidence: Confiança mínima para encontrar a imagem (0.0 a 1.0) - requer OpenCV
        timeout: Tempo máximo de espera em segundos
        description: Descrição do elemento (para logs)
    
    Returns:
        bool: True se clicou, False se não encontrou
    """
    if not os.path.exists(image_path):
        print(f'✗ Arquivo de imagem não encontrado: {image_path}')
        return False
    
    desc = description or os.path.basename(image_path)
    print(f'Procurando imagem: {desc}...')
    
    start_time = time.time()
    while time.time() - start_time < timeout:
        try:
            # Tenta com confidence se OpenCV estiver disponível
            try:
                location = pyautogui.locateCenterOnScreen(image_path, confidence=confidence)
            except TypeError:
                # OpenCV não disponível, usa sem confidence
                location = pyautogui.locateCenterOnScreen(image_path)
            
            if location:
                print(f'✓ Imagem encontrada em: {location}')
                pyautogui.click(location)
                time.sleep(0.5)
                return True
        except pyautogui.ImageNotFoundException:
            pass
        except Exception as e:
            # Mostra erro apenas na primeira tentativa
            if time.time() - start_time < 1:
                print(f'Aviso: {e}')
                print('Continuando sem parâmetro confidence...')
        time.sleep(0.5)
    
    print(f'✗ Imagem não encontrada após {timeout}s: {desc}')
    return False

def wait_for_element(app, class_name: str = None, title: str = None, timeout: int = 60, check_interval: float = 0.5):
    """Aguarda até que um elemento apareça na tela
    
    Args:
        app: Objeto Application do pywinauto
        class_name: Nome da classe do elemento (ex: 'TsPanel')
        title: Título do elemento
        timeout: Tempo máximo de espera em segundos
        check_interval: Intervalo entre verificações em segundos
        
    Returns:
        True se encontrou o elemento, False se timeout
    """
    print(f'\nAguardando elemento aparecer...')
    if class_name:
        print(f'  classe: {class_name}')
    if title:
        print(f'  título: {title}')
    print(f'  timeout: {timeout}s')
    
    elapsed = 0
    while elapsed < timeout:
        try:
            dlg = app.top_window()
            
            # Tenta encontrar o elemento
            if class_name and title:
                element = dlg.child_window(class_name=class_name, title=title)
            elif class_name:
                element = dlg.child_window(class_name=class_name)
            elif title:
                element = dlg.child_window(title=title)
            else:
                raise ValueError('Especifique class_name ou title')
            
            # Verifica se existe
            if element.exists():
                print(f'\r✓ Elemento encontrado após {elapsed:.1f}s!          ')
                return True
                
        except Exception:
            pass
        
        # Mostra progresso
        remaining = timeout - elapsed
        print(f'\rProcurando... {remaining:.0f}s restantes', end='', flush=True)
        
        time.sleep(check_interval)
        elapsed += check_interval
    
    print(f'\r✗ Timeout: elemento não encontrado após {timeout}s')
    return False

def clean_planilha_folder():
    """Limpa a pasta planilha"""
    try:
        print('\nLimpando pasta planilha...')
        
        if not os.path.exists(PLANILHA_PATH):
            os.makedirs(PLANILHA_PATH)
            print(f'✓ Pasta criada: {PLANILHA_PATH}')
            return
        
        for file in os.listdir(PLANILHA_PATH):
            file_path = os.path.join(PLANILHA_PATH, file)
            if os.path.isfile(file_path):
                os.unlink(file_path)
                print(f'✓ Arquivo removido: {file}')
        
        print('✓ Pasta planilha limpa com sucesso!')
    except Exception as e:
        print(f'✗ Erro ao limpar pasta planilha: {e}')
        raise

def close_controle_processos_windows():
    """Fecha o aplicativo CPJ se estiver em execução"""
    try:
        print('\nFechando aplicativo CPJ...')
        
        # Tenta conectar ao processo existente
        try:
            app = Application(backend='win32').connect(path=CPJ_PATH, timeout=5)
            print('✓ Aplicativo CPJ encontrado em execução')
            time.sleep(2)
            # Tenta fechar normalmente
            try:
                app.kill()
                print('✓ Aplicativo CPJ fechado com sucesso')
            except Exception as e:
                print(f'⚠ Erro ao fechar: {e}')
                
        except Exception:
            print('✓ Aplicativo CPJ não está em execução')
            
    except Exception as e:
        print(f'✗ Erro ao fechar aplicativo: {e}')

def clean_pdf_merge_folder():
    """Limpa a pasta pdf_merge"""
    try:
        merge_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdf_merge')
        print(f'\nLimpando pasta pdf_merge: {merge_folder}...')
        
        if not os.path.exists(merge_folder):
            os.makedirs(merge_folder)
            print(f'✓ Pasta criada: {merge_folder}')
            return
        
        for item in os.listdir(merge_folder):
            item_path = os.path.join(merge_folder, item)
            try:
                if os.path.isfile(item_path):
                    os.unlink(item_path)
                    print(f'  ✓ Arquivo removido: {item}')
                elif os.path.isdir(item_path):
                    shutil.rmtree(item_path)
                    print(f'  ✓ Pasta removida: {item}')
            except Exception as e:
                print(f'  ✗ Erro ao remover {item}: {e}')
        
        print('✓ Pasta pdf_merge limpa com sucesso!')
    except Exception as e:
        print(f'✗ Erro ao limpar pasta pdf_merge: {e}')
        raise

def open_cpj_application():
    """Abre o aplicativo CPJ"""
    try:
        print('Abrindo aplicativo CPJ...')
        # Inicia o aplicativo
        app = Application(backend='win32').start(CPJ_PATH)
        print('✓ Aplicativo CPJ iniciado!')
        
        print('Aguardando aplicativo carregar...')
        time.sleep(5)
        
        # Retorna o objeto Application
        return app
        
    except Exception as e:
        print(f'✗ Erro ao abrir aplicativo: {e}')
        # Tenta conectar se já estiver rodando
        try:
            print('Tentando conectar ao processo existente...')
            app = Application(backend='win32').connect(path=CPJ_PATH, timeout=10)
            return app
        except:
            raise

def close_cpj_application(app):
    """Fecha o aplicativo CPJ normalmente"""
    try:
        print('\nFechando aplicativo CPJ...')
        dlg = app.top_window()
        dlg.close()
        print('✓ Aplicativo CPJ fechado!')
    except Exception as e:
        print(f'✗ Erro ao fechar aplicativo: {e}')

def perform_login(app):
    """Realiza o login no sistema"""
    try:
        print('\nIniciando processo de login...')
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Tenta acessar os campos TsEdit diretamente pela classe
        try:
            print('Tentando acessar campos TsEdit...')
            # Busca todos os controles TsEdit
            edits = dlg.children(class_name="TsEdit")
            
            if len(edits) >= 2:
                print(f'✓ Encontrados {len(edits)} campos TsEdit')
                
                # Primeiro campo = Login
                print('Digitando login no primeiro campo...')
                edits[0].set_focus()
                time.sleep(0.2)
                edits[0].set_edit_text(LOGIN)
                time.sleep(0.3)
                
                # Segundo campo = Senha
                print('Digitando senha no segundo campo...')
                edits[1].set_focus()
                time.sleep(0.2)
                edits[1].set_edit_text(PASSWORD)
                time.sleep(0.3)
                
                # O botão Ok é um TsPanel, não um Button padrão
                # Método 1: Tentar encontrar o TsPanel com texto "Ok"
                print('Procurando botão Ok...')
                try:
                    # Método 2: Simplesmente pressionar Enter (mais confiável)
                    print('Usando Enter para confirmar...')
                    dlg.type_keys('{ENTER}')
                    print('✓ Enter pressionado')
                    
                except:

                    ok_button = dlg.child_window(title="Ok", class_name="TsSplitView")
                    ok_button.click()
                    print('✓ Botão Ok clicado (TsSplitView)')
                   
                
            else:
                raise Exception(f'Apenas {len(edits)} campos TsEdit encontrados')
                
        except Exception as e:
            print(f'Método direto falhou: {e}')
            print('Usando método alternativo com type_keys...')
            
            # Fallback: usa digitação sequencial
            dlg.set_focus()
            time.sleep(0.5)
            dlg.type_keys(LOGIN, with_spaces=True)
            time.sleep(0.2)
            dlg.type_keys('{TAB}')
            time.sleep(0.2)
            dlg.type_keys(PASSWORD, with_spaces=True)
            time.sleep(0.2)
            dlg.type_keys('{ENTER}')
        
        print('✓ Login realizado com sucesso!')
        
        # Aguarda até que TsSplitView apareça (indica que o sistema carregou)
        if wait_for_element(app, class_name="TsSplitView", timeout=60):
            print('✓ Sistema carregado!')
        else:
            print('⚠ Timeout ao aguardar sistema carregar, continuando mesmo assim...')
        
        return dlg
        
    except Exception as e:
        print(f'✗ Erro ao realizar login: {e}')
        import traceback
        traceback.print_exc()
        
        # Debug: mostra controles disponíveis
        try:
            print('\n=== DEBUG: Controles disponíveis ===')
            dlg.print_control_identifiers()
        except:
            pass
        raise

def click_menu(app):
    """Clica no menu e navega para CC de clientes/fornecedores"""
    try:
        time.sleep(20)
        print('\nAbrindo menu...')
        
        # Usa Desktop para procurar TsSplitView em todas as janelas
        desktop = Desktop(backend='win32')
        all_windows = desktop.windows()
        
        menu_found = None
        print('Procurando TsSplitView em todas as janelas...')
        
        for window in all_windows:
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsSplitView' and desc.is_visible():
                            print(f'✓ TsSplitView encontrado na janela: "{window.window_text()}"')
                            menu_found = desc
                            break
                    except:
                        pass
                if menu_found:
                    break
            except:
                pass
        click(coords=(38, 31))
        time.sleep(1)
        if not menu_found:
            print('✗ TsSplitView não encontrado!')
            raise Exception('TsSplitView não encontrado')
        
        # Clica no TsSplitView encontrado
        print('Clicando no TsSplitView...')
        menu_found.click_input()
        print('✓ TsSplitView clicado!')
        time.sleep(1)
        
        # Clica na posição específica do item do menu
        print('Clicando em CC de clientes/fornecedores (164, 503)...')
        click(coords=(164, 503))
        print('✓ Item clicado!')
        
        # Aguarda a tela carregar
        time.sleep(2)
        print('✓ Menu navegado com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao clicar no menu: {e}')
        import traceback
        traceback.print_exc()
        raise

def click_printer(app):
    """Pressiona F5 para acionar a impressão"""
    try:
        print('\nAguardando tela carregar...')
        time.sleep(2)
        
        # Pressiona F5 (botão de impressão)
        print('Pressionando F5 (botão de impressão)...')
        send_keys('{F5}')
        print('✓ F5 pressionado!')
        time.sleep(2)
        
        return True
            
    except Exception as e:
        print(f'✗ Erro ao listar elementos: {e}')
        import traceback
        traceback.print_exc()
        raise

def navigate_to_analise_lancamento(app):
    """Navega para Análise de Lançamento procurando TsGroupButton"""
    try:
        print('\nNavegando para Análise de Lançamento...')
        time.sleep(2)
        
        # Usa Desktop para procurar em todas as janelas
        desktop = Desktop(backend='win32')
        
        # Procura o TsGroupButton em TODAS as janelas abertas
        print('Procurando TsGroupButton em todas as janelas...')
        
        all_windows = desktop.windows()
        print(f'Total de janelas: {len(all_windows)}')
        
        found_button = None
        
        for window in all_windows:
            try:
                window_text = window.window_text()
                window_class = window.class_name()
                
                # Procura em cada janela
                descendants = window.descendants()
                
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsGroupButton':
                            btn_text = desc.window_text()
                            print(f'  TsGroupButton encontrado: "{btn_text}" na janela "{window_text}" (classe: {window_class})')
                            
                            # Tenta diferentes formas de comparação
                            if ('Análise de lançamentos' in btn_text or 
                                'Analise de lancamentos' in btn_text or
                                'análise de lançamentos' in btn_text.lower()):
                                found_button = desc
                                print(f'✓ Botão MATCH encontrado: "{btn_text}"')
                                print(f'  Janela: {window_text}')
                                print(f'  Classe da janela: {window_class}')
                                print(f'  Retângulo: {desc.rectangle()}')
                                print(f'  Visível: {desc.is_visible()}')
                                print(f'  Habilitado: {desc.is_enabled()}')
                                break
                    except:
                        pass
                
                if found_button:
                    break
                    
            except:
                pass
        
        if not found_button:
            print('✗ TsGroupButton "Análise de lançamentos" não encontrado em nenhuma janela')
            raise Exception('TsGroupButton não encontrado')
        
        # Tenta clicar no botão encontrado
        print('\nTentando clicar no botão...')
        try:
            found_button.click_input()
            print('✓ Botão clicado com click_input()!')
            time.sleep(2)
        except Exception as e:
            print(f'click_input() falhou: {e}')
            try:
                print('Tentando click()...')
                found_button.click()
                print('✓ Botão clicado com click()!')
                time.sleep(2)
            except Exception as e2:
                print(f'click() falhou: {e2}')
                # Último recurso: clica na posição
                rect = found_button.rectangle()
                center_x = (rect.left + rect.right) // 2
                center_y = (rect.top + rect.bottom) // 2
                print(f'Tentando clicar na posição ({center_x}, {center_y})...')
                click(coords=(center_x, center_y))
                print('✓ Clicado na posição!')
                time.sleep(2)
        
        # Agora procura e clica no botão OK (TsBitBtn)
        print('\nProcurando botão OK (TsBitBtn)...')
        
        for window in all_windows:
            try:
                descendants = window.descendants()
                
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsBitBtn':
                            btn_text = desc.window_text()
                            if 'OK' in btn_text or 'Ok' in btn_text:
                                print(f'✓ Botão OK encontrado: "{btn_text}"')
                                print(f'  Retângulo: {desc.rectangle()}')
                                desc.click_input()
                                print('✓ Botão OK clicado!')
                                time.sleep(2)
                                return True
                    except:
                        pass
            except:
                pass
        
        print('⚠ Botão OK não encontrado, mas continuando...')
        return True
        
    except Exception as e:
        print(f'✗ Erro ao navegar para Análise de Lançamento: {e}')
        import traceback
        traceback.print_exc()
        raise

def navigate_to_relatorio_conta(app):
    """Navega para Relatório da Conta"""
    try:
        print('\nNavegando para Relatório da Conta...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Duas setas para baixo
        dlg.type_keys('{DOWN}')
        time.sleep(0.5)
        dlg.type_keys('{DOWN}')
        print('✓ Relatório da Conta selecionado!')
        time.sleep(1)
        
        print('Confirmando com Enter...')
        dlg.type_keys('{ENTER}')
        print('✓ Relatório da Conta confirmado!')
        time.sleep(2)
        
    except Exception as e:
        print(f'✗ Erro ao navegar para Relatório da Conta: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_cliente_and_dates(app):
    """Preenche dados do cliente e datas"""
    try:
        print('\nPreenchendo dados do cliente...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Cliente
        print('Digitando número do cliente (1397)...')
        dlg.type_keys('1397')
        time.sleep(0.5)
        
        # Tab 2x
        print('Navegando para campo de data...')
        dlg.type_keys('{TAB}{TAB}')
        time.sleep(0.5)
        
        # Data de ontem
        #data_ontem = input("Digite a data inicial (DD/MM/AAAA) ou deixe vazio para usar ontem: ")
        #data_ontem = None
        data_ontem = DATA_INICIAL_PESQUISA
        if not data_ontem:
            ontem = datetime.now() - timedelta(days=4)
            data_ontem = ontem.strftime('%d/%m/%Y')
        print(f'Digitando data inicial - ontem ({data_ontem})...')
        dlg.type_keys(data_ontem)
        time.sleep(0.5)
        
        # Tab
        dlg.type_keys('{TAB}')
        time.sleep(0.5)
        
        # Data de hoje
        #data_hoje = input("Digite a data final (DD/MM/AAAA) ou deixe vazio para usar hoje: ")
        #data_hoje = None
        data_hoje = DATA_FINAL_PESQUISA
        if not data_hoje:
            hoje = datetime.now()
            data_hoje = hoje.strftime('%d/%m/%Y')
        print(f'Digitando data final - hoje ({data_hoje})...')
        dlg.type_keys(data_hoje)
        time.sleep(0.5)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        dlg.type_keys('{ENTER}')
        print('✓ Dados preenchidos e confirmados!')
        
        # Aguarda até que o botão Salvar apareça (indica que o relatório foi processado)
        if wait_for_element(app, class_name="TsBitBtn", title="Salvar", timeout=120):
            print('✓ Relatório processado e pronto para exportação!')
        else:
            print('⚠ Timeout ao aguardar relatório processar, continuando mesmo assim...')
        
    except Exception as e:
        print(f'✗ Erro ao preencher dados: {e}')
        import traceback
        traceback.print_exc()
        raise

def export_file(app):
    """Exporta o arquivo CSV"""
    try:
        print('\nIniciando exportação do arquivo...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Clica no botão Salvar
        print('Clicando no botão Salvar...')
        salvar_button = dlg.child_window(class_name="TsBitBtn", title="Salvar")
        salvar_button.click_input()
        print('✓ Botão Salvar clicado!')
        time.sleep(5)
        
        # Digita o caminho do arquivo (janela nativa do Windows)
        print('Digitando caminho do arquivo...')
        send_keys(CSV_PATH)
        print(f'✓ Caminho digitado: {CSV_PATH}')
        time.sleep(1)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        send_keys('{ENTER}')
        print('✓ Exportação confirmada!')
        time.sleep(2)
        
        print('✓ Arquivo exportado com sucesso!')
        time.sleep(2)
        
    except Exception as e:
        print(f'✗ Erro ao exportar arquivo: {e}')
        import traceback
        traceback.print_exc()
        raise

def close_window(app):
    """Fecha a janela atual"""
    try:
        print('\nFechando janela...')
        time.sleep(2)
        
        # Alt+F4 para fechar a janela
        print('Pressionando Alt+F4...')
        send_keys('%{F4}')
        print('✓ Alt+F4 pressionado!')
        time.sleep(2)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        send_keys('{ENTER}')
        print('✓ Enter pressionado!')
        time.sleep(2)
        
        print('✓ Janela fechada com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao fechar janela: {e}')
        import traceback
        traceback.print_exc()
        raise

def to_snake_case(text: str) -> str:
    """Converte string para snake_case"""
    import unicodedata
    
    # Valida entrada
    if text is None or not isinstance(text, str):
        return ''
    
    # Remove acentos
    text = unicodedata.normalize('NFD', text)
    text = ''.join(char for char in text if unicodedata.category(char) != 'Mn')
    
    # Converte para lowercase e substitui
    text = text.lower()
    text = text.replace('.', '_').replace(' ', '_')
    text = ''.join(c if c.isalnum() or c == '_' else '' for c in text)
    text = '_'.join(filter(None, text.split('_')))
    
    return text

def process_csv():
    """Processa o CSV e gera JSON"""
    try:
        print('\nProcessando arquivo CSV...')
        
        # Lê o CSV
        with open(CSV_PATH, 'r', encoding='latin1') as f:
            reader = csv.DictReader(f, delimiter=';')
            records = list(reader)
        
        if len(records) == 0:
            print('⚠ Aviso: CSV está vazio!')

            # Lê o CSV
            with open(CSV_PATH, 'r', encoding='latin1') as f:
                reader = csv.DictReader(f, delimiter=',')
                records = list(reader)

        print(f'✓ CSV lido: {len(records)} registros encontrados')
        
        # Remove duplicados por "PRO.Número de integração"
        grouped = {}
        registros_filtrados_correios = 0
        
        for record in records:

            numero_integracao = record.get('PRO.Número de integração')
            historico = record.get('Histórico', '')
            valor_credito_conta = record.get('Crédito na moeda','0')

            # Filtra registros com valor 0,00
            if '0,00' not in valor_credito_conta.strip():
                continue
            
            # Filtra registros que contenham "CORREIOS" no histórico
            if 'CORREIOS' in historico.upper():
                registros_filtrados_correios += 1
                continue

            if numero_integracao and numero_integracao not in grouped:
                grouped[numero_integracao] = record
        
        print(f'✓ Removidas duplicatas: {len(grouped)} registros únicos')
        
        # Calcula soma e converte para snake_case
        total_valor = 0
        registros_snake_case = []
        
        for record in grouped.values():
            valor_str = record.get('Valor em R$', '0')
            valor = float(valor_str.strip().replace(',', '.'))
            total_valor += valor
            
            # Converte chaves para snake_case
            record_snake = {}
            for key, value in record.items():
                snake_key = to_snake_case(key)
                
                # Remove caracteres especiais do número do processo
                if snake_key == 'pro_numero_do_processo':
                    value = ''.join(c for c in value if c.isdigit())
                
                record_snake[snake_key] = value
            
            registros_snake_case.append(record_snake)
        
        # Formata valor
        valor_formatado = f"{total_valor:.2f}".replace('.', ',')
        
        # Cria objeto final
        resultado = {
            'contagem': len(grouped),
            'valor_somado': valor_formatado,
            'registros': registros_snake_case
        }
        
        print(f'✓ Contagem de registros únicos: {resultado["contagem"]}')
        print(f'✓ Valor total somado: R$ {resultado["valor_somado"]}')
        
        # Salva JSON
        with open(JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump(resultado, f, indent=2, ensure_ascii=False)
        
        print(f'✓ JSON salvo em: {JSON_PATH}')
        return resultado
        
    except Exception as e:
        print(f'✗ Erro ao processar CSV: {e}')
        raise

def open_search_area(app):
    """Abre a área de busca"""
    try:
        print('\nAbrindo área de busca...')
        
        # Clique 1
        click(coords=(91, 270))
        print('✓ Clicado em (91, 270)')
        time.sleep(2)
        
        # Usa Desktop para procurar o TsComboBox em todas as janelas
        print('Procurando TsComboBox...')
        desktop = Desktop(backend='win32')
        
        combo_found = None
        for window in desktop.windows():
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsComboBox':
                            print(f'✓ TsComboBox encontrado na janela: "{window.window_text()}"')
                            combo_found = desc
                            break
                    except:
                        pass
                if combo_found:
                    break
            except:
                pass
        
        if combo_found:
            print('Clicando no TsComboBox...')
            combo_found.click_input()
            print('✓ TsComboBox clicado!')
            time.sleep(0.5)
        else:
            print('⚠ TsComboBox não encontrado, continuando...')
        
        # Pressiona Ctrl+I
        print('Pressionando Ctrl+I...')
        send_keys('^i')
        print('✓ Ctrl+I pressionado!')
        time.sleep(1)
        
        print('✓ Área de busca aberta!')
        
    except Exception as e:
        print(f'✗ Erro ao abrir área de busca: {e}')
        raise

def fill_planilha_modelo():
    """Preenche a planilha modelo com dados do JSON"""
    try:
        print('\nPreenchendo planilha modelo...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Abre a planilha existente para ler o cabeçalho
        workbook_read = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        sheet_read = workbook_read.sheet_by_index(0)
        
        # Cria nova planilha
        workbook = Workbook()
        sheet = workbook.add_sheet('Plan1')
        
        # Copia o cabeçalho (primeira linha)
        print('Copiando cabeçalho...')
        for col_idx in range(sheet_read.ncols):
            value = sheet_read.cell_value(0, col_idx)
            sheet.write(0, col_idx, value)
        
        print('✓ Cabeçalho copiado')

        # Fecha o workbook de leitura para liberar o arquivo
        del workbook_read
        del sheet_read
        
        # Preenche os dados a partir da linha 1 (linha 0 é o cabeçalho)
        print('Preenchendo dados...')
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):
            integracao = registro.get('pro_numero_de_integracao', '')
            processo = registro.get('pro_numero_do_processo', '')
            valor = registro.get('debito_na_moeda', '')
            
            # Coluna A: pro_numero_de_integracao
            sheet.write(idx, 0, integracao)
            
            # Coluna B: pro_numero_do_processo
            sheet.write(idx, 1, processo)
            
            # Coluna C: texto fixo
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            
            # Coluna D: valor
            sheet.write(idx, 3, valor)
            
            # Coluna E: data de hoje
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] {integracao} - {processo} - {valor}')
        
        # Adiciona linha final com valor somado na coluna D
        # ultima_linha = len(data['registros']) + 1
        # valor_somado = data.get('valor_somado', '')
        # sheet.write(ultima_linha, 3, valor_somado)
        # print(f'\n✓ Valor total na linha {ultima_linha + 1}: {valor_somado}')
        
        # Salva a planilha
        workbook.save(PLANILHA_MODELO_PATH)
        print(f'✓ Planilha salva em: {PLANILHA_MODELO_PATH}')
        print(f'✓ Total de registros preenchidos: {len(data["registros"])}')
        
        # Libera explicitamente os objetos e força coleta de lixo
        del workbook
        del sheet
        import gc
        gc.collect()
        
        # Aguarda um pouco para garantir que o arquivo foi liberado pelo sistema operacional
        time.sleep(1)
        print('✓ Arquivo fechado e liberado')

    except Exception as e:
        print(f'✗ Erro ao preencher planilha modelo: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_planilha_modelo_v2():
    """Preenche a planilha modelo com dados do JSON - Versão 2
    
    Copia a planilha modelo do Excel para a pasta de destino antes de preencher
    """
    try:
        print('\nPreenchendo planilha modelo (v2)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        planilha_excel_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), 
            'planilha_modelo_excel', 
            'planilha_modelo_custas.xls'
        )
        
        # Verifica se a planilha modelo existe
        if not os.path.exists(planilha_excel_path):
            raise FileNotFoundError(f'Planilha modelo não encontrada: {planilha_excel_path}')
        
        print(f'✓ Planilha modelo encontrada: {planilha_excel_path}')
        
        # Cria a pasta de destino se não existir
        destino_folder = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), 
            'planilha_modelo', 
            'bmg'
        )
        os.makedirs(destino_folder, exist_ok=True)
        
        # Copia a planilha modelo para o destino
        print(f'Copiando planilha para: {PLANILHA_MODELO_PATH}')
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada com sucesso!')
        
        # Aguarda um pouco para garantir que o arquivo foi copiado
        time.sleep(0.5)
        
        # Abre a planilha copiada para ler o cabeçalho
        workbook_read = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        sheet_read = workbook_read.sheet_by_index(0)
        
        # Cria nova planilha
        workbook = Workbook()
        sheet = workbook.add_sheet('Plan1')
        
        # # Copia o cabeçalho (primeira linha)
        # print('Copiando cabeçalho...')
        # for col_idx in range(sheet_read.ncols):
        #     value = sheet_read.cell_value(0, col_idx)
        #     sheet.write(0, col_idx, value)
        
        # print('✓ Cabeçalho copiado')

        # Fecha o workbook de leitura para liberar o arquivo
        del workbook_read
        del sheet_read
        
        # Preenche os dados a partir da linha 1 (linha 0 é o cabeçalho)
        print('Preenchendo dados...')
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):
            integracao = registro.get('pro_numero_de_integracao', '')
            processo = registro.get('pro_numero_do_processo', '')
            valor = registro.get('debito_na_moeda', '')
            
            # Coluna A: pro_numero_de_integracao
            sheet.write(idx, 0, integracao)
            
            # Coluna B: pro_numero_do_processo
            sheet.write(idx, 1, processo)
            
            # Coluna C: texto fixo
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            
            # Coluna D: valor
            sheet.write(idx, 3, valor)
            
            # Coluna E: data de hoje
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] {integracao} - {processo} - {valor}')
        #pdb.set_trace()
        # Salva a planilha
        #workbook.save(PLANILHA_MODELO_PATH)
        print(f'✓ Planilha salva em: {PLANILHA_MODELO_PATH}')
        print(f'✓ Total de registros preenchidos: {len(data["registros"])}')
        
        # Libera explicitamente os objetos e força coleta de lixo
        del workbook
        del sheet
        import gc
        gc.collect()
        
        # Aguarda um pouco para garantir que o arquivo foi liberado pelo sistema operacional
        time.sleep(1)
        print('✓ Arquivo fechado e liberado')

    except Exception as e:
        print(f'✗ Erro ao preencher planilha modelo (v2): {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_planilha_modelo_v4():
    """Preenche a planilha modelo preservando formatação - Versão 4"""
    try: 
        print('\nPreenchendo planilha modelo (v4 - xlutils)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        planilha_excel_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), 
            'planilha_modelo_excel', 
            'planilha_modelo_custas_modelo.xls'
        )
        
        # Copia a planilha para o destino
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada!')
        
        # Abre a planilha original
        rb = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        
        # Cria uma cópia que preserva formatação
        wb = xl_copy(rb)
        sheet = wb.get_sheet(0)
        
        # Preenche os dados
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):  # Linha 1 (pula cabeçalho na linha 0)
            sheet.write(idx, 0, registro.get('pro_numero_de_integracao', ''))
            sheet.write(idx, 1, registro.get('pro_numero_do_processo', ''))
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            sheet.write(idx, 3, registro.get('debito_na_moeda', ''))
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] Linha preenchida')
        
        # Salva (preserva GRANDE PARTE da formatação)
        wb.save(PLANILHA_MODELO_PATH)
        
        print('✓ Planilha salva com formatação preservada!')
        
        del wb
        del rb
        import gc
        gc.collect()
        time.sleep(1)
        
    except Exception as e:
        print(f'✗ Erro: {e}')
        traceback.print_exc()
        raise

def fill_planilha_modelo_v5():
    """Preenche a planilha modelo preservando formatação - Versão 4"""
    try: 
        print('\nPreenchendo planilha modelo (v4 - xlutils)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        planilha_excel_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), 
            'planilha_modelo_excel', 
            'planilha_modelo_custas_modelo.xls'
        )
        
        # Copia a planilha para o destino
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada!')
        
        # IMPORTANTE: Usar context manager implícito ao não guardar referência
        # Abre a planilha original
        rb = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        
        # Cria uma cópia que preserva formatação
        wb = xl_copy(rb)
        sheet = wb.get_sheet(0)
        
        # Preenche os dados
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):  # Linha 1 (pula cabeçalho na linha 0)
            sheet.write(idx, 0, registro.get('pro_numero_de_integracao', ''))
            sheet.write(idx, 1, registro.get('pro_numero_do_processo', ''))
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            sheet.write(idx, 3, registro.get('debito_na_moeda', ''))
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] Linha preenchida')
        
        # Salva (preserva GRANDE PARTE da formatação)
        wb.save(PLANILHA_MODELO_PATH)
        
        print('✓ Planilha salva com formatação preservada!')
        
        # FORÇA FECHAMENTO EXPLÍCITO
        # Fecha todos os handles
        try:
            if hasattr(wb, 'close'):
                wb.close()
        except:
            pass
        
        try:
            if hasattr(rb, 'release_resources'):
                rb.release_resources()
        except:
            pass
        
        # Remove referências
        del sheet
        del wb
        del rb
        
        # Força coleta de lixo agressiva
        import gc
        gc.collect()
        gc.collect()  # Duas vezes para garantir
        
        # Aguarda mais tempo para o SO liberar o arquivo
        print('Aguardando liberação do arquivo...')
        time.sleep(2)  # Aumentei de 1 para 2 segundos
        
        # VALIDAÇÃO: Tenta abrir e fechar o arquivo para garantir que está liberado
        try:
            with open(PLANILHA_MODELO_PATH, 'rb') as test_file:
                test_file.read(1)
            print('✓ Arquivo validado e totalmente liberado!')
        except Exception as e:
            print(f'⚠ Aviso: Arquivo ainda pode estar travado: {e}')
            time.sleep(2)  # Aguarda mais um pouco
        
    except Exception as e:
        print(f'✗ Erro: {e}')
        traceback.print_exc()
        raise

def process_record_for_pdf(record: dict, index: int, primeira_procura: bool = False):
    """Processa um registro para gerar PDF"""
    try:
        integracao = record['pro_numero_de_integracao']

        #or 'CIV1428248' not in integracao
        # if 'CIV1482215' not in integracao:
        #     print(f'\n[{index + 1}] Ignorando registro.')
        #     return

        print(f'\n[{index + 1}] Processando integração: {integracao}')
        
        # Procura e clica no campo de busca TsEdit
        print('Procurando campos TsEdit...')
        desktop = Desktop(backend='win32')

        # Procurar TsComboBox habilitados
        if index == 0:
            selects_found = []
            for window in desktop.windows():
                try:
                    descendants = window.descendants()
                    for desc in descendants:
                        try:
                            if desc.class_name() == 'TsComboBox' and desc.is_enabled() == True and desc.is_visible() == True:
                                selects_found.append(desc)
                        except:
                            pass
                except:
                    pass
            
            print(f'✓ Total de TsComboBox encontrados: {len(selects_found)}')
            selects_found[0].click_input()
            time.sleep(0.5)
            print('✓ TsComboBox clicado!')
            send_keys('{DOWN}')
            time.sleep(0.5)
            send_keys('{DOWN}')
            time.sleep(0.5)
            send_keys('{ENTER}')
            time.sleep(0.5)
        
        edits_found = []
        for window in desktop.windows():
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsEdit':
                            edits_found.append(desc)
                    except:
                        pass
            except:
                pass
        
        print(f'✓ Total de TsEdit encontrados: {len(edits_found)}')

        edits_visible = [e for e in edits_found if e.is_visible()]
        print(f'✓ TsEdit visíveis: {len(edits_visible)}')

        for i, edit in enumerate(edits_visible):
            rect = edit.rectangle()
            print(f'  [{i}] Pos: ({rect.left}, {rect.top}), Tamanho: {rect.width()}x{rect.height()}')
         
        if len(edits_visible) >= 2:
            # Usa o segundo TsEdit visível (índice 1)
            edit_found = edits_visible[1]
            print(f'Usando o TsEdit visível na posição 1...')
            edit_found.click_input()
            print('✓ TsEdit clicado!')
            time.sleep(0.5)
            
            # Garante foco no elemento
            edit_found.set_focus()
            time.sleep(0.3)
            
            # Limpa o campo antes de digitar
            print('Limpando campo...')
            send_keys('^a')  # Ctrl+A para selecionar tudo
            send_keys('{DELETE}')  # Delete para limpar
            time.sleep(0.3)
            
            # Digita integração usando send_keys
            print(f'Digitando integração: {integracao}')
            send_keys(integracao)
            print('✓ Integração digitada!')
            time.sleep(0.5)
        
        # Enter para buscar
        print('Pressionando Enter para buscar...')
        send_keys('{ENTER}')
        print('✓ Busca realizada')
        time.sleep(1)
        
        # Aguarda até que TsDBCtrlGrid tenha elementos
        print('Aguardando resultados da busca...')
        max_wait = 30  # máximo 30 segundos
        elapsed = 0
        found_results = False
        grid_element = None
        
        while elapsed < max_wait:
            try:
                desktop = Desktop(backend='win32')
                for window in desktop.windows():
                    try:
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TsDBCtrlGrid':
                                    # Verifica se o grid tem filhos (resultados)
                                    children = desc.children()
                                    if len(children) == 1:
                                        print(f'✓ Resultados encontrados! Total de elementos no grid: {len(children)}')
                                        found_results = True
                                        grid_element = desc
                                        break
                            except:
                                pass
                        if found_results:
                            break
                    except:
                        pass
                if found_results:
                    break
            except:
                pass
            
            remaining = max_wait - elapsed
            print(f'\rAguardando resultados... {remaining}s restantes', end='', flush=True)
            time.sleep(1)
            elapsed += 1

        if not found_results:
            print(f'\r⚠ Timeout: Nenhum resultado encontrado após {max_wait}s')
        else:
            print('\r✓ Resultados carregados com sucesso!              ')
            time.sleep(6)
            
            # Pressiona Enter para abrir o resultado encontrado
            print('Pressionando Enter para abrir o resultado...')
            if primeira_procura == False:
                send_keys('{DOWN}')

            time.sleep(3)
            send_keys('{ENTER}')
            print('✓ Enter pressionado!')
            
            # Aguarda até que TsPanel fique visível
            print('Aguardando TsPanel ficar visível...')
            max_wait_panel = 30
            elapsed_panel = 0
            panel_visible = False
            panel_visible_lembrete = False
            
            while elapsed_panel < max_wait_panel:
                try:
                    desktop = Desktop(backend='win32')
                    
                    # Verifica se existe janela "Cadastro de Processos - Lembretes"
                    for window in desktop.windows():
                        try:
                            #print(f'Verificando janela: {window.window_text()}')
                            if 'Cadastro de Processos - Lembretes' in window.window_text() and window.is_visible():
                                print('✓ Janela "Cadastro de Processos - Lembretes" encontrada, pressionando F4...')
                                send_keys('{F4}')
                                time.sleep(1)
                                break
                        except:
                            pass
                    
                    for window in desktop.windows():
                        try:
                            descendants = window.descendants()
                            for desc in descendants:
                                try:
                                    #print(f'Verificando elemento: {desc.class_name()} - Visível: {desc.is_visible()}')
                                    if desc.class_name() == 'TsPanel' and desc.is_visible():
                                        print(f'✓ TsPanel visível encontrado!')
                                        panel_visible = True
                                        break
                                except:
                                    pass
                            if panel_visible:
                                break
                        except:
                            pass
                    if panel_visible:
                        break
                except:
                    pass
                
                remaining_panel = max_wait_panel - elapsed_panel
                print(f'\rAguardando TsPanel... {remaining_panel}s restantes', end='', flush=True)
                time.sleep(1)
                elapsed_panel += 1

            if panel_visible:
                print('\r✓ TsPanel visível! Executando Ctrl+L e digitando SPF...              ')
                send_keys('^l')
                time.sleep(0.5)
                send_keys('SPF')
                print('✓ Ctrl+L pressionado e SPF digitado no TsPanel!')
                time.sleep(1)
                send_keys('{ENTER}')
                print('✓ Enter pressionado após digitar SPF')
                time.sleep(3)  # Aguarda mais tempo para a janela SPF carregar
                send_keys('{ENTER}')
                print('✓ Enter pressionado após digitar SPF')
                time.sleep(3) 
                
                # Procura e clica no elemento Grid (TPage) da janela SPF
                print('Procurando elemento TPage em todas as janelas...')
                desktop = Desktop(backend='win32')
                
                # Lista todos os TPage encontrados
                tpages_found = []
                for window in desktop.windows():
                    try:
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TPage' and desc.is_visible():
                                    window_title = window.window_text()
                                    desc_name = desc.window_text()
                                    rect = desc.rectangle()
                                    tpages_found.append({
                                        'element': desc,
                                        'name': desc_name,
                                        'window': window_title,
                                        'rect': rect
                                    })
                                    print(f'  TPage visível: "{desc_name}" na janela "{window_title}" em {rect}')
                            except:
                                pass
                    except:
                        pass
                
                print(f'✓ Total de TPage visíveis: {len(tpages_found)}')
                
                # Garante que a janela SPF está em foco
                print('\nColocando janela SPF em foco...')
                spf_window = None
                for window in desktop.windows():
                    try:
                        window_title = window.window_text()
                        if 'SPF' in window_title and 'Solicitação' in window_title:
                            spf_window = window
                            print(f'✓ Janela SPF encontrada: "{window_title}"')
                            break
                    except:
                        pass
                
                if spf_window is None:
                    print('⚠ Janela SPF não encontrada!')
                    send_keys('%{F4}')  # Alt+F4 para fechar
                    time.sleep(2)
                    send_keys('{RIGHT}')  # Right para nao fechar
                    time.sleep(1)
                    send_keys('{ENTER}')  # Enter para confirmar
                    return process_record_for_pdf(record, index, primeira_procura=True)

                if spf_window:
                    try:
                        # Traz a janela para frente e dá foco
                        spf_window.set_focus()
                        print('✓ Janela SPF em foco')
                        time.sleep(1)
                    except Exception as e:
                        print(f'⚠ Erro ao dar foco: {e}')
                    
                    # Procura TsComboBox "Ambas" diretamente
                    print('\nProcurando TsComboBox "Ambas"...')
                    descendants = spf_window.descendants()
                    elemento_encontrado = None
                    
                    for desc in descendants:
                        try:
                            if desc.is_visible() and desc.class_name() == 'TsComboBox':
                                desc_name = desc.window_text()
                                if desc_name == 'Ambas':
                                    rect = desc.rectangle()
                                    # Posiciona no início do elemento (canto superior esquerdo)
                                    start_x = rect.left
                                    start_y = rect.top + 55 
                                    
                                    print(f'✓ TsComboBox "Ambas" encontrado em ({start_x}, {start_y})')
                                    print('Posicionando mouse no início do elemento...')
                                    pyautogui.moveTo(start_x, start_y)
                                    time.sleep(2)
                                    pyautogui.click(start_x, start_y)
                                    # Clica no elemento
                                    print('Clicando no elemento...')
                                    print('✓ Clicado com pyautogui!')
                                    time.sleep(1)
                                    break
                        except:
                            pass

        print('Pressionando seta para baixo...')
        send_keys('{DOWN}')
        time.sleep(0.5)
        send_keys('{DOWN}')
        time.sleep(0.5)
        
        print('Pressionando Enter para confirmar...')                     
        send_keys('{ENTER}')
        time.sleep(2)
        
        print('Digitando Downloads...')
        send_keys('Downloads')
        time.sleep(0.5)
        
        print('Pressionando seta para direita...')
        send_keys('{RIGHT}')
        time.sleep(0.5)
        
        print('Digitando pdf_cpj...')
        send_keys('pdf_cpj')
        time.sleep(0.5)
        
        print('Pressionando Enter...')
        send_keys('{ENTER}')

        # Obtém o nome do usuário atual do Windows

        print(f'\n✓ PDFs sendo salvos em: {PDF_FOLDER}')
        
        # Aguarda e clica no botão OK da janela de confirmação
        print('\nProcurando botão OK...')
        desktop = Desktop(backend='win32')
        botao_ok = False

        while botao_ok == False:
            time.sleep(5)
            for window in desktop.windows():
                try:
                    window_title = window.window_text()
                    if 'Controle de Processos Jurídicos' in window_title and window.is_visible():
                        #print(f'✓ Janela encontrada: "{window_title}"')
                        
                        # Procura o botão OK
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TButton' and desc.window_text() == 'OK':
                                    print('Clicando no botão OK...')
                                    desc.click_input()
                                    print('✓ Botão OK clicado!')
                                    time.sleep(1)
                                    botao_ok = True
                                    break
                            except:
                                botao_ok = False
                                pass
                        break
                except:
                    botao_ok = False
                    pass
        
        # Lista e faz merge dos PDFs
        print(f'\nListando PDFs em: {PDF_FOLDER}')
        
        if os.path.exists(PDF_FOLDER):
            # Busca recursivamente por PDFs em todas as subpastas
            pdf_files = []
            for root, dirs, files in os.walk(PDF_FOLDER):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_path = os.path.join(root, file)
                        pdf_files.append(pdf_path)
            
            # Ordena os PDFs por caminho
            pdf_files.sort()
            
            print(f'✓ Total de PDFs encontrados: {len(pdf_files)}')
            
            for pdf in pdf_files:
                print(f'  - {pdf}')
            
            if pdf_files:
                # Cria pasta pdf_merge se não existir
                merge_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdf_merge')
                os.makedirs(merge_folder, exist_ok=True)
                print(f'✓ Pasta de merge: {merge_folder}')
                
                # Constrói nome do arquivo com trim
                integracao = record['pro_numero_de_integracao'].strip()
                processo = record['pro_numero_do_processo'].strip()
                debito = record['debito_na_moeda'].strip()
                
                merged_filename = f'{integracao}_{processo}_{debito}.pdf'
                merged_path = os.path.join(merge_folder, merged_filename)
                
                print(f'\nFazendo merge dos PDFs...')
                print(f'Nome do arquivo: {merged_filename}')
                
                # Faz o merge
                merger = PdfMerger()
                
                for pdf_path in pdf_files:
                    merger.append(pdf_path)
                    print(f'  + {os.path.basename(pdf_path)}')
                
                # Salva o PDF mesclado
                merger.write(merged_path)
                merger.close()
                
                print(f'✓ PDF mesclado salvo em: {merged_path}')
                
                # Deleta todo o conteúdo de PDF_FOLDER
                print(f'\nLimpando pasta {PDF_FOLDER}...')
                
                for item in os.listdir(PDF_FOLDER):
                    item_path = os.path.join(PDF_FOLDER, item)
                    try:
                        if os.path.isfile(item_path):
                            os.unlink(item_path)
                            print(f'  ✓ Arquivo deletado: {item}')
                        elif os.path.isdir(item_path):
                            shutil.rmtree(item_path)
                            print(f'  ✓ Pasta deletada: {item}')
                    except Exception as e:
                        print(f'  ✗ Erro ao deletar {item}: {e}')
                
                print(f'✓ Pasta limpa com sucesso!')
            else:
                print('⚠ Nenhum PDF encontrado para fazer merge')
        else:
            print(f'✗ Pasta não encontrada: {PDF_FOLDER}')
        
        # Pressiona F4 duas vezes para finalizar
        print('\nFinalizando processo...')
        print('Pressionando F4...')
        send_keys('{F4}')
        time.sleep(1)
        print('Pressionando F4 novamente...')
        send_keys('{F4}')
        time.sleep(1)
        print('✓ Processo finalizado!')

        print(f'\n✓ Registro {index + 1} processado com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao processar registro {index + 1}: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_descritivo_pdf():

    #global NUMERO_RECIBO

    doc = Document(DOCX_MODELO)

    with open(JSON_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Dados para substituição
    #NUMERO_RECIBO = numero_recibo = input('Digite o número do recibo: ').strip()
    data_hoje = datetime.now().strftime('%d/%m/%Y')
    mes_hoje = datetime.now().strftime('%m/%Y')
    valor_somado = data['valor_somado']

    # Se o documento tem uma tabela para preencher
    if len(doc.tables) > 0:

        for table_idx, table in enumerate(doc.tables):
            print(f'\nProcessando tabela {table_idx + 1}...')
            
            # Percorre todas as células da tabela
            for row in table.rows:
                for cell in row.cells:
                    # Verifica se a célula contém os marcadores
                    if r'{{NUMERO_RECIBO}}' in cell.text:
                        #numero_recibo = input('Digite o número do recibo: ').strip()
                        numero_recibo = NUMERO_RECIBO
                        cell.text = cell.text.replace('{{NUMERO_RECIBO}}', numero_recibo)
                        print(f'  ✓ {{{{NUMERO_RECIBO}}}} → {numero_recibo}')
                    
                    if r'{{DATA_HOJE}}' in cell.text:
                        cell.text = cell.text.replace('{{DATA_HOJE}}', data_hoje)
                        print(f'  ✓ {{{{DATA_HOJE}}}} → {data_hoje}')
                    
                    if r'{{MES_HOJE}}' in cell.text:
                        cell.text = cell.text.replace('{{MES_HOJE}}', mes_hoje)
                        print(f'  ✓ {{{{MES_HOJE}}}} → {mes_hoje}')
                    
                    if r'{{VALOR_SOMADO}}' in cell.text:
                        cell.text = cell.text.replace('{{VALOR_SOMADO}}', valor_somado)
                        print(f'  ✓ {{{{VALOR_SOMADO}}}} → {valor_somado}')

        table = doc.tables[2]  # Primeira tabela do documento

        for idx, registro in enumerate(data['registros']):
            row_idx = idx + 1

            # Adiciona nova linha se necessário
            if row_idx >= len(table.rows):
                table.add_row()
            
            row = table.rows[row_idx]
            row.cells[0].text = 'Custas'
            row.cells[1].text = datetime.now().strftime('%d/%m/%Y')
            row.cells[2].text = registro['pro_numero_de_integracao']
            row.cells[3].text = 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU'
            row.cells[4].text = registro['debito_na_moeda'].strip()
    

    # Se o documento tem uma tabela para preencher
    if len(doc.tables) > 0:
        table = doc.tables[2]  # Primeira tabela do documento
        
        # Preenche os registros (começa na linha 1, pois linha 0 é cabeçalho)
        for idx, registro in enumerate(data['registros']):
            row_idx = idx + 1
            
            # Adiciona nova linha se necessário
            if row_idx >= len(table.rows):
                table.add_row()
            
            row = table.rows[row_idx]
            row.cells[0].text = 'Custas'
            row.cells[1].text = datetime.now().strftime('%d/%m/%Y')
            row.cells[2].text = registro['pro_numero_de_integracao']
            row.cells[3].text = 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU'
            row.cells[4].text = registro['debito_na_moeda'].strip()

    doc.save(DOCX_OUTPUT)
    print(f'✓ Documento Word salvo em: {DOCX_OUTPUT}')
    
    # Converte para PDF
    print('\nConvertendo para PDF...')
    convert(DOCX_OUTPUT, PDF_OUTPUT)
    print(f'✓ PDF salvo em: {PDF_OUTPUT}')
    
    # Faz merge com o primeiro PDF da pasta pdf_merge
    # print('\nFazendo merge com PDFs da pasta pdf_merge...')
    # pdf_merge_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdf_merge')
    
    # if os.path.exists(pdf_merge_folder):
    #     # Lista todos os PDFs da pasta pdf_merge
    #     pdf_files = sorted([f for f in os.listdir(pdf_merge_folder) if f.lower().endswith('.pdf')])
        
    #     if pdf_files:
    #         primeiro_pdf = os.path.join(pdf_merge_folder, pdf_files[0])
    #         print(f'✓ Primeiro PDF encontrado: {pdf_files[0]}')
            
    #         # Cria nome do arquivo final mantendo o nome do primeiro PDF + "_COM_DESCRITIVO"
    #         nome_base = os.path.splitext(pdf_files[0])[0]  # Remove a extensão .pdf
    #         nome_final = f'{nome_base}_COM_DESCRITIVO.pdf'
    #         merged_final_path = os.path.join(pdf_merge_folder, nome_final)
            
    #         merger = PdfMerger()
            
    #         # Adiciona o descritivo primeiro
    #         merger.append(PDF_OUTPUT)
    #         print(f'  + {os.path.basename(PDF_OUTPUT)}')
            
    #         # Adiciona o primeiro PDF da pasta
    #         merger.append(primeiro_pdf)
    #         print(f'  + {os.path.basename(primeiro_pdf)}')
            
    #         # Salva o PDF final
    #         merger.write(merged_final_path)
    #         merger.close()
            
    #         print(f'✓ PDF final com merge salvo em: {merged_final_path}')
            
    #         # Apaga o arquivo original que foi usado no merge
    #         os.unlink(primeiro_pdf)
    #         print(f'✓ Arquivo original removido: {pdf_files[0]}')
    #     else:
    #         print('⚠ Nenhum PDF encontrado na pasta pdf_merge')
    # else:
    #     print(f'⚠ Pasta não encontrada: {pdf_merge_folder}')

def open_chrome_browser():
    """Abre o navegador Chrome usando Selenium"""
    try:
        print('\nAbrindo navegador Chrome...')
        
        # Configura opções do Chrome
        chrome_options = Options()
        # chrome_options.add_argument('--headless')  # Descomente para rodar sem interface gráfica
        chrome_options.add_argument('--start-maximized')
        chrome_options.add_argument('--disable-blink-features=AutomationControlled')
        chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])
        
        # Método 1: Tentar sem especificar o service (usa ChromeDriver do PATH)
        try:
            print('Tentando usar ChromeDriver do PATH do sistema...')
            driver = webdriver.Chrome(options=chrome_options)
            print('✓ Chrome aberto com sucesso usando ChromeDriver do PATH!')
            return driver
        except Exception as path_error:
            print(f'⚠ ChromeDriver não encontrado no PATH: {path_error}')
        
        # Método 2: Tentar com webdriver-manager (download automático)
        try:
            print('Tentando instalar ChromeDriver automaticamente com webdriver-manager...')
            
            # Limpa o cache primeiro para forçar novo download
            cache_path = os.path.join(os.path.expanduser('~'), '.wdm')
            if os.path.exists(cache_path):
                try:
                    shutil.rmtree(cache_path)
                    print('✓ Cache limpo, baixando ChromeDriver novamente...')
                except Exception as cache_error:
                    print(f'⚠ Erro ao limpar cache: {cache_error}')
            
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)
            print('✓ Chrome aberto com sucesso usando webdriver-manager!')
            return driver
            
        except Exception as manager_error:
            print(f'✗ Erro com webdriver-manager: {manager_error}')
        
        # Se chegou aqui, nenhum método funcionou
        raise Exception('Não foi possível inicializar o ChromeDriver')
        
    except Exception as e:
        print(f'\n✗ ERRO CRÍTICO: Não foi possível abrir o Chrome')
        print(f'Erro: {e}')
        print('\n' + '='*70)
        print('SOLUÇÃO MANUAL - Siga estes passos:')
        print('='*70)
        print('\n1. Verifique a versão do seu Chrome:')
        print('   - Abra o Chrome')
        print('   - Vá em: chrome://settings/help')
        print('   - Anote a versão (ex: 131.0.6778.86)')
        print('\n2. Baixe o ChromeDriver correspondente:')
        print('   - Acesse: https://googlechromelabs.github.io/chrome-for-testing/')
        print('   - Baixe a versão EXATA do seu Chrome para Windows 64-bit')
        print('   - Escolha o arquivo chromedriver-win64.zip')
        print('\n3. Extraia e configure:')
        print('   - Extraia o arquivo chromedriver.exe')
        print('   - Coloque em: C:\\chromedriver\\chromedriver.exe')
        print('   - Ou adicione a pasta ao PATH do sistema')
        print('\n4. Reinicie este script')
        print('='*70)
        
        import traceback
        traceback.print_exc()
        raise

def login_web_exyon_bmg(driver):
    """Faz login no sistema web"""
    try:
        print('\nRealizando login no sistema...')

        # Navegar para a URL de login
        print('\nAcessando sistema jurídico...')
        driver.get('https://juridico.bancobmg.com.br/Paginas/Principal/_FSet_Abertura.asp?Pagina=Logon')
        print('✓ Página carregada')
        
        # Aguarda a página carregar
        wait = WebDriverWait(driver, 10)
        
        # Encontra e preenche o campo de login
        print('Preenchendo campo de login...')
        login_input = wait.until(EC.presence_of_element_located((By.ID, 'txtcd_Logon')))
        login_input.clear()
        login_input.send_keys('rafael.dias')
        print('✓ Login digitado: rafael.dias')
        
        # Encontra e preenche o campo de senha
        print('Preenchendo campo de senha...')
        senha_input = driver.find_element(By.ID, 'txtcd_Pwd')
        senha_input.clear()
        senha_input.send_keys('yIt21qZVKanwNuKf')
        print('✓ Senha digitada')
        
        # Clica no botão OK
        print('Clicando no botão OK...')
        btn_ok = driver.find_element(By.ID, 'btOK')
        btn_ok.click()
        print('✓ Botão OK clicado!')
        
        # Aguarda um pouco para o login processar
        time.sleep(2)
        print('✓ Login realizado com sucesso!')
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao fazer login: {e}')
        import traceback
        traceback.print_exc()
        return False

def selecionar_unidade(driver):
    """Seleciona a unidade após o login"""
    try:
        print('\nAguardando seleção de unidade...')
        
        # Aguarda o elemento cmbUnidade aparecer
        wait = WebDriverWait(driver, 10)
        print('Aguardando elemento cmbUnidade aparecer...')
        select_element = wait.until(EC.presence_of_element_located((By.ID, 'cmbUnidade')))
        print('✓ Elemento cmbUnidade encontrado')
        
        # Cria objeto Select para manipular o dropdown
        select = Select(select_element)
        
        # Seleciona a opção com value 106
        print('Selecionando unidade (value=106)...')
        select.select_by_value('106')
        print('✓ Unidade 106 selecionada')
        
        # Clica no botão Entrar (btOk)
        print('Clicando no botão Entrar...')
        btn_entrar = driver.find_element(By.ID, 'btOk')
        btn_entrar.click()
        print('✓ Botão Entrar clicado!')
        
        # Aguarda um pouco para processar
        time.sleep(2)
        print('✓ Unidade selecionada com sucesso!')
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao selecionar unidade: {e}')
        import traceback
        traceback.print_exc()
        return False

def executar_script_menu(driver):
    """Executa o script JavaScript para mudar o menu"""
    try:
        print('\nExecutando script de menu...')
        
        # Aguarda um pouco para garantir que a página carregou
        time.sleep(2)
        
        # Muda para o frame FraMenu
        print('Mudando para o frame FraMenu...')
        wait = WebDriverWait(driver, 10)
        frame = wait.until(EC.presence_of_element_located((By.ID, 'FraMenu')))
        driver.switch_to.frame(frame)
        print('✓ Frame FraMenu encontrado e selecionado')
        
        # Executa o script JavaScript
        script = "MudaClasseMenu('Menu_1', '6');Carrega_Combo('FI');"
        print(f'Executando: {script}')
        driver.execute_script(script)
        print('✓ Script executado com sucesso!')
        
        # Volta para o contexto principal
        driver.switch_to.default_content()
        print('✓ Voltou ao contexto principal')
        
        # Aguarda o script processar
        time.sleep(2)
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao executar script: {e}')
        import traceback
        traceback.print_exc()
        return False

def executar_script_menu_lateral(driver):
    """Executa o script JavaScript do menu lateral"""
    try:
        print('\nExecutando script de menu lateral...')
        
        # Aguarda um pouco
        time.sleep(2)
        
        # Muda para o frame FraMenu
        print('Mudando para o frame FraMenu...')
        wait = WebDriverWait(driver, 10)
        frame = wait.until(EC.presence_of_element_located((By.ID, 'FraMenu')))
        driver.switch_to.frame(frame)
        print('✓ Frame FraMenu encontrado e selecionado')
        
        # Executa o script JavaScript
        script = "MudaClasseMenuLateral('LNF');Mnu_Click_Handler('LNF');parent.document.getElementById('__submenu').style.display='none';parent.document.getElementById('__submenu').innerHTML = '';document.getElementById('__setinha').src = 'imagens/icone-select.png';"
        print(f'Executando script do menu lateral...')
        driver.execute_script(script)
        print('✓ Script do menu lateral executado com sucesso!')
        
        # Volta para o contexto principal
        driver.switch_to.default_content()
        print('✓ Voltou ao contexto principal')
        
        # Aguarda o script processar
        time.sleep(2)
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao executar script do menu lateral: {e}')
        import traceback
        traceback.print_exc()
        return False

def executar_script_classificacao(driver):
    """Executa o script de classificação financeira e seleciona CUSTAS"""
    try:
        print('\nExecutando script de classificação financeira...')
        
        # Aguarda um pouco
        time.sleep(2)
        
        # Muda para o frame FraDetalhe
        print('Mudando para o frame FraDetalhe...')
        wait = WebDriverWait(driver, 10)
        frame = wait.until(EC.presence_of_element_located((By.ID, 'FraDetalhe')))
        driver.switch_to.frame(frame)
        print('✓ Frame FraDetalhe encontrado e selecionado')
        
        # Guarda a janela principal
        janela_principal = driver.current_window_handle
        
        # Executa o script que abre o popup
        script = "searcher('ClassificacaoFinanceira_List_pop.asp');"
        print(f'Executando: {script}')
        driver.execute_script(script)
        print('✓ Script executado - popup deve ter aberto')
        
        # Aguarda o popup abrir
        time.sleep(2)
        
        # Muda para a nova janela popup
        print('Procurando janela popup...')
        todas_janelas = driver.window_handles
        for janela in todas_janelas:
            if janela != janela_principal:
                driver.switch_to.window(janela)
                print('✓ Popup encontrado e selecionado')
                break
        
        # Aguarda o frame do popup carregar e muda para ele
        print('Aguardando frame do popup...')
        time.sleep(1)
        frame_popup = wait.until(EC.presence_of_element_located((By.XPATH, '/html/frameset/frame')))
        driver.switch_to.frame(frame_popup)
        print('✓ Frame do popup encontrado e selecionado')
        
        # Executa o script select no popup
        script_select = "select('34','CUSTAS');"
        print(f'Executando no popup: {script_select}')
        driver.execute_script(script_select)
        print('✓ Script select executado no popup!')
        
        # Aguarda processar
        time.sleep(2)
        
        # Volta para a janela principal
        driver.switch_to.window(janela_principal)
        print('✓ Voltou à janela principal')
        
        # Volta para o contexto padrão
        driver.switch_to.default_content()
        print('✓ Voltou ao contexto padrão')
        
        # Muda para o frame FraDetalhe para preencher a data
        print('Mudando para o frame FraDetalhe...')
        frame_detalhe = wait.until(EC.presence_of_element_located((By.ID, 'FraDetalhe')))
        driver.switch_to.frame(frame_detalhe)
        print('✓ Frame FraDetalhe selecionado')
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao executar script de classificação: {e}')
        import traceback
        traceback.print_exc()
        return False

def executar_preenchimento_formulario(driver, valor_somado):
    """Executa o script de preenchimento do formulário"""
    try:
        print('\nExecutando preenchimento do formulário..')
        
        # Aguarda um pouco
        time.sleep(2)
        
        # Preenche a data de emissão
        print('Preenchendo data de emissão...')
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        dt_emissao = driver.find_element(By.ID, 'dt_Emissao')
        dt_emissao.clear()
        dt_emissao.click()
        dt_emissao.send_keys(data_hoje)
        print(f'✓ Data de emissão preenchida: {data_hoje}')
        
        # Aguarda processar
        time.sleep(1)
        
        # Preenche o número da NF
        print('Solicitando número da NF...')
        #numero_nf = input('Digite o número da NF: ').strip()
        numero_nf = NUMERO_RECIBO
        
        print(f'Preenchendo número da NF: {numero_nf}...')
        nu_nf = driver.find_element(By.ID, 'nu_NF')
        nu_nf.clear()
        nu_nf.click()
        nu_nf.send_keys(numero_nf)
        print(f'✓ Número da NF preenchido: {numero_nf}')
        
        # Aguarda processar
        time.sleep(1)
        
        # Preenche o valor da nota
        print(f'Preenchendo valor da nota: {valor_somado}...')
        vl_nota = driver.find_element(By.ID, 'vl_Nota')
        vl_nota.clear()
        vl_nota.click()
        vl_nota.send_keys(valor_somado)
        print(f'✓ Valor da nota preenchido: {valor_somado}')
        
        # Aguarda processar
        # Faz upload do arquivo
        
        print(f'Fazendo upload do arquivo: {PLANILHA_MODELO_PATH}...')
        arquivo_input = driver.find_element(By.ID, 'ArqCarga')
        arquivo_input.send_keys(PLANILHA_MODELO_PATH)
        print(f'✓ Arquivo enviado: {PLANILHA_MODELO_PATH}')
        
        # Aguarda processar
        time.sleep(1)
        print('Abrindo arquivo no Excel...')
        os.startfile(PLANILHA_MODELO_PATH)

        # Aguarda o Excel abrir
        time.sleep(3)

        # Clica no botão Enviar
        print('Clicando no botão Enviar...')
        btn_enviar = driver.find_element(By.ID, 'btn_Enviar')
        btn_enviar.click()
        print('✓ Botão Enviar clicado!')
        
        # Aguarda alert aparecer
        time.sleep(1)
        
        # Aceita o alert que aparecer
        try:
            print('Aguardando alert...')
            alert = WebDriverWait(driver, 5).until(EC.alert_is_present())
            print('✓ Alert detectado, aceitando...')
            alert.accept()
            print('✓ Alert aceito!')
        except:
            print('⚠ Nenhum alert detectado')

        #Agora fecha o Excel programaticamente
        print('Fechando Excel...')
        os.system(f'taskkill /F /IM EXCEL.EXE')
                
        # Aguarda processar
        time.sleep(1)
        
        # Aguarda elemento aparecer
        print('Aguardando elemento aparecer...')

        fieldset_imgs = len(driver.find_elements(By.XPATH, '/html/body/fieldset/form/div[1]/fieldset/div/center/table/tbody/tr[3]/td/center/img'))
        itercaoes = 0
        while fieldset_imgs == 1:
            print('Elemento ainda não encontrado, aguardando mais 5 segundos...')
            time.sleep(5)
            fieldset_imgs = len(driver.find_elements(By.XPATH, '/html/body/fieldset/form/div[1]/fieldset/div/center/table/tbody/tr[3]/td/center/img'))
            itercaoes += 1
            if itercaoes >= 12:  # Limite de 1 minuto
                print('✗ Tempo esgotado aguardando o elemento aparecer')
                return False

        print('✓ Elemento encontrado!')
        
        # Aguarda processar
        time.sleep(2)
        
        return True
        
    except Exception as e:
        print(f'✗ Erro ao executar script de classificação: {e}')
        import traceback
        traceback.print_exc()
        return False

def anexar_pdfs_formulario(driver):
    """Anexa os PDFs da pasta pdf_merge nos inputs correspondentes do formulário"""
    try:
        print('\nAnexando PDFs ao formulário...')

        # Caminho da pasta pdf_merge
        pdf_merge_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'pdf_merge')
        
        if not os.path.exists(pdf_merge_folder):
            print(f'✗ Pasta não encontrada: {pdf_merge_folder}')
            return False
        
        # Lista todos os PDFs da pasta
        pdf_files = [f for f in os.listdir(pdf_merge_folder) if f.lower().endswith('.pdf')]
        print(f'✓ Total de PDFs encontrados: {len(pdf_files)}')
        
        if not pdf_files:
            print('⚠ Nenhum PDF encontrado na pasta pdf_merge')
            return False
        
        # Processa cada PDF
        for pdf_file in pdf_files:
            print(f'\nProcessando PDF: {pdf_file}')
            
            # Remove a extensão .pdf
            nome_sem_extensao = os.path.splitext(pdf_file)[0]
            
            # Faz split por underline
            partes = nome_sem_extensao.split('_')
            
            if len(partes) < 3:
                print(f'⚠ Nome do arquivo inválido: {pdf_file}')
                continue
            
            # Extrai numero_integracao (posição 0) e valor_pleito (posição 2)
            numero_integracao = partes[0]
            valor_pleito = partes[2].strip()
            
            print(f'  Número Integração: {numero_integracao}')
            print(f'  Valor Pleito: {valor_pleito}')
            
            # Procura o input correspondente no formulário
            encontrado = False
            indice = 1
            indice_simbolo = 2
            #/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[2]/td[1] 
            #/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[6]/td[1]
            #/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[10]/td[1]
            #/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[2]/td[1]
            
            while True:
                try:
                    #pula quando problema para importar
                    if indice_simbolo != 2:
                        indice_simbolo +=  4
                    
                    elemento = driver.find_element(By.XPATH, f'/html/body/fieldset/form/div[1]/fieldset/table[1]/tbody/tr[{indice_simbolo}]')
                    html_interno = elemento.get_attribute('innerHTML')

                    if '/images/check_green.gif' not in html_interno:
                        print('XXXXXXXXXXXX TROCAR NA PLANILHA DE IMPORTAÇÃO  XXXXXXXXXXXX')
                        continue

                    # Tenta encontrar o input cau_X
                    cau_input = driver.find_element(By.ID, f'cau_{indice}')
                    cau_value = cau_input.get_attribute('value')
                    
                    # Tenta encontrar o input Pleito_X
                    pleito_input = driver.find_element(By.ID, f'Pleito_{indice}')
                    pleito_value = pleito_input.get_attribute('value').strip().replace('.','')
                    
                    # Verifica se os valores correspondem
                    if cau_value == numero_integracao and pleito_value == valor_pleito:
                        print(f'  ✓ Correspondência encontrada no índice {indice}')
                        
                        pdf_path = os.path.join(pdf_merge_folder, pdf_file)
                        
                        # Se for o primeiro índice (cau_1), faz merge com descritivo
                        if indice == 1:
                            print('  → Primeira posição: fazendo merge com descritivo...')
                            
                            # Nome do arquivo com descritivo
                            nome_base = os.path.splitext(pdf_file)[0]
                            pdf_com_descritivo = f'{nome_base}_COM_DESCRITIVO.pdf'
                            pdf_final_path = os.path.join(pdf_merge_folder, pdf_com_descritivo)
                            
                            # Faz o merge
                            merger = PdfMerger()
                            
                            # Adiciona o descritivo primeiro
                            merger.append(PDF_OUTPUT)
                            print(f'    + {os.path.basename(PDF_OUTPUT)}')
                            
                            # Adiciona o PDF original
                            merger.append(pdf_path)
                            print(f'    + {pdf_file}')
                            
                            # Salva o PDF mesclado
                            merger.write(pdf_final_path)
                            merger.close()
                            
                            print(f'  ✓ PDF com descritivo criado: {pdf_com_descritivo}')
                            
                            # Usa o PDF mesclado para anexar
                            pdf_path = pdf_final_path
                        
                        # Anexa o PDF no input docLancfin_X
                        doc_input = driver.find_element(By.ID, f'docLancfin_{indice}')
                        doc_input.send_keys(pdf_path)
                        
                        print(f'  ✓ PDF anexado: {os.path.basename(pdf_path)}')
                        encontrado = True
                        time.sleep(0.5)
                        break
                    
                    indice += 1
                    
                except:
                    # Se não encontrar mais inputs, sai do loop
                    break
            
            if not encontrado:
                print(f'  ✗ Correspondência não encontrada para: {pdf_file}')
        
        print('\n✓ Todos os PDFs foram processados!')
        return True
        
    except Exception as e:
        print(f'✗ Erro ao anexar PDFs: {e}')
        import traceback
        traceback.print_exc()
        return False

def finalizar_processo(driver, valor_somado):
    print('\nVerificando valor total...')
    try:
        # Pega o valor do campo txtTotalLiquidoTela
        total_tela = driver.find_element(By.ID, 'txtTotalLiquidoTela').get_attribute('value')
        print(f'Valor na tela: {total_tela}')
        
        # Normaliza os valores removendo pontos
        valor_somado_normalizado = valor_somado
        total_tela_normalizado = total_tela.replace('.', '')
        
        print(f'Valor esperado: {valor_somado_normalizado}')
        print(f'Valor na tela (normalizado): {total_tela_normalizado}')
        
        # Compara os valores
        if valor_somado_normalizado == total_tela_normalizado:
            print('\n✓✓✓ Pronto para submeter ✓✓✓')
        else:
            print(f'\n⚠ ATENÇÃO: Valores não conferem!')
            print(f'   Esperado: {valor_somado_normalizado}')
            print(f'   Na tela: {total_tela_normalizado}')
    except Exception as e:
        print(f'⚠ Erro ao verificar valor total: {e}')

def main():
    """Executa o fluxo completo"""
    try:
        ## Etapa 1: Login e exportação do CSV

        etapa_cpj = True
        etapa_arquivos = False  

        if etapa_cpj and not etapa_arquivos:
            close_controle_processos_windows()
            clean_pdf_merge_folder()
            clean_planilha_folder()
            app = open_cpj_application()
            dlg = app.top_window()

            perform_login(app)   

            try:   
                click_menu(app)
            except Exception as e:
                close_cpj_application(app)
                print(f'Erro ao clicar no menu: {e}')
                dlg.type_keys('{ENTER}')
                time.sleep(5)
                app = open_cpj_application()
                perform_login(app) 
                
            click_printer(app)
            navigate_to_analise_lancamento(app)
            navigate_to_relatorio_conta(app)
            fill_cliente_and_dates(app)

            try:
                export_file(app)
                close_window(app)

                ## Etapa 2: Processar CSV
                process_csv()

            except Exception as e:
                print(f'Erro ao exportar arquivo: {e}')
                if 'No such file or directory' in str(e):
                    print('Não há arquivo para exportar...')
                    send_keys('%{F4}')
                    time.sleep(2)
                    send_keys('{ENTER}')
                    return

            # Etapa 3: Preparar PDFs
            #app = Application(backend='win32').connect(path=CPJ_PATH) 
            open_search_area(app)        
        
        
        if(etapa_arquivos):
            app = Application(backend='win32').connect(path=CPJ_PATH)

            ## Etapa 2: Processar CSV
            process_csv() 
            open_search_area(app)

            etapa_cpj = True
            
        ## Lê JSON processadoo
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)

        print(f'\nTotal de registros: {data["contagem"]}')
        print('Processando todos os registros...\n')
        
        ## Processa todos os registros
        if etapa_cpj:
            if data['registros']:
                for index, registro in enumerate(data['registros']):
                    process_record_for_pdf(registro, index)

        ## Etapa 4: Preencher planilha modelo
        fill_planilha_modelo_v4()

        ## Etapa 5: Preencher o descritivo e gerar
        fill_descritivo_pdf()

        ## Etapa 6: Abrir Chrome com Selenium
        driver = open_chrome_browser()
        
        # Fazer login
        login_web_exyon_bmg(driver)
        
        # Selecionar unidade
        selecionar_unidade(driver)
        
        # Executar script de menu
        executar_script_menu(driver)
        
        # Executar script do menu lateral
        executar_script_menu_lateral(driver)
        
        # Executar script de classificação financeira
        executar_script_classificacao(driver)

        #Executa preenchimento formulario
        retorno = executar_preenchimento_formulario(driver, data['valor_somado'])

        if(retorno == False):
            print('\n✗ Falha no preenchimento do formulário!')
            pdb.set_trace()
        
        # Anexa os PDFs no formulário
        anexar_pdfs_formulario(driver)
        
        # Verifica se o valor total está correto
        finalizar_processo(driver, data['valor_somado'])
        
        # Aguarda para verificar
        pdb.set_trace()
        
        # Não esqueça de fechar o navegador ao final
        # driver.quit()

        print('\n✓ Automação concluída com sucesso!')
        
    except Exception as e:
        print(f'\n✗ Falha na automação: {e}')
        import traceback
        traceback.print_exc()
        raise


if __name__ == '__main__':
    main()

# Funções CPJ App
import time, os, csv, json, shutil, ctypes, tkinter as tk, pyautogui, pdb, traceback
from datetime import datetime, timedelta
from pywinauto import Application, Desktop
from pywinauto.keyboard import send_keys
from pywinauto.mouse import click
from PyPDF2 import PdfMerger
import xlrd
from xlwt import Workbook
from xlutils.copy import copy as xl_copy
from docx import Document
from docx2pdf import convert
CPJ_PATH=None;LOGIN=None;PASSWORD=None;PLANILHA_PATH=None;CSV_PATH=None;JSON_PATH=None
BASE_DPI=144;BASE_WIDTH=2560;BASE_HEIGHT=1600;IMAGES_PATH=None;PDF_FOLDER=None;PDF_MERGE_PATH=None
PLANILHA_MODELO_PATH=None;PLANILHA_MODELO_PATH_2=None;PLANILHA_MODELO_EXCEL_PATH=None;DOCX_MODELO=None
DOCX_OUTPUT=None;PDF_OUTPUT=None;DATA_INICIAL_PESQUISA=None;DATA_FINAL_PESQUISA=None;NUMERO_RECIBO=None
def set_cpj_config(**config):
    global CPJ_PATH,LOGIN,PASSWORD,PLANILHA_PATH,CSV_PATH,JSON_PATH,BASE_DPI,BASE_WIDTH,BASE_HEIGHT
    global IMAGES_PATH,PDF_FOLDER,PDF_MERGE_PATH,PLANILHA_MODELO_PATH,PLANILHA_MODELO_PATH_2,PLANILHA_MODELO_EXCEL_PATH
    global DOCX_MODELO,DOCX_OUTPUT,PDF_OUTPUT,DATA_INICIAL_PESQUISA,DATA_FINAL_PESQUISA,NUMERO_RECIBO
    for k,v in config.items(): globals()[k.upper()]=v
    print('✓ Config atualizada')
def get_system_dpi():
    user32 = ctypes.windll.user32
    user32.SetProcessDPIAware()
    return user32.GetDpiForSystem()

def get_resolution():
    root = tk.Tk()
    root.withdraw()
    return root.winfo_screenwidth(), root.winfo_screenheight()

def auto_scale_coords(x, y):
    dpi = get_system_dpi()
    width, height = get_resolution()

    dpi_factor = dpi / BASE_DPI
    width_factor = width / BASE_WIDTH
    height_factor = height / BASE_HEIGHT

    final_x = int(x * dpi_factor * width_factor)
    final_y = int(y * dpi_factor * height_factor)

    return final_x, final_y

# wrapper do click
def smart_click(x, y):
    sx, sy = auto_scale_coords(x, y)
    return sx, sy 

def valida_data(data_str: str) -> datetime:
    try:
        return datetime.strptime(data_str, "%d/%m/%Y")
    except ValueError:
        raise ValueError(
            f"Data inválida: {data_str}. Formato obrigatório: dd/mm/YYYY"
        )

def sleep_with_countdown(seconds: int, message: str = 'Aguardando'):
    """Aguarda com countdown visual"""
    for i in range(seconds, 0, -1):
        print(f'\r{message}: {i}s restantes...', end='', flush=True)
        time.sleep(1)
    print(f'\r{message}: concluído!          ')

def click_image(image_path: str, confidence: float = 0.8, timeout: int = 10, description: str = ''):
    """Procura uma imagem na tela e clica no centro dela
    
    Args:
        image_path: Caminho para o arquivo de imagem
        confidence: Confiança mínima para encontrar a imagem (0.0 a 1.0) - requer OpenCV
        timeout: Tempo máximo de espera em segundos
        description: Descrição do elemento (para logs)
    
    Returns:
        bool: True se clicou, False se não encontrou
    """
    if not os.path.exists(image_path):
        print(f'✗ Arquivo de imagem não encontrado: {image_path}')
        return False
    
    desc = description or os.path.basename(image_path)
    print(f'Procurando imagem: {desc}...')
    
    start_time = time.time()
    while time.time() - start_time < timeout:
        try:
            # Tenta com confidence se OpenCV estiver disponível
            try:
                location = pyautogui.locateCenterOnScreen(image_path, confidence=confidence)
            except TypeError:
                # OpenCV não disponível, usa sem confidence
                location = pyautogui.locateCenterOnScreen(image_path)
            
            if location:
                print(f'✓ Imagem encontrada em: {location}')
                pyautogui.click(location)
                time.sleep(0.5)
                return True
        except pyautogui.ImageNotFoundException:
            pass
        except Exception as e:
            # Mostra erro apenas na primeira tentativa
            if time.time() - start_time < 1:
                print(f'Aviso: {e}')
                print('Continuando sem parâmetro confidence...')
        time.sleep(0.5)
    
    print(f'✗ Imagem não encontrada após {timeout}s: {desc}')
    return False

def wait_for_element(app, class_name: str = None, title: str = None, timeout: int = 60, check_interval: float = 0.5):
    """Aguarda até que um elemento apareça na tela
    
    Args:
        app: Objeto Application do pywinauto
        class_name: Nome da classe do elemento (ex: 'TsPanel')
        title: Título do elemento
        timeout: Tempo máximo de espera em segundos
        check_interval: Intervalo entre verificações em segundos
        
    Returns:
        True se encontrou o elemento, False se timeout
    """
    print(f'\nAguardando elemento aparecer...')
    if class_name:
        print(f'  classe: {class_name}')
    if title:
        print(f'  título: {title}')
    print(f'  timeout: {timeout}s')
    
    elapsed = 0
    while elapsed < timeout:
        try:
            dlg = app.top_window()
            
            # Tenta encontrar o elemento
            if class_name and title:
                element = dlg.child_window(class_name=class_name, title=title)
            elif class_name:
                element = dlg.child_window(class_name=class_name)
            elif title:
                element = dlg.child_window(title=title)
            else:
                raise ValueError('Especifique class_name ou title')
            
            # Verifica se existe
            if element.exists():
                print(f'\r✓ Elemento encontrado após {elapsed:.1f}s!          ')
                return True
                
        except Exception:
            pass
        
        # Mostra progresso
        remaining = timeout - elapsed
        print(f'\rProcurando... {remaining:.0f}s restantes', end='', flush=True)
        
        time.sleep(check_interval)
        elapsed += check_interval
    
    print(f'\r✗ Timeout: elemento não encontrado após {timeout}s')
    return False

def clean_planilha_folder():
    """Limpa a pasta planilha"""
    try:
        print('\nLimpando pasta planilha...')
        
        if not os.path.exists(PLANILHA_PATH):
            os.makedirs(PLANILHA_PATH)
            print(f'✓ Pasta criada: {PLANILHA_PATH}')
            return
        
        for file in os.listdir(PLANILHA_PATH):
            file_path = os.path.join(PLANILHA_PATH, file)
            if os.path.isfile(file_path):
                os.unlink(file_path)
                print(f'✓ Arquivo removido: {file}')
        
        print('✓ Pasta planilha limpa com sucesso!')
    except Exception as e:
        print(f'✗ Erro ao limpar pasta planilha: {e}')
        raise

def close_controle_processos_windows():
    """Fecha o aplicativo CPJ se estiver em execução"""
    try:
        print('\nFechando aplicativo CPJ...')
        
        # Tenta conectar ao processo existente
        try:
            app = Application(backend='win32').connect(path=CPJ_PATH, timeout=5)
            print('✓ Aplicativo CPJ encontrado em execução')
            time.sleep(2)
            # Tenta fechar normalmente
            try:
                app.kill()
                print('✓ Aplicativo CPJ fechado com sucesso')
            except Exception as e:
                print(f'⚠ Erro ao fechar: {e}')
                
        except Exception:
            print('✓ Aplicativo CPJ não está em execução')
            
    except Exception as e:
        print(f'✗ Erro ao fechar aplicativo: {e}')

def clean_pdf_merge_folder():
    """Limpa a pasta pdf_merge"""
    try:
        if PDF_MERGE_PATH is None:
            raise ValueError('PDF_MERGE_PATH não foi configurado. Use set_cpj_config()')
        
        merge_folder = PDF_MERGE_PATH
        print(f'\nLimpando pasta pdf_merge: {merge_folder}...')
        
        if not os.path.exists(merge_folder):
            os.makedirs(merge_folder)
            print(f'✓ Pasta criada: {merge_folder}')
            return
        
        for item in os.listdir(merge_folder):
            item_path = os.path.join(merge_folder, item)
            try:
                if os.path.isfile(item_path):
                    os.unlink(item_path)
                    print(f'  ✓ Arquivo removido: {item}')
                elif os.path.isdir(item_path):
                    shutil.rmtree(item_path)
                    print(f'  ✓ Pasta removida: {item}')
            except Exception as e:
                print(f'  ✗ Erro ao remover {item}: {e}')
        
        print('✓ Pasta pdf_merge limpa com sucesso!')
    except Exception as e:
        print(f'✗ Erro ao limpar pasta pdf_merge: {e}')
        raise

def open_cpj_application():
    """Abre o aplicativo CPJ"""
    try:
        print('Abrindo aplicativo CPJ...')
        # Inicia o aplicativo
        app = Application(backend='win32').start(CPJ_PATH)
        print('✓ Aplicativo CPJ iniciado!')
        
        print('Aguardando aplicativo carregar...')
        time.sleep(5)
        
        # Retorna o objeto Application
        return app
        
    except Exception as e:
        print(f'✗ Erro ao abrir aplicativo: {e}')
        # Tenta conectar se já estiver rodando
        try:
            print('Tentando conectar ao processo existente...')
            app = Application(backend='win32').connect(path=CPJ_PATH, timeout=10)
            return app
        except:
            raise

def close_cpj_application(app):
    """Fecha o aplicativo CPJ normalmente"""
    try:
        print('\nFechando aplicativo CPJ...')
        dlg = app.top_window()
        dlg.close()
        print('✓ Aplicativo CPJ fechado!')
    except Exception as e:
        print(f'✗ Erro ao fechar aplicativo: {e}')

def perform_login(app):
    """Realiza o login no sistema"""
    try:
        print('\nIniciando processo de login...')
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Tenta acessar os campos TsEdit diretamente pela classe
        try:
            print('Tentando acessar campos TsEdit...')
            # Busca todos os controles TsEdit
            edits = dlg.children(class_name="TsEdit")
            
            if len(edits) >= 2:
                print(f'✓ Encontrados {len(edits)} campos TsEdit')
                
                # Primeiro campo = Login
                print('Digitando login no primeiro campo...')
                edits[0].set_focus()
                time.sleep(0.2)
                edits[0].set_edit_text(LOGIN)
                time.sleep(0.3)
                
                # Segundo campo = Senha
                print('Digitando senha no segundo campo...')
                edits[1].set_focus()
                time.sleep(0.2)
                edits[1].set_edit_text(PASSWORD)
                time.sleep(0.3)
                
                # O botão Ok é um TsPanel, não um Button padrão
                # Método 1: Tentar encontrar o TsPanel com texto "Ok"
                print('Procurando botão Ok...')
                try:
                    # Método 2: Simplesmente pressionar Enter (mais confiável)
                    print('Usando Enter para confirmar...')
                    dlg.type_keys('{ENTER}')
                    print('✓ Enter pressionado')
                    
                except:

                    ok_button = dlg.child_window(title="Ok", class_name="TsSplitView")
                    ok_button.click()
                    print('✓ Botão Ok clicado (TsSplitView)')
                   
                
            else:
                raise Exception(f'Apenas {len(edits)} campos TsEdit encontrados')
                
        except Exception as e:
            print(f'Método direto falhou: {e}')
            print('Usando método alternativo com type_keys...')
            
            # Fallback: usa digitação sequencial
            dlg.set_focus()
            time.sleep(0.5)
            dlg.type_keys(LOGIN, with_spaces=True)
            time.sleep(0.2)
            dlg.type_keys('{TAB}')
            time.sleep(0.2)
            dlg.type_keys(PASSWORD, with_spaces=True)
            time.sleep(0.2)
            dlg.type_keys('{ENTER}')
        
        print('✓ Login realizado com sucesso!')
        
        # Aguarda até que TsSplitView apareça (indica que o sistema carregou)
        if wait_for_element(app, class_name="TsSplitView", timeout=60):
            print('✓ Sistema carregado!')
        else:
            print('⚠ Timeout ao aguardar sistema carregar, continuando mesmo assim...')
        
        return dlg
        
    except Exception as e:
        print(f'✗ Erro ao realizar login: {e}')
        import traceback
        traceback.print_exc()
        
        # Debug: mostra controles disponíveis
        try:
            print('\n=== DEBUG: Controles disponíveis ===')
            dlg.print_control_identifiers()
        except:
            pass
        raise

def click_menu(app):
    """Clica no menu e navega para CC de clientes/fornecedores"""
    try:
        time.sleep(20)
        print('\nAbrindo menu...')
        
        # Usa Desktop para procurar TsSplitView em todas as janelas
        desktop = Desktop(backend='win32')
        all_windows = desktop.windows()
        
        menu_found = None
        print('Procurando TsSplitView em todas as janelas...')
        
        for window in all_windows:
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsSplitView' and desc.is_visible():
                            print(f'✓ TsSplitView encontrado na janela: "{window.window_text()}"')
                            menu_found = desc
                            break
                    except:
                        pass
                if menu_found:
                    break
            except:
                pass
        click(coords=(38, 31))
        time.sleep(1)
        if not menu_found:
            print('✗ TsSplitView não encontrado!')
            raise Exception('TsSplitView não encontrado')
        
        # Clica no TsSplitView encontrado
        print('Clicando no TsSplitView...')
        menu_found.click_input()
        print('✓ TsSplitView clicado!')
        time.sleep(1)
        
        # Clica na posição específica do item do menu
        print('Clicando em CC de clientes/fornecedores (164, 503)...')
        #click(coords=(164, 503))
        click(coords=smart_click(164,503))
        print('✓ Item clicado!')
        
        # Aguarda a tela carregar
        time.sleep(2)
        print('✓ Menu navegado com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao clicar no menu: {e}')
        import traceback
        traceback.print_exc()
        raise

def click_printer(app):
    """Pressiona F5 para acionar a impressão"""
    try:
        print('\nAguardando tela carregar...')
        time.sleep(2)
        
        # Pressiona F5 (botão de impressão)
        print('Pressionando F5 (botão de impressão)...')
        send_keys('{F5}')
        print('✓ F5 pressionado!')
        time.sleep(2)
        
        return True
            
    except Exception as e:
        print(f'✗ Erro ao listar elementos: {e}')
        import traceback
        traceback.print_exc()
        raise

def navigate_to_analise_lancamento(app):
    """Navega para Análise de Lançamento procurando TsGroupButton"""
    try:
        print('\nNavegando para Análise de Lançamento...')
        time.sleep(2)
        
        # Usa Desktop para procurar em todas as janelas
        desktop = Desktop(backend='win32')
        
        # Procura o TsGroupButton em TODAS as janelas abertas
        print('Procurando TsGroupButton em todas as janelas...')
        
        all_windows = desktop.windows()
        print(f'Total de janelas: {len(all_windows)}')
        
        found_button = None
        
        for window in all_windows:
            try:
                window_text = window.window_text()
                window_class = window.class_name()
                
                # Procura em cada janela
                descendants = window.descendants()
                
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsGroupButton':
                            btn_text = desc.window_text()
                            print(f'  TsGroupButton encontrado: "{btn_text}" na janela "{window_text}" (classe: {window_class})')
                            
                            # Tenta diferentes formas de comparação
                            if ('Análise de lançamentos' in btn_text or 
                                'Analise de lancamentos' in btn_text or
                                'análise de lançamentos' in btn_text.lower()):
                                found_button = desc
                                print(f'✓ Botão MATCH encontrado: "{btn_text}"')
                                print(f'  Janela: {window_text}')
                                print(f'  Classe da janela: {window_class}')
                                print(f'  Retângulo: {desc.rectangle()}')
                                print(f'  Visível: {desc.is_visible()}')
                                print(f'  Habilitado: {desc.is_enabled()}')
                                break
                    except:
                        pass
                
                if found_button:
                    break
                    
            except:
                pass
        
        if not found_button:
            print('✗ TsGroupButton "Análise de lançamentos" não encontrado em nenhuma janela')
            raise Exception('TsGroupButton não encontrado')
        
        # Tenta clicar no botão encontrado
        print('\nTentando clicar no botão...')
        try:
            found_button.click_input()
            print('✓ Botão clicado com click_input()!')
            time.sleep(2)
        except Exception as e:
            print(f'click_input() falhou: {e}')
            try:
                print('Tentando click()...')
                found_button.click()
                print('✓ Botão clicado com click()!')
                time.sleep(2)
            except Exception as e2:
                print(f'click() falhou: {e2}')
                # Último recurso: clica na posição
                rect = found_button.rectangle()
                center_x = (rect.left + rect.right) // 2
                center_y = (rect.top + rect.bottom) // 2
                print(f'Tentando clicar na posição ({center_x}, {center_y})...')
                click(coords=(center_x, center_y))
                print('✓ Clicado na posição!')
                time.sleep(2)
        
        # Agora procura e clica no botão OK (TsBitBtn)
        print('\nProcurando botão OK (TsBitBtn)...')
        
        for window in all_windows:
            try:
                descendants = window.descendants()
                
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsBitBtn':
                            btn_text = desc.window_text()
                            if 'OK' in btn_text or 'Ok' in btn_text:
                                print(f'✓ Botão OK encontrado: "{btn_text}"')
                                print(f'  Retângulo: {desc.rectangle()}')
                                desc.click_input()
                                print('✓ Botão OK clicado!')
                                time.sleep(2)
                                return True
                    except:
                        pass
            except:
                pass
        
        print('⚠ Botão OK não encontrado, mas continuando...')
        return True
        
    except Exception as e:
        print(f'✗ Erro ao navegar para Análise de Lançamento: {e}')
        import traceback
        traceback.print_exc()
        raise

def navigate_to_relatorio_conta(app):
    """Navega para Relatório da Conta"""
    try:
        print('\nNavegando para Relatório da Conta...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Duas setas para baixo
        dlg.type_keys('{DOWN}')
        time.sleep(0.5)
        dlg.type_keys('{DOWN}')
        print('✓ Relatório da Conta selecionado!')
        time.sleep(1)
        
        print('Confirmando com Enter...')
        dlg.type_keys('{ENTER}')
        print('✓ Relatório da Conta confirmado!')
        time.sleep(2)
        
    except Exception as e:
        print(f'✗ Erro ao navegar para Relatório da Conta: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_cliente_and_dates(app):
    """Preenche dados do cliente e datas"""
    try:
        print('\nPreenchendo dados do cliente...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Cliente
        print('Digitando número do cliente (1397)...')
        dlg.type_keys('1397')
        time.sleep(0.5)
        
        # Tab 2x
        print('Navegando para campo de data...')
        dlg.type_keys('{TAB}{TAB}')
        time.sleep(0.5)
        
        # Data de ontem
        #data_ontem = input("Digite a data inicial (DD/MM/AAAA) ou deixe vazio para usar ontem: ")
        #data_ontem = None
        data_ontem = DATA_INICIAL_PESQUISA
        if not data_ontem:
            ontem = datetime.now() - timedelta(days=4)
            data_ontem = ontem.strftime('%d/%m/%Y')
        print(f'Digitando data inicial - ontem ({data_ontem})...')
        dlg.type_keys(data_ontem)
        time.sleep(0.5)
        
        # Tab
        dlg.type_keys('{TAB}')
        time.sleep(0.5)
        
        # Data de hoje
        #data_hoje = input("Digite a data final (DD/MM/AAAA) ou deixe vazio para usar hoje: ")
        #data_hoje = None
        data_hoje = DATA_FINAL_PESQUISA
        if not data_hoje:
            hoje = datetime.now()
            data_hoje = hoje.strftime('%d/%m/%Y')
        print(f'Digitando data final - hoje ({data_hoje})...')
        dlg.type_keys(data_hoje)
        time.sleep(0.5)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        dlg.type_keys('{ENTER}')
        print('✓ Dados preenchidos e confirmados!')
        
        # Aguarda até que o botão Salvar apareça (indica que o relatório foi processado)
        if wait_for_element(app, class_name="TsBitBtn", title="Salvar", timeout=120):
            print('✓ Relatório processado e pronto para exportação!')
        else:
            print('⚠ Timeout ao aguardar relatório processar, continuando mesmo assim...')
        
    except Exception as e:
        print(f'✗ Erro ao preencher dados: {e}')
        import traceback
        traceback.print_exc()
        raise

def export_file(app):
    """Exporta o arquivo CSV"""
    try:
        print('\nIniciando exportação do arquivo...')
        time.sleep(2)
        
        # Pega a janela principal
        dlg = app.top_window()
        dlg.set_focus()
        time.sleep(0.5)
        
        # Clica no botão Salvar
        print('Clicando no botão Salvar...')
        salvar_button = dlg.child_window(class_name="TsBitBtn", title="Salvar")
        salvar_button.click_input()
        print('✓ Botão Salvar clicado!')
        time.sleep(5)
        
        # Digita o caminho do arquivo (janela nativa do Windows)
        print('Digitando caminho do arquivo...')
        send_keys(CSV_PATH)
        print(f'✓ Caminho digitado: {CSV_PATH}')
        time.sleep(1)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        send_keys('{ENTER}')
        print('✓ Exportação confirmada!')
        time.sleep(2)
        
        print('✓ Arquivo exportado com sucesso!')
        time.sleep(2)
        
    except Exception as e:
        print(f'✗ Erro ao exportar arquivo: {e}')
        import traceback
        traceback.print_exc()
        raise

def close_window(app):
    """Fecha a janela atual"""
    try:
        print('\nFechando janela...')
        time.sleep(2)
        
        # Alt+F4 para fechar a janela
        print('Pressionando Alt+F4...')
        send_keys('%{F4}')
        print('✓ Alt+F4 pressionado!')
        time.sleep(2)
        
        # Enter para confirmar
        print('Confirmando com Enter...')
        send_keys('{ENTER}')
        print('✓ Enter pressionado!')
        time.sleep(2)
        
        print('✓ Janela fechada com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao fechar janela: {e}')
        import traceback
        traceback.print_exc()
        raise

def to_snake_case(text: str) -> str:
    """Converte string para snake_case"""
    import unicodedata
    
    # Valida entrada
    if text is None or not isinstance(text, str):
        return ''
    
    # Remove acentos
    text = unicodedata.normalize('NFD', text)
    text = ''.join(char for char in text if unicodedata.category(char) != 'Mn')
    
    # Converte para lowercase e substitui
    text = text.lower()
    text = text.replace('.', '_').replace(' ', '_')
    text = ''.join(c if c.isalnum() or c == '_' else '' for c in text)
    text = '_'.join(filter(None, text.split('_')))
    
    return text

def process_csv():
    """Processa o CSV e gera JSON"""
    try:
        print('\nProcessando arquivo CSV...')

        # Lê o CSV
        with open(CSV_PATH, 'r', encoding='latin1') as f:
            reader = csv.DictReader(f, delimiter=';')
            records = list(reader)
        #pdb.set_trace()
        if len(records) == 0:
            print('⚠ Aviso: CSV está vazio!')

            # Lê o CSV
            with open(CSV_PATH, 'r', encoding='latin1') as f:
                reader = csv.DictReader(f, delimiter=',')
                records = list(reader)

        print(f'✓ CSV lido: {len(records)} registros encontrados')
        
        # Filtra registros
        filtered_records = []
        registros_filtrados_correios = 0
        
        for record in records:

            historico = record.get('Histórico', '')
            valor_credito_conta = record.get('Crédito na moeda','0')

            # Filtra registros com valor 0,00
            if '0,00' not in valor_credito_conta.strip() and '0' not in valor_credito_conta.strip():
                continue
            
            # Filtra registros que contenham "CORREIOS" no histórico
            if 'CORREIOS' in historico.upper():
                registros_filtrados_correios += 1
                continue

            filtered_records.append(record)
        
        print(f'✓ Registros filtrados: {len(filtered_records)} registros')
        
        # Calcula soma e converte para snake_case
        total_valor = 0
        registros_snake_case = []
        
        for record in filtered_records:
            valor_str = record.get('Valor em R$', '0')
            valor = float(valor_str.strip().replace(',', '.'))
            total_valor += valor
            
            # Converte chaves para snake_case
            record_snake = {}
            for key, value in record.items():
                snake_key = to_snake_case(key)
                
                # Remove caracteres especiais do número do processo
                if snake_key == 'pro_numero_do_processo':
                    value = ''.join(c for c in value if c.isdigit())
                
                record_snake[snake_key] = value
            
            registros_snake_case.append(record_snake)
        
        # Formata valor
        valor_formatado = f"{total_valor:.2f}".replace('.', ',')
        
        # Cria objeto final
        resultado = {
            'contagem': len(filtered_records),
            'valor_somado': valor_formatado,
            'registros': registros_snake_case
        }
        
        print(f'✓ Contagem de registros: {resultado["contagem"]}')
        print(f'✓ Valor total somado: R$ {resultado["valor_somado"]}')
        
        # Salva JSON
        with open(JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump(resultado, f, indent=2, ensure_ascii=False)
        
        print(f'✓ JSON salvo em: {JSON_PATH}')
        return resultado
        
    except Exception as e:
        print(f'✗ Erro ao processar CSV: {e}')
        raise

def open_search_area(app):
    """Abre a área de busca"""
    try:
        print('\nAbrindo área de busca...')
        
        # Clique 1
        click(coords=smart_click(91, 270))
        print('✓ Clicado em (91, 270)')
        time.sleep(2)
        
        # Usa Desktop para procurar o TsComboBox em todas as janelas
        print('Procurando TsComboBox...')
        desktop = Desktop(backend='win32')
        
        combo_found = None
        for window in desktop.windows():
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsComboBox':
                            print(f'✓ TsComboBox encontrado na janela: "{window.window_text()}"')
                            combo_found = desc
                            break
                    except:
                        pass
                if combo_found:
                    break
            except:
                pass
        
        if combo_found:
            print('Clicando no TsComboBox...')
            combo_found.click_input()
            print('✓ TsComboBox clicado!')
            time.sleep(0.5)
        else:
            print('⚠ TsComboBox não encontrado, continuando...')
        
        # Pressiona Ctrl+I
        # print('Pressionando Ctrl+I...')
        # send_keys('^i')
        # print('✓ Ctrl+I pressionado!')
        # time.sleep(1)

        # Pressiona Ctrl+I
        print('Pressionando Ctrl+4...')
        send_keys('^4')
        print('✓ Ctrl+4 pressionado!')
        time.sleep(1)
        
        print('✓ Área de busca aberta!')
        
    except Exception as e:
        print(f'✗ Erro ao abrir área de busca: {e}')
        raise

def fill_planilha_modelo():
    """Preenche a planilha modelo com dados do JSON"""
    try:
        print('\nPreenchendo planilha modelo...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Abre a planilha existente para ler o cabeçalho
        workbook_read = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        sheet_read = workbook_read.sheet_by_index(0)
        
        # Cria nova planilha
        workbook = Workbook()
        sheet = workbook.add_sheet('Plan1')
        
        # Copia o cabeçalho (primeira linha)
        print('Copiando cabeçalho...')
        for col_idx in range(sheet_read.ncols):
            value = sheet_read.cell_value(0, col_idx)
            sheet.write(0, col_idx, value)
        
        print('✓ Cabeçalho copiado')

        # Fecha o workbook de leitura para liberar o arquivo
        del workbook_read
        del sheet_read
        
        # Preenche os dados a partir da linha 1 (linha 0 é o cabeçalho)
        print('Preenchendo dados...')
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):
            integracao = registro.get('pro_numero_de_integracao', '')
            processo = registro.get('pro_numero_do_processo', '')
            valor = registro.get('debito_na_moeda', '')
            
            # Coluna A: pro_numero_de_integracao
            sheet.write(idx, 0, integracao)
            
            # Coluna B: pro_numero_do_processo
            sheet.write(idx, 1, processo)
            
            # Coluna C: texto fixo
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            
            # Coluna D: valor
            sheet.write(idx, 3, valor)
            
            # Coluna E: data de hoje
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] {integracao} - {processo} - {valor}')
        
        # Adiciona linha final com valor somado na coluna D
        # ultima_linha = len(data['registros']) + 1
        # valor_somado = data.get('valor_somado', '')
        # sheet.write(ultima_linha, 3, valor_somado)
        # print(f'\n✓ Valor total na linha {ultima_linha + 1}: {valor_somado}')
        
        # Salva a planilha
        workbook.save(PLANILHA_MODELO_PATH)
        print(f'✓ Planilha salva em: {PLANILHA_MODELO_PATH}')
        print(f'✓ Total de registros preenchidos: {len(data["registros"])}')
        
        # Libera explicitamente os objetos e força coleta de lixo
        del workbook
        del sheet
        import gc
        gc.collect()
        
        # Aguarda um pouco para garantir que o arquivo foi liberado pelo sistema operacional
        time.sleep(1)
        print('✓ Arquivo fechado e liberado')

    except Exception as e:
        print(f'✗ Erro ao preencher planilha modelo: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_planilha_modelo_v2():
    """Preenche a planilha modelo com dados do JSON - Versão 2
    
    Copia a planilha modelo do Excel para a pasta de destino antes de preencher
    """
    try:
        print('\nPreenchendo planilha modelo (v2)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        if PLANILHA_MODELO_EXCEL_PATH is None:
            raise ValueError('PLANILHA_MODELO_EXCEL_PATH não foi configurado. Use set_cpj_config()')
        
        planilha_excel_path = PLANILHA_MODELO_EXCEL_PATH
        
        # Verifica se a planilha modelo existe
        if not os.path.exists(planilha_excel_path):
            raise FileNotFoundError(f'Planilha modelo não encontrada: {planilha_excel_path}')
        
        print(f'✓ Planilha modelo encontrada: {planilha_excel_path}')
        
        # Cria a pasta de destino se não existir
        destino_folder = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), 
            'planilha_modelo', 
            'bmg'
        )
        os.makedirs(destino_folder, exist_ok=True)
        
        # Copia a planilha modelo para o destino
        print(f'Copiando planilha para: {PLANILHA_MODELO_PATH}')
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada com sucesso!')
        
        # Aguarda um pouco para garantir que o arquivo foi copiado
        time.sleep(0.5)
        
        # Abre a planilha copiada para ler o cabeçalho
        workbook_read = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        sheet_read = workbook_read.sheet_by_index(0)
        
        # Cria nova planilha
        workbook = Workbook()
        sheet = workbook.add_sheet('Plan1')
        
        # # Copia o cabeçalho (primeira linha)
        # print('Copiando cabeçalho...')
        # for col_idx in range(sheet_read.ncols):
        #     value = sheet_read.cell_value(0, col_idx)
        #     sheet.write(0, col_idx, value)
        
        # print('✓ Cabeçalho copiado')

        # Fecha o workbook de leitura para liberar o arquivo
        del workbook_read
        del sheet_read
        
        # Preenche os dados a partir da linha 1 (linha 0 é o cabeçalho)
        print('Preenchendo dados...')
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):
            integracao = registro.get('pro_numero_de_integracao', '')
            processo = registro.get('pro_numero_do_processo', '')
            valor = registro.get('debito_na_moeda', '')
            
            # Coluna A: pro_numero_de_integracao
            sheet.write(idx, 0, integracao)
            
            # Coluna B: pro_numero_do_processo
            sheet.write(idx, 1, processo)
            
            # Coluna C: texto fixo
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            
            # Coluna D: valor
            sheet.write(idx, 3, valor)
            
            # Coluna E: data de hoje
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] {integracao} - {processo} - {valor}')
        #pdb.set_trace()
        # Salva a planilha
        #workbook.save(PLANILHA_MODELO_PATH)
        print(f'✓ Planilha salva em: {PLANILHA_MODELO_PATH}')
        print(f'✓ Total de registros preenchidos: {len(data["registros"])}')
        
        # Libera explicitamente os objetos e força coleta de lixo
        del workbook
        del sheet
        import gc
        gc.collect()
        
        # Aguarda um pouco para garantir que o arquivo foi liberado pelo sistema operacional
        time.sleep(1)
        print('✓ Arquivo fechado e liberado')

    except Exception as e:
        print(f'✗ Erro ao preencher planilha modelo (v2): {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_planilha_modelo_v4():
    """Preenche a planilha modelo preservando formatação - Versão 4"""
    try: 
        print('\nPreenchendo planilha modelo (v4 - xlutils)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        if PLANILHA_MODELO_EXCEL_PATH is None:
            raise ValueError('PLANILHA_MODELO_EXCEL_PATH não foi configurado. Use set_cpj_config()')
        
        planilha_excel_path = PLANILHA_MODELO_EXCEL_PATH
        
        # Copia a planilha para o destino
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada!')
        
        # Abre a planilha original
        rb = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        
        # Cria uma cópia que preserva formatação
        wb = xl_copy(rb)
        sheet = wb.get_sheet(0)
        
        # Preenche os dados
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):  # Linha 1 (pula cabeçalho na linha 0)
            integracao = registro.get('pro_numero_de_integracao', '')
            integracao = integracao.split('.')[0] if '.' in integracao else integracao
            sheet.write(idx, 0, integracao)
            sheet.write(idx, 1, registro.get('pro_numero_do_processo', ''))
            sheet.write(idx, 2, registro.get('historico', 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU'))
            sheet.write(idx, 3, registro.get('debito_na_moeda', ''))
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] Linha preenchida')
        
        # Salva (preserva GRANDE PARTE da formatação)
        wb.save(PLANILHA_MODELO_PATH)
        
        print('✓ Planilha salva com formatação preservada!')
        
        del wb
        del rb
        import gc
        gc.collect()
        time.sleep(1)
        
    except Exception as e:
        print(f'✗ Erro: {e}')
        traceback.print_exc()
        raise

def fill_planilha_modelo_v5():
    """Preenche a planilha modelo preservando formatação - Versão 4"""
    try: 
        print('\nPreenchendo planilha modelo (v4 - xlutils)...')
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        
        # Caminho da planilha modelo no Excel
        if PLANILHA_MODELO_EXCEL_PATH is None:
            raise ValueError('PLANILHA_MODELO_EXCEL_PATH não foi configurado. Use set_cpj_config()')
        
        planilha_excel_path = PLANILHA_MODELO_EXCEL_PATH
        
        # Copia a planilha para o destino
        shutil.copy2(planilha_excel_path, PLANILHA_MODELO_PATH)
        print('✓ Planilha copiada!')
        
        # IMPORTANTE: Usar context manager implícito ao não guardar referência
        # Abre a planilha original
        rb = xlrd.open_workbook(PLANILHA_MODELO_PATH, formatting_info=True)
        
        # Cria uma cópia que preserva formatação
        wb = xl_copy(rb)
        sheet = wb.get_sheet(0)
        
        # Preenche os dados
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        
        for idx, registro in enumerate(data['registros'], start=1):  # Linha 1 (pula cabeçalho na linha 0)
            sheet.write(idx, 0, registro.get('pro_numero_de_integracao', ''))
            sheet.write(idx, 1, registro.get('pro_numero_do_processo', ''))
            sheet.write(idx, 2, 'CUSTAS DIVERSAS E TAXAS JUDICIAIS / TRIBUNAL DE JU')
            sheet.write(idx, 3, registro.get('debito_na_moeda', ''))
            sheet.write(idx, 4, data_hoje)
            
            print(f'  [{idx}] Linha preenchida')
        
        # Salva (preserva GRANDE PARTE da formatação)
        wb.save(PLANILHA_MODELO_PATH)
        
        print('✓ Planilha salva com formatação preservada!')
        
        # FORÇA FECHAMENTO EXPLÍCITO
        # Fecha todos os handles
        try:
            if hasattr(wb, 'close'):
                wb.close()
        except:
            pass
        
        try:
            if hasattr(rb, 'release_resources'):
                rb.release_resources()
        except:
            pass
        
        # Remove referências
        del sheet
        del wb
        del rb
        
        # Força coleta de lixo agressiva
        import gc
        gc.collect()
        gc.collect()  # Duas vezes para garantir
        
        # Aguarda mais tempo para o SO liberar o arquivo
        print('Aguardando liberação do arquivo...')
        time.sleep(2)  # Aumentei de 1 para 2 segundos
        
        # VALIDAÇÃO: Tenta abrir e fechar o arquivo para garantir que está liberado
        try:
            with open(PLANILHA_MODELO_PATH, 'rb') as test_file:
                test_file.read(1)
            print('✓ Arquivo validado e totalmente liberado!')
        except Exception as e:
            print(f'⚠ Aviso: Arquivo ainda pode estar travado: {e}')
            time.sleep(2)  # Aguarda mais um pouco
        
    except Exception as e:
        print(f'✗ Erro: {e}')
        traceback.print_exc()
        raise

def process_record_for_pdf(record: dict, index: int, primeira_procura: bool = False):
    """Processa um registro para gerar PDF"""
    try:
        integracao = record['pro_numero_de_integracao']
        n_processo = record['pro_numero_do_processo']

        #or 'CIV1428248' not in integracao
        # if 'CIV1468595' not in integracao:
        #     print(f'\n[{index + 1}] Ignorando registro.')
        #     return

        #print(f'\n[{index + 1}] Processando integração: {integracao}')
        print(f'\n[{index + 1}] Processando numero processo: {n_processo}')
        
        # Procura e clica no campo de busca TsEdit
        print('Procurando campos TsEdit...')
        desktop = Desktop(backend='win32')

        # Procurar TsComboBox habilitados
        if index == 0:
            selects_found = []
            for window in desktop.windows():
                try:
                    descendants = window.descendants()
                    for desc in descendants:
                        try:
                            if desc.class_name() == 'TsComboBox' and desc.is_enabled() == True and desc.is_visible() == True:
                                selects_found.append(desc)
                        except:
                            pass
                except:
                    pass
            #pdb.set_trace()
            print(f'✓ Total de TsComboBox encontrados: {len(selects_found)}')
            selects_found[0].click_input()
            time.sleep(0.5)
            print('✓ TsComboBox clicado!')
            send_keys('{DOWN}')
            time.sleep(0.5)
            send_keys('{DOWN}')
            time.sleep(0.5)
            send_keys('{ENTER}')
            time.sleep(0.5)

            #desmarca o segundo checkbox TsCheckBox habilitado
            checks_found = []
            for window in desktop.windows():
                try:
                    descendants = window.descendants()
                    for desc in descendants:
                        try:
                            if desc.class_name() == 'TsCheckBox' and desc.is_visible() == True and desc.is_enabled() == True:
                                checks_found.append(desc)
                        except:
                            pass
                except:
                    pass

            if len(checks_found) >= 2:
                # Usa o segundo TsCheckBox visível (índice 1)
                check_found = checks_found[1]
                print(f'Usando o TsCheckBox visível na posição 1...')
                check_found.click_input()
                print('✓ TsCheckBox clicado!')
                time.sleep(0.5)
        
        edits_found = []
        for window in desktop.windows():
            try:
                descendants = window.descendants()
                for desc in descendants:
                    try:
                        if desc.class_name() == 'TsEdit':
                            edits_found.append(desc)
                    except:
                        pass
            except:
                pass
        
        print(f'✓ Total de TsEdit encontrados: {len(edits_found)}')

        edits_visible = [e for e in edits_found if e.is_visible()]
        print(f'✓ TsEdit visíveis: {len(edits_visible)}')

        for i, edit in enumerate(edits_visible):
            rect = edit.rectangle()
            print(f'  [{i}] Pos: ({rect.left}, {rect.top}), Tamanho: {rect.width()}x{rect.height()}')

        if len(edits_visible) >= 2:
            # Usa o segundo TsEdit visível (índice 1)
            edit_found = edits_visible[2]
            print(f'Usando o TsEdit visível na posição 1...')
            edit_found.click_input()
            print('✓ TsEdit clicado!')
            time.sleep(0.5)
            
            # Garante foco no elemento
            edit_found.set_focus()
            time.sleep(0.3)
            
            # Limpa o campo antes de digitar
            print('Limpando campo...')
            send_keys('^a')  # Ctrl+A para selecionar tudo
            send_keys('{DELETE}')  # Delete para limpar
            time.sleep(0.3)
            
            # Digita integração usando send_keys
            # print(f'Digitando integração: {integracao}')
            # send_keys(integracao)
            # print('✓ Integração digitada!')
            # time.sleep(0.5)

            # Digita integração usando send_keys
            print(f'Digitando numero processo: {n_processo}')
            send_keys(n_processo)
            print('✓ Numero processo digitado!')
            time.sleep(0.5)

        # Enter para buscar
        print('Pressionando Enter para buscar...')
        send_keys('{ENTER}')
        print('✓ Busca realizada')
        time.sleep(1)
        
        # Aguarda até que TsDBCtrlGrid tenha elementos
        print('Aguardando resultados da busca...')
        max_wait = 30  # máximo 30 segundos
        elapsed = 0
        found_results = False
        grid_element = None
        
        while elapsed < max_wait:
            try:
                desktop = Desktop(backend='win32')
                for window in desktop.windows():
                    try:
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TsDBCtrlGrid':
                                    # Verifica se o grid tem filhos (resultados)
                                    children = desc.children()
                                    if len(children) == 1:
                                        print(f'✓ Resultados encontrados! Total de elementos no grid: {len(children)}')
                                        found_results = True
                                        grid_element = desc
                                        break
                            except:
                                pass
                        if found_results:
                            break
                    except:
                        pass
                if found_results:
                    break
            except:
                pass
            
            remaining = max_wait - elapsed
            print(f'\rAguardando resultados... {remaining}s restantes', end='', flush=True)
            time.sleep(1)
            elapsed += 1

        if not found_results:
            print(f'\r⚠ Timeout: Nenhum resultado encontrado após {max_wait}s')
        else:
            print('\r✓ Resultados carregados com sucesso!              ')
            time.sleep(6)
            
            # Pressiona Enter para abrir o resultado encontrado
            print('Pressionando Enter para abrir o resultado...')
            if primeira_procura == False:
                send_keys('{DOWN}')

            time.sleep(3)
            send_keys('{ENTER}')
            print('✓ Enter pressionado!')
            
            # Aguarda até que TsPanel fique visível
            print('Aguardando TsPanel ficar visível...')
            max_wait_panel = 30
            elapsed_panel = 0
            panel_visible = False
            panel_visible_lembrete = False
            
            while elapsed_panel < max_wait_panel:
                try:
                    desktop = Desktop(backend='win32')
                    
                    # Verifica se existe janela "Cadastro de Processos - Lembretes"
                    for window in desktop.windows():
                        try:
                            #print(f'Verificando janela: {window.window_text()}')
                            if 'Cadastro de Processos - Lembretes' in window.window_text() and window.is_visible():
                                print('✓ Janela "Cadastro de Processos - Lembretes" encontrada, pressionando F4...')
                                send_keys('{F4}')
                                time.sleep(1)
                                break
                        except:
                            pass
                    
                    for window in desktop.windows():
                        try:
                            descendants = window.descendants()
                            for desc in descendants:
                                try:
                                    #print(f'Verificando elemento: {desc.class_name()} - Visível: {desc.is_visible()}')
                                    if desc.class_name() == 'TsPanel' and desc.is_visible():
                                        print(f'✓ TsPanel visível encontrado!')
                                        panel_visible = True
                                        break
                                except:
                                    pass
                            if panel_visible:
                                break
                        except:
                            pass
                    if panel_visible:
                        break
                except:
                    pass
                
                remaining_panel = max_wait_panel - elapsed_panel
                print(f'\rAguardando TsPanel... {remaining_panel}s restantes', end='', flush=True)
                time.sleep(1)
                elapsed_panel += 1

            if panel_visible:
                print('\r✓ TsPanel visível! Executando Ctrl+L e digitando SPF...              ')
                send_keys('^l')
                time.sleep(0.5)
                send_keys('SPF')
                print('✓ Ctrl+L pressionado e SPF digitado no TsPanel!')
                time.sleep(1)
                send_keys('{ENTER}')
                print('✓ Enter pressionado após digitar SPF')
                time.sleep(3)  # Aguarda mais tempo para a janela SPF carregar
                send_keys('{ENTER}')
                print('✓ Enter pressionado após digitar SPF')
                time.sleep(3) 
                
                # Procura e clica no elemento Grid (TPage) da janela SPF
                print('Procurando elemento TPage em todas as janelas...')
                desktop = Desktop(backend='win32')
                
                # Lista todos os TPage encontrados
                tpages_found = []
                for window in desktop.windows():
                    try:
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TPage' and desc.is_visible():
                                    window_title = window.window_text()
                                    desc_name = desc.window_text()
                                    rect = desc.rectangle()
                                    tpages_found.append({
                                        'element': desc,
                                        'name': desc_name,
                                        'window': window_title,
                                        'rect': rect
                                    })
                                    print(f'  TPage visível: "{desc_name}" na janela "{window_title}" em {rect}')
                            except:
                                pass
                    except:
                        pass
                
                print(f'✓ Total de TPage visíveis: {len(tpages_found)}')
                
                # Garante que a janela SPF está em foco
                print('\nColocando janela SPF em foco...')
                spf_window = None
                for window in desktop.windows():
                    try:
                        window_title = window.window_text()
                        if 'SPF' in window_title and 'Solicitação' in window_title:
                            spf_window = window
                            print(f'✓ Janela SPF encontrada: "{window_title}"')
                            break
                    except:
                        pass
                
                if spf_window is None:
                    print('⚠ Janela SPF não encontrada!')
                    send_keys('%{F4}')  # Alt+F4 para fechar
                    time.sleep(2)
                    send_keys('{RIGHT}')  # Right para nao fechar
                    time.sleep(1)
                    send_keys('{ENTER}')  # Enter para confirmar
                    return process_record_for_pdf(record, index, primeira_procura=True)

                if spf_window:
                    try:
                        # Traz a janela para frente e dá foco
                        spf_window.set_focus()
                        print('✓ Janela SPF em foco')
                        time.sleep(1)
                    except Exception as e:
                        print(f'⚠ Erro ao dar foco: {e}')
                    
                    # Procura TsComboBox "Ambas" diretamente
                    print('\nProcurando TsComboBox "Ambas"...')
                    descendants = spf_window.descendants()
                    elemento_encontrado = None
                    
                    for desc in descendants:
                        try:
                            if desc.is_visible() and desc.class_name() == 'TsComboBox':
                                desc_name = desc.window_text()
                                if desc_name == 'Ambas':
                                    rect = desc.rectangle()
                                    # Posiciona no início do elemento (canto superior esquerdo)
                                    start_x = rect.left
                                    start_y = rect.top + 55 
                                    
                                    print(f'✓ TsComboBox "Ambas" encontrado em ({start_x}, {start_y})')
                                    print('Posicionando mouse no início do elemento...')
                                    pyautogui.moveTo(start_x, start_y)
                                    time.sleep(2)
                                    pyautogui.click(start_x, start_y)
                                    # Clica no elemento
                                    print('Clicando no elemento...')
                                    print('✓ Clicado com pyautogui!')
                                    time.sleep(1)
                                    break
                        except:
                            pass

        print('Pressionando seta para baixo...')
        send_keys('{DOWN}')
        time.sleep(0.5)
        send_keys('{DOWN}')
        time.sleep(0.5)
        
        print('Pressionando Enter para confirmar...')                     
        send_keys('{ENTER}')
        time.sleep(2)
        
        print('Digitando Downloads...')
        send_keys('Downloads')
        time.sleep(0.5)
        
        print('Pressionando seta para direita...')
        send_keys('{RIGHT}')
        time.sleep(0.5)
        
        print('Digitando pdf_cpj...')
        send_keys('pdf_cpj')
        time.sleep(0.5)
        
        print('Pressionando Enter...')
        send_keys('{ENTER}')

        # Obtém o nome do usuário atual do Windows

        print(f'\n✓ PDFs sendo salvos em: {PDF_FOLDER}')
        
        # Aguarda e clica no botão OK da janela de confirmação
        print('\nProcurando botão OK...')
        desktop = Desktop(backend='win32')
        botao_ok = False

        while botao_ok == False:
            time.sleep(5)
            for window in desktop.windows():
                try:
                    window_title = window.window_text()
                    if 'Controle de Processos Jurídicos' in window_title and window.is_visible():
                        #print(f'✓ Janela encontrada: "{window_title}"')
                        
                        # Procura o botão OK
                        descendants = window.descendants()
                        for desc in descendants:
                            try:
                                if desc.class_name() == 'TButton' and desc.window_text() == 'OK':
                                    print('Clicando no botão OK...')
                                    desc.click_input()
                                    print('✓ Botão OK clicado!')
                                    time.sleep(1)
                                    botao_ok = True
                                    break
                            except:
                                botao_ok = False
                                pass
                        break
                except:
                    botao_ok = False
                    pass
        
        # Lista e faz merge dos PDFs
        print(f'\nListando PDFs em: {PDF_FOLDER}')
        
        if os.path.exists(PDF_FOLDER):
            # Busca recursivamente por PDFs em todas as subpastas
            pdf_files = []
            for root, dirs, files in os.walk(PDF_FOLDER):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        pdf_path = os.path.join(root, file)
                        pdf_files.append(pdf_path)
            
            # Ordena os PDFs por caminho
            pdf_files.sort()
            
            print(f'✓ Total de PDFs encontrados: {len(pdf_files)}')
            
            for pdf in pdf_files:
                print(f'  - {pdf}')
            
            if pdf_files:
                # Cria pasta pdf_merge se não existir
                if PDF_MERGE_PATH is None:
                    raise ValueError('PDF_MERGE_PATH não foi configurado. Use set_cpj_config()')
                
                merge_folder = PDF_MERGE_PATH
                os.makedirs(merge_folder, exist_ok=True)
                print(f'✓ Pasta de merge: {merge_folder}')
                
                # Constrói nome do arquivo com trim
                integracao = record['pro_numero_de_integracao'].strip()
                processo = record['pro_numero_do_processo'].strip()
                debito = record['debito_na_moeda'].strip()
                
                merged_filename = f'{integracao}_{processo}_{debito}.pdf'
                merged_path = os.path.join(merge_folder, merged_filename)
                
                print(f'\nFazendo merge dos PDFs...')
                print(f'Nome do arquivo: {merged_filename}')
                
                # Faz o merge
                merger = PdfMerger()
                
                for pdf_path in pdf_files:
                    merger.append(pdf_path)
                    print(f'  + {os.path.basename(pdf_path)}')
                
                # Salva o PDF mesclado
                merger.write(merged_path)
                merger.close()
                
                print(f'✓ PDF mesclado salvo em: {merged_path}')
                
                # Deleta todo o conteúdo de PDF_FOLDER
                print(f'\nLimpando pasta {PDF_FOLDER}...')
                
                for item in os.listdir(PDF_FOLDER):
                    item_path = os.path.join(PDF_FOLDER, item)
                    try:
                        if os.path.isfile(item_path):
                            os.unlink(item_path)
                            print(f'  ✓ Arquivo deletado: {item}')
                        elif os.path.isdir(item_path):
                            shutil.rmtree(item_path)
                            print(f'  ✓ Pasta deletada: {item}')
                    except Exception as e:
                        print(f'  ✗ Erro ao deletar {item}: {e}')
                
                print(f'✓ Pasta limpa com sucesso!')
            else:
                print('⚠ Nenhum PDF encontrado para fazer merge')
        else:
            print(f'✗ Pasta não encontrada: {PDF_FOLDER}')
        
        # Pressiona F4 duas vezes para finalizar
        print('\nFinalizando processo...')
        print('Pressionando F4...')
        send_keys('{F4}')
        time.sleep(1)
        print('Pressionando F4 novamente...')
        send_keys('{F4}')
        time.sleep(1)
        print('✓ Processo finalizado!')

        print(f'\n✓ Registro {index + 1} processado com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao processar registro {index + 1}: {e}')
        import traceback
        traceback.print_exc()
        raise

def fill_descritivo_pdf():
    """Preenche o documento descritivo e converte para PDF"""
    try:
        print('\nPreenchendo documento descritivo...')
        
        # Abre o documento modelo
        doc = Document(DOCX_MODELO)
        
        # Lê o JSON
        with open(JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Dados para substituição
        data_hoje = datetime.now().strftime('%d/%m/%Y')
        mes_hoje = datetime.now().strftime('%m/%Y')
        valor_somado = data['valor_somado']
        
        print(f'✓ JSON lido: {data["contagem"]} registros')
        print(f'✓ Valor somado: R$ {valor_somado}')
        print(f'✓ Número recibo: {NUMERO_RECIBO}')
        
        # Substitui placeholders nas tabelas
        if len(doc.tables) > 0:
            for table_idx, table in enumerate(doc.tables):
                print(f'Processando tabela {table_idx + 1}...')
                
                # Percorre todas as células da tabela
                for row in table.rows:
                    for cell in row.cells:
                        # Substitui placeholders no texto de cada célula
                        if r'{{NUMERO_RECIBO}}' in cell.text:
                            cell.text = cell.text.replace(r'{{NUMERO_RECIBO}}', str(NUMERO_RECIBO))
                        if r'{{DATA_HOJE}}' in cell.text:
                            cell.text = cell.text.replace(r'{{DATA_HOJE}}', data_hoje)
                        if r'{{MES_HOJE}}' in cell.text:
                            cell.text = cell.text.replace(r'{{MES_HOJE}}', mes_hoje)
                        if r'{{VALOR_SOMADO}}' in cell.text:
                            cell.text = cell.text.replace(r'{{VALOR_SOMADO}}', valor_somado)
            

            # Preenche a tabela com os registros (geralmente é a terceira tabela - índice 2)
            if len(doc.tables) > 2:
                table = doc.tables[2]
                print(f'Preenchendo registros na tabela 3...')
                # pdb.set_trace()
                # for idx, registro in enumerate(data['registros']):
                #     # Adiciona nova linha
                #     row = table.add_row()
                    
                #     # Preenche as células
                #     row.cells[0].text = str(idx + 1)  # Número sequencial
                #     row.cells[1].text = registro.get('historico', '')  # Histórico
                #     row.cells[2].text = registro.get('debito_na_moeda', '')  # Valor
                    
                #     print(f'  [{idx + 1}] {registro.get("historico", "")}')

                for idx, registro in enumerate(data['registros']):
                    row_idx = idx + 1

                    # Adiciona nova linha se necessário
                    if row_idx >= len(table.rows):
                        table.add_row()
                    
                    row = table.rows[row_idx]
                    row.cells[0].text = 'Custas'
                    row.cells[1].text = datetime.now().strftime('%d/%m/%Y')
                    row.cells[2].text = registro['pro_numero_de_integracao']
                    row.cells[3].text = registro['historico'].strip()
                    row.cells[4].text = registro['debito_na_moeda'].strip()
        
        # Salva o documento Word
        doc.save(DOCX_OUTPUT)
        print(f'✓ Documento Word salvo em: {DOCX_OUTPUT}')
        
        # Converte para PDF
        print('Convertendo para PDF...')
        convert(DOCX_OUTPUT, PDF_OUTPUT)
        print(f'✓ PDF salvo em: {PDF_OUTPUT}')
        print('✓ Descritivo preenchido com sucesso!')
        
    except Exception as e:
        print(f'✗ Erro ao preencher descritivo: {e}')
        import traceback
        traceback.print_exc()
        raise

